"""
Suite H: Parser (chunk_method) coverage with end-to-end retrieval + QA.

For every ragflow chunk_method we can feed with the available test data, this
suite:
  1. creates a KB with that chunk_method (parser_id),
  2. uploads N docs whose FORMAT and CONTENT match what the parser expects,
  3. parses them and verifies chunks were produced,
  4. for QA-targetable parsers, asks the encoded questions through a chat
     assistant and checks the answer SEMANTICALLY (LLM-as-judge over the
     configured chat model), confirming retrieval + generation actually work.

Facts come from generate_kb_test_data.TOPICS[*]["qa_pairs"], so the same
question/answer set is encoded into every format and can be asked uniformly.

Env knobs:
  H_DOCS_PER_PARSER   docs uploaded per parser KB (default 10; >=3 recommended)
  H_PARSERS           comma list to restrict parsers (default: all feasible)
  H_QA_QUESTIONS      questions asked per QA-targetable parser (default 3)

audio is intentionally excluded: no speech2text model is configured/available
(ragflow 0.18 only ships whisper-1 under OpenAI/Azure; SILICONFLOW exposes TTS
only). Add an ASR model + set the tenant asr_id to cover it.
"""
import os
import sys
import time
import tempfile
from io import BytesIO
from pathlib import Path

import pytest
import requests

sys.path.insert(0, "/tests")
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from generate_kb_test_data import TOPICS  # noqa: E402

RAGFLOW_URL = os.environ.get("RAGFLOW_BASE_URL", "http://localhost:9380").rstrip("/")
RAGFLOW_API_KEY = os.environ.get("RAGFLOW_API_KEY", "")
HEADERS = {"Authorization": f"Bearer {RAGFLOW_API_KEY}", "Content-Type": "application/json"}
UPLOAD_HEADERS = {"Authorization": f"Bearer {RAGFLOW_API_KEY}"}

DOCS_PER_PARSER = int(os.environ.get("H_DOCS_PER_PARSER", "10"))
QA_QUESTIONS = int(os.environ.get("H_QA_QUESTIONS", "3"))
CHAT_MODEL = os.environ.get("H_CHAT_MODEL", "deepseek-chat")

# parser_id -> (gen_kind, qa_targetable)
# gen_kind selects the data generator below; qa_targetable=False means we only
# assert chunks were created (the content is not a topic-QA target, e.g. tag,
# or the chunks are structured data that retrieval/generation can't answer from
# conversationally, e.g. table).
#
# These 12 are exactly the chunk_methods the dataset API accepts. ragflow also
# defines resume / knowledge_graph / audio internally, but the create-dataset
# API rejects them ("not in [...]"), so they cannot be covered here:
#   - resume/knowledge_graph: not exposed by the dataset API
#   - audio: also needs a speech2text model (none configured)
PARSER_PLAN = [
    ("naive", "prose_txt", True),
    ("paper", "prose_pdf", True),
    ("book", "prose_docx", True),
    ("laws", "law_txt", True),
    ("manual", "prose_pdf", True),
    ("presentation", "prose_pdf", True),
    ("one", "prose_txt", True),
    ("qa", "qa_xlsx", True),
    ("table", "table_xlsx", False),
    ("tag", "tag_xlsx", False),
    ("picture", "picture_png", True),
    ("email", "email_eml", True),
]


def _fact_pairs():
    """Flatten all topics' qa_pairs into a single (q, a, keywords) list."""
    out = []
    for tinfo in TOPICS.values():
        for p in tinfo.get("qa_pairs", []):
            out.append((p["question"], p["answer"], p.get("keywords", [])))
    return out


FACTS = _fact_pairs()


# ----------------------------------------------------------------------------
# Data generators — each returns (filename, bytes). They encode FACTS so the
# same questions are answerable regardless of the parser's required format.
# ----------------------------------------------------------------------------
def _facts_slice(seq, n=6):
    if not FACTS:
        return []
    start = (seq * n) % len(FACTS)
    return [FACTS[(start + i) % len(FACTS)] for i in range(min(n, len(FACTS)))]


def gen_prose_txt(seq):
    lines = [f"军事装备技术资料 第{seq + 1}册", ""]
    for q, a, _ in _facts_slice(seq):
        lines.append(f"问题：{q}")
        lines.append(f"解答：{a}")
        lines.append("")
    return f"doc_{seq:03d}.txt", ("\n".join(lines)).encode("utf-8")


def gen_prose_pdf(seq):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font = "STSong-Light"
    except Exception:
        font = "Helvetica"
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    y = 800
    c.setFont(font, 14)
    c.drawString(50, y, f"军事装备技术论文 No.{seq + 1}")
    y -= 30
    c.setFont(font, 11)
    for q, a, _ in _facts_slice(seq):
        for seg in (f"问题：{q}", f"解答：{a}", ""):
            c.drawString(50, y, seg[:60])
            y -= 18
            if y < 60:
                c.showPage(); c.setFont(font, 11); y = 800
    c.showPage(); c.save()
    return f"doc_{seq:03d}.pdf", buf.getvalue()


def gen_prose_docx(seq):
    from docx import Document
    doc = Document()
    doc.add_heading(f"军事装备手册 第{seq + 1}章", level=1)
    for q, a, _ in _facts_slice(seq):
        doc.add_heading(q, level=2)
        doc.add_paragraph(a)
    buf = BytesIO(); doc.save(buf)
    return f"doc_{seq:03d}.docx", buf.getvalue()


def gen_law_txt(seq):
    """laws parser needs numbered article structure (第N条) to chunk."""
    lines = [f"装备技术规范汇编 第{seq + 1}章 通用要求", ""]
    for i, (q, a, _) in enumerate(_facts_slice(seq, n=8), 1):
        topic = q.rstrip("？?").replace("是什么", "").replace("是多少", "")
        lines.append(f"第{i}条 {topic}应符合下列规定：{a}。")
    return f"law_{seq:03d}.txt", ("\n".join(lines)).encode("utf-8")


def gen_qa_xlsx(seq):
    """qa parser (excel): 2 columns question | answer."""
    from openpyxl import Workbook
    wb = Workbook(); ws = wb.active
    for q, a, _ in _facts_slice(seq, n=8):
        ws.append([q, a])
    buf = BytesIO(); wb.save(buf)
    return f"qa_{seq:03d}.xlsx", buf.getvalue()


def gen_table_xlsx(seq):
    """table parser (excel WITH header): rows encode parameter facts."""
    from openpyxl import Workbook
    wb = Workbook(); ws = wb.active
    ws.append(["问题", "答案", "序号"])
    for i, (q, a, _) in enumerate(_facts_slice(seq, n=8)):
        ws.append([q, a, seq * 100 + i])
    buf = BytesIO(); wb.save(buf)
    return f"table_{seq:03d}.xlsx", buf.getvalue()


def gen_tag_xlsx(seq):
    """tag parser (excel, NO header): 2 columns content | tags."""
    from openpyxl import Workbook
    wb = Workbook(); ws = wb.active
    for q, a, kws in _facts_slice(seq, n=8):
        ws.append([a, ",".join(kws) if kws else "军事"])
    buf = BytesIO(); wb.save(buf)
    return f"tag_{seq:03d}.xlsx", buf.getvalue()


def gen_resume_docx(seq):
    """resume parser: resume-structured docx (not a topic-QA target)."""
    from docx import Document
    doc = Document()
    doc.add_paragraph(f"姓名：技术员{seq + 1}")
    doc.add_paragraph("性别：男")
    doc.add_paragraph("学历：硕士")
    doc.add_paragraph("专业：军事装备工程")
    doc.add_paragraph("工作经历：装备维护与技术保障 5 年")
    doc.add_paragraph("技能：" + "、".join(
        kw for _, _, kws in _facts_slice(seq, n=3) for kw in kws))
    buf = BytesIO(); doc.save(buf)
    return f"resume_{seq:03d}.docx", buf.getvalue()


def gen_picture_png(seq):
    """picture parser: PNG with rendered Chinese text for built-in OCR."""
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new("RGB", (900, 500), "white")
    d = ImageDraw.Draw(img)
    font = None
    for fp in ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
               "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"):
        try:
            font = ImageFont.truetype(fp, 22); break
        except Exception:
            continue
    y = 20
    for q, a, _ in _facts_slice(seq, n=4):
        d.text((20, y), f"{q}", fill="black", font=font); y += 50
        d.text((20, y), f"{a}", fill="black", font=font); y += 70
    buf = BytesIO(); img.save(buf, format="PNG")
    return f"pic_{seq:03d}.png", buf.getvalue()


def gen_email_eml(seq):
    from email.message import EmailMessage
    msg = EmailMessage()
    msg["Subject"] = f"装备技术通报 No.{seq + 1}"
    msg["From"] = "tech@unit.mil"
    msg["To"] = "ops@unit.mil"
    body = ["技术问答通报：", ""]
    for q, a, _ in _facts_slice(seq):
        body.append(f"问：{q}")
        body.append(f"答：{a}")
        body.append("")
    msg.set_content("\n".join(body))
    # ragflow's email parser feeds the HTML part through HtmlParser and chokes
    # on a missing one — always include an HTML alternative.
    html_body = "<html><body>" + "".join(f"<p>{x}</p>" for x in body if x) + "</body></html>"
    msg.add_alternative(html_body, subtype="html")
    return f"mail_{seq:03d}.eml", bytes(msg)


GENERATORS = {
    "prose_txt": gen_prose_txt, "prose_pdf": gen_prose_pdf, "prose_docx": gen_prose_docx,
    "law_txt": gen_law_txt,
    "qa_xlsx": gen_qa_xlsx, "table_xlsx": gen_table_xlsx, "tag_xlsx": gen_tag_xlsx,
    "resume_docx": gen_resume_docx, "picture_png": gen_picture_png, "email_eml": gen_email_eml,
}


# ----------------------------------------------------------------------------
# ragflow helpers
# ----------------------------------------------------------------------------
def _create_kb(name, chunk_method):
    r = requests.post(f"{RAGFLOW_URL}/api/v1/datasets", headers=HEADERS,
                      json={"name": name, "chunk_method": chunk_method,
                            "language": "Chinese"}, timeout=30)
    d = r.json()
    assert d.get("code") == 0, f"create KB({chunk_method}) failed: {d.get('message')}"
    return d["data"]["id"]


def _upload(ds_id, filename, content):
    r = requests.post(f"{RAGFLOW_URL}/api/v1/datasets/{ds_id}/documents",
                      headers=UPLOAD_HEADERS, files={"file": (filename, content)}, timeout=120)
    d = r.json()
    if d.get("code") != 0:
        return None
    data = d["data"]
    return (data[0]["id"] if isinstance(data, list) else data["id"])


def _parse(ds_id, doc_ids):
    requests.post(f"{RAGFLOW_URL}/api/v1/datasets/{ds_id}/chunks", headers=HEADERS,
                  json={"document_ids": doc_ids}, timeout=30)


def _wait_parse(ds_id, timeout=900):
    start = time.time()
    while time.time() - start < timeout:
        r = requests.get(f"{RAGFLOW_URL}/api/v1/datasets/{ds_id}/documents?page_size=500",
                         headers=HEADERS, timeout=30)
        docs = r.json().get("data", {}).get("docs", [])
        running = [d for d in docs if d.get("run") in ("RUNNING", "UNSTART")]
        if not running:
            done = sum(1 for d in docs if d.get("run") in ("DONE", "SUCCESS"))
            fail = sum(1 for d in docs if d.get("run") == "FAIL")
            chunks = sum(d.get("chunk_count", 0) for d in docs)
            return done, fail, chunks, docs
        time.sleep(8)
    docs = requests.get(f"{RAGFLOW_URL}/api/v1/datasets/{ds_id}/documents?page_size=500",
                        headers=HEADERS, timeout=30).json().get("data", {}).get("docs", [])
    chunks = sum(d.get("chunk_count", 0) for d in docs)
    return -1, -1, chunks, docs


def _delete_kb(ds_id):
    try:
        requests.delete(f"{RAGFLOW_URL}/api/v1/datasets", headers=HEADERS,
                        json={"ids": [ds_id]}, timeout=30)
    except Exception:
        pass


def _ask(ds_id, question):
    """Create a transient assistant on the KB, ask one question, return answer."""
    name = f"h_ask_{int(time.time() * 1000) % 100000}"
    r = requests.post(f"{RAGFLOW_URL}/api/v1/chats", headers=HEADERS,
                      json={"name": name, "dataset_ids": [ds_id],
                            "llm": {"model_name": CHAT_MODEL, "temperature": 0.1},
                            "prompt": {"similarity_threshold": 0.1, "top_n": 6,
                                       "empty_response": ""}}, timeout=30)
    d = r.json()
    if d.get("code") != 0:
        return None, f"chat-create failed: {d.get('message')}"
    chat_id = d["data"]["id"]
    try:
        s = requests.post(f"{RAGFLOW_URL}/api/v1/chats/{chat_id}/sessions", headers=HEADERS,
                          json={"name": "s"}, timeout=30).json()
        sess_id = s["data"]["id"]
        c = requests.post(f"{RAGFLOW_URL}/api/v1/chats/{chat_id}/completions", headers=HEADERS,
                          json={"question": question, "session_id": sess_id, "stream": False},
                          timeout=120).json()
        ans = ""
        if c.get("code") == 0:
            data = c.get("data", {})
            ans = data.get("answer", "") if isinstance(data, dict) else ""
        return ans, None
    finally:
        try:
            requests.delete(f"{RAGFLOW_URL}/api/v1/chats", headers=HEADERS,
                            json={"ids": [chat_id]}, timeout=30)
        except Exception:
            pass


def _judge_semantic(question, expected, got, ds_id, keywords=None):
    """LLM-as-judge: is `got` semantically consistent with `expected`?

    ragflow 0.18 requires dataset_ids on chat create, so the judge chat is
    bound to the KB (retrieval is harmless — the verdict is decided purely from
    the explicit expected/got text in the prompt). Falls back to keyword
    overlap if the judge model is unusable, so QA never silently skips.
    """
    prompt = (f"判断下面的'实际回答'是否在语义上正确回答了问题、且与'参考答案'一致。"
              f"只输出 YES 或 NO。\n问题：{question}\n参考答案：{expected}\n实际回答：{got}")
    r = requests.post(f"{RAGFLOW_URL}/api/v1/chats", headers=HEADERS,
                      json={"name": f"h_judge_{int(time.time()*1000)%100000}",
                            "dataset_ids": [ds_id],
                            "llm": {"model_name": CHAT_MODEL, "temperature": 0.0}}, timeout=30)
    d = r.json()
    if d.get("code") != 0:
        # Fallback: keyword overlap (>=half the expected keywords present).
        if keywords:
            hit = sum(1 for k in keywords if k and k in (got or ""))
            return hit >= max(1, len(keywords) // 2)
        return None
    chat_id = d["data"]["id"]
    try:
        s = requests.post(f"{RAGFLOW_URL}/api/v1/chats/{chat_id}/sessions", headers=HEADERS,
                          json={"name": "j"}, timeout=30).json()
        sess_id = s["data"]["id"]
        c = requests.post(f"{RAGFLOW_URL}/api/v1/chats/{chat_id}/completions", headers=HEADERS,
                          json={"question": prompt, "session_id": sess_id, "stream": False},
                          timeout=120).json()
        verdict = ""
        if c.get("code") == 0 and isinstance(c.get("data"), dict):
            verdict = c["data"].get("answer", "")
        return "YES" in verdict.upper()
    finally:
        try:
            requests.delete(f"{RAGFLOW_URL}/api/v1/chats", headers=HEADERS,
                            json={"ids": [chat_id]}, timeout=30)
        except Exception:
            pass


def _active_parsers():
    only = os.environ.get("H_PARSERS", "").strip()
    if only:
        wanted = {p.strip() for p in only.split(",")}
        return [x for x in PARSER_PLAN if x[0] in wanted]
    return PARSER_PLAN


@pytest.mark.api
@pytest.mark.parametrize("parser_id,gen_kind,qa_targetable",
                         _active_parsers(),
                         ids=[x[0] for x in _active_parsers()])
@pytest.mark.timeout(1800)
def test_parser_coverage(parser_id, gen_kind, qa_targetable):
    """Per chunk_method: build KB, parse format-matched docs, verify chunks + QA."""
    gen = GENERATORS[gen_kind]
    ds_id = _create_kb(f"h_{parser_id}_{int(time.time())}", parser_id)
    try:
        doc_ids = []
        for seq in range(DOCS_PER_PARSER):
            fn, content = gen(seq)
            did = _upload(ds_id, fn, content)
            if did:
                doc_ids.append(did)
        assert doc_ids, f"[{parser_id}] no documents uploaded"
        _parse(ds_id, doc_ids)
        done, fail, chunks, docs = _wait_parse(ds_id)
        print(f"  [{parser_id}] uploaded={len(doc_ids)} done={done} fail={fail} chunks={chunks}")
        assert done != -1, f"[{parser_id}] parsing timed out"
        assert fail == 0, f"[{parser_id}] {fail} docs FAILED to parse"
        assert chunks > 0, f"[{parser_id}] parsed but produced 0 chunks"

        if not qa_targetable:
            return  # chunk-level coverage only (tag/resume content isn't a QA target)

        # End-to-end retrieval + QA: ask encoded questions, judge semantically.
        asked = 0
        passed = 0
        details = []
        for q, expected, kws in FACTS[:QA_QUESTIONS]:
            ans, err = _ask(ds_id, q)
            if err:
                pytest.skip(f"[{parser_id}] QA unavailable: {err}")
            asked += 1
            verdict = _judge_semantic(q, expected, ans or "", ds_id, keywords=kws)
            if verdict is None:
                pytest.skip(f"[{parser_id}] judge unavailable")
            passed += 1 if verdict else 0
            details.append((q, expected, (ans or "")[:60], verdict))
        for q, e, a, v in details:
            print(f"    QA[{parser_id}] {'OK' if v else 'XX'} q={q[:24]} exp={e[:24]} got={a}")
        assert passed >= max(1, asked // 2), \
            f"[{parser_id}] QA semantic match {passed}/{asked} below threshold"
    finally:
        _delete_kb(ds_id)

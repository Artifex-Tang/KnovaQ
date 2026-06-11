#!/usr/bin/env python3
"""Generate bulk DARPA-style test files in many variants.

Creates ~30 files per format (PDF, DOCX, XLSX, TXT, MD, JSON) with unique
military / technical content covering radar, communications, missiles, armor,
logistics, medical, engineering, CBRN, EW, UAV and more.

Usage:
    python generate_bulk_test_data.py
    python generate_bulk_test_data.py --count 10 --formats pdf,docx
"""

import argparse
import json
import os
import random
import sys
import textwrap
import time
from pathlib import Path

# ---------------------------------------------------------------------------
# Make imports work whether run from repo root or from the script's directory
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
REPO_TEST_RUNNER = SCRIPT_DIR
FIXTURES_DIR = REPO_TEST_RUNNER / "fixtures"
OUTPUT_ROOT = REPO_TEST_RUNNER / "test_data" / "bulk"

sys.path.insert(0, str(FIXTURES_DIR.parent))

# Import base helpers when available (non-fatal if missing)
try:
    from fixtures.test_data_factory import (
        generate_pdf,
        generate_docx,
        generate_xlsx,
        generate_txt,
        generate_qa_json,
    )
except ImportError:
    # Fallback implementations
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as pdf_canvas
    from docx import Document
    from openpyxl import Workbook

    def generate_pdf(path: Path, content: str) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        c = pdf_canvas.Canvas(str(path), pagesize=A4)
        w, h = A4
        margin, y, lh = 50, h - 50, 18
        for raw_line in content.split("\n"):
            if y < margin:
                c.showPage()
                y = h - margin
            line = raw_line.strip()
            if not line:
                y -= lh
                continue
            while len(line) > 80:
                c.drawString(margin, y, line[:80])
                y -= lh
                if y < margin:
                    c.showPage()
                    y = h - margin
                line = line[80:]
            c.drawString(margin, y, line)
            y -= lh
        c.save()
        return path

    def generate_docx(path: Path, content: str) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for raw_line in content.split("\n"):
            line = raw_line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
        doc.save(str(path))
        return path

    def generate_txt(path: Path, content: str) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        return path


# Re-export for JSON (same logic either way)
try:
    from fixtures.test_data_factory import generate_qa_json  # noqa: F811
except ImportError:
    pass

# ===========================================================================
# CONTENT TEMPLATES  –  10 military domains, 3+ topics each
# ===========================================================================

CATEGORIES = {
    "radar": {
        "label_cn": "雷达",
        "label_en": "radar",
        "topics": [
            ("搜索算法", "Search Algorithm"),
            ("频率规划", "Frequency Planning"),
            ("信号处理", "Signal Processing"),
            ("目标跟踪", "Target Tracking"),
            ("抗干扰技术", "Anti-Jamming"),
            ("杂波抑制", "Clutter Suppression"),
            ("波形设计", "Waveform Design"),
            ("相控阵天线", "Phased Array Antenna"),
            ("合成孔径", "SAR Imaging"),
            ("双基地雷达", "Bistatic Radar"),
            ("超视距探测", "Over-The-Horizon Detection"),
            ("低截获概率", "LPI Techniques"),
            ("气象雷达", "Weather Radar"),
            ("敌我识别", "IFF Systems"),
        ],
    },
    "comm": {
        "label_cn": "通信",
        "label_en": "communications",
        "topics": [
            ("频率规划", "Frequency Planning"),
            ("加密体制", "Encryption System"),
            ("跳频通信", "FHSS Communication"),
            ("卫星链路", "Satellite Link"),
            ("数据链协议", "Data Link Protocol"),
            ("短波通信", "HF Communication"),
            ("软件无线电", "SDR Architecture"),
            ("量子通信", "Quantum Comm"),
            ("战术互联网", "Tactical Internet"),
            ("抗毁组网", "Survivable Networking"),
            ("认知无线电", "Cognitive Radio"),
            ("量子密钥分发", "QKD"),
        ],
    },
    "missile": {
        "label_cn": "导弹",
        "label_en": "missiles",
        "topics": [
            ("制导系统", "Guidance System"),
            ("推进技术", "Propulsion"),
            ("弹道计算", "Ballistic Computation"),
            ("末段突防", "Terminal Penetration"),
            ("反舰导弹", "Anti-Ship Missile"),
            ("防空导弹", "SAM Systems"),
            ("巡航导弹", "Cruise Missile"),
            ("固体火箭", "Solid Rocket Motor"),
            ("多模导引头", "Multi-Mode Seeker"),
            ("导弹防御", "Missile Defense"),
        ],
    },
    "armor": {
        "label_cn": "装甲",
        "label_en": "armor",
        "topics": [
            ("复合装甲", "Composite Armor"),
            ("主动防护", "Active Protection"),
            ("火控系统", "Fire Control"),
            ("动力系统", "Powertrain"),
            ("夜视装备", "Night Vision"),
            ("步兵战车", "IFV"),
            ("两栖装甲", "Amphibious Armor"),
            ("装甲抢修", "ARV Recovery"),
            ("地雷防护", "Mine Protection"),
            ("城市作战", "Urban Combat"),
        ],
    },
    "logistics": {
        "label_cn": "后勤",
        "label_en": "logistics",
        "topics": [
            ("补给链管理", "Supply Chain"),
            ("野战修理", "Field Repair"),
            ("油料保障", "Fuel Supply"),
            ("弹药管理", "Ammunition Mgmt"),
            ("运输调度", "Transport Scheduling"),
            ("卫勤保障", "Medical Support"),
            ("军需物资", "Quartermaster"),
            ("仓储管理", "Warehouse Mgmt"),
            ("应急投送", "Emergency Airlift"),
            ("装备退役", "Equipment Retirement"),
        ],
    },
    "medical": {
        "label_cn": "卫生",
        "label_en": "medical",
        "topics": [
            ("战伤救治", "Combat Casualty Care"),
            ("核生化防护", "CBRN Medical"),
            ("心理卫勤", "Mental Health"),
            ("野战医院", "Field Hospital"),
            ("血液保障", "Blood Supply"),
            ("传染病防控", "Infectious Disease"),
            ("航空医学", "Aviation Medicine"),
            ("潜水医学", "Diving Medicine"),
            ("药材管理", "Pharmaceutical Mgmt"),
            ("卫生统计", "Medical Statistics"),
        ],
    },
    "engineer": {
        "label_cn": "工程",
        "label_en": "engineering",
        "topics": [
            ("渡河架桥", "River Crossing"),
            ("地雷战", "Mine Warfare"),
            ("筑城构筑", "Fortification"),
            ("爆破技术", "Demolition"),
            ("伪装工程", "Camouflage"),
            ("给水工程", "Water Supply"),
            ("道路抢修", "Road Repair"),
            ("核防护工程", "Nuclear Shelter"),
            ("工程机械", "Engineering Vehicles"),
            ("障碍物清除", "Obstacle Clearance"),
        ],
    },
    "cbrn": {
        "label_cn": "防化",
        "label_en": "CBRN",
        "topics": [
            ("毒剂侦检", "Agent Detection"),
            ("洗消技术", "Decontamination"),
            ("核辐射监测", "Radiation Monitoring"),
            ("生物检测", "Bio Detection"),
            ("防护装备", "Protective Equipment"),
            ("化学报警", "Chemical Alarm"),
            ("核爆探测", "Nuclear Detection"),
            ("气溶胶采样", "Aerosol Sampling"),
            ("烟幕施放", "Smoke Screening"),
            ("喷火器操作", "Flamethrower Ops"),
        ],
    },
    "ew": {
        "label_cn": "电子战",
        "label_en": "EW",
        "topics": [
            ("电子侦察", "ESM Recon"),
            ("电子干扰", "ECM Jamming"),
            ("电子防护", "ECCM"),
            ("频谱管理", "Spectrum Mgmt"),
            ("雷达告警", "RWR"),
            ("反辐射导弹", "ARM"),
            ("通信干扰", "Comm Jamming"),
            ("GPS对抗", "GPS Countermeasures"),
            ("光电对抗", "EO Countermeasures"),
            ("网络电磁战", "Cyber-EW"),
        ],
    },
    "uav": {
        "label_cn": "无人机",
        "label_en": "UAV",
        "topics": [
            ("侦察无人机", "Recon UAV"),
            ("攻击无人机", "Strike UAV"),
            ("蜂群控制", "Swarm Control"),
            ("自主导航", "Autonomous Nav"),
            ("数据链路", "Data Link"),
            ("目标识别", "Target Recognition"),
            ("无人机反制", "Anti-UAV"),
            ("太阳能长航时", "Solar HALE"),
            ("垂直起降", "VTOL"),
            ("编队协同", "Formation Coordination"),
        ],
    },
}

# ===========================================================================
# CONTENT GENERATORS
# ===========================================================================

# --- Shared data pools -----------------------------------------------------

EQUIPMENT_MODELS = [
    ("AN/TPQ-53", "反炮兵雷达", "Counter-battery radar"),
    ("AN/APG-81", "机载有源相控阵雷达", "AESA radar"),
    ("HQ-9B", "地空导弹系统", "SAM system"),
    ("Type 99A", "主战坦克", "MBT"),
    ("ZBD-04A", "步兵战车", "IFV"),
    ("Z-20", "通用直升机", "Utility helicopter"),
    ("Y-20", "战略运输机", "Strategic airlifter"),
    ("DF-26", "中程弹道导弹", "IRBM"),
    ("CJ-100", "巡航导弹", "Cruise missile"),
    ("PL-15", "空空导弹", "AAM"),
    ("HQ-19", "反导拦截弹", "ABM interceptor"),
    ("Type 055", "驱逐舰", "Destroyer"),
    ("J-20", "隐身战斗机", "Stealth fighter"),
    ("WZ-7", "无人侦察机", "UAV recon"),
    ("GJ-11", "隐身攻击无人机", "Stealth UCAV"),
]

FAILURE_CODES = [
    ("ERR-001", "发射机功率不足", "Transmitter power low"),
    ("ERR-002", "接收机噪声过高", "Receiver noise high"),
    ("ERR-003", "目标跟踪丢失", "Target track lost"),
    ("ERR-004", "通信链路中断", "Comm link failure"),
    ("ERR-005", "电源模块故障", "Power module fault"),
    ("ERR-006", "温度超限报警", "Temperature alarm"),
    ("ERR-007", "液压系统泄漏", "Hydraulic leak"),
    ("ERR-008", "导航精度下降", "Nav accuracy degraded"),
    ("ERR-009", "火控解算异常", "FCS computation error"),
    ("ERR-010", "数据总线错误", "Data bus error"),
]

MAINTENANCE_LEVELS = [
    ("日常维护", "Daily Maintenance", "每日"),
    ("周维护", "Weekly Maintenance", "每周"),
    ("月度维护", "Monthly Maintenance", "每月"),
    ("季度维护", "Quarterly Maintenance", "每季度"),
    ("年度维护", "Annual Maintenance", "每年"),
]

SPECS = {
    "freq": [
        ("工作频段", "Operating Frequency", "8-12", "GHz"),
        ("载波频率", "Carrier Frequency", "2.4", "GHz"),
        ("中频带宽", "IF Bandwidth", "20", "MHz"),
        ("跳频速率", "Hop Rate", "500", "hop/s"),
        ("调频带宽", "FM Bandwidth", "100", "MHz"),
    ],
    "power": [
        ("峰值功率", "Peak Power", "45", "kW"),
        ("平均功率", "Average Power", "5", "kW"),
        ("输出功率", "Output Power", "50", "W"),
        ("发射功率", "TX Power", "20", "W"),
        ("备用功率", "Standby Power", "2", "W"),
    ],
    "range": [
        ("探测距离", "Detection Range", "60", "km"),
        ("通信距离", "Comm Range", "30", "km"),
        ("作用距离", "Effective Range", "150", "km"),
        ("拦截距离", "Intercept Range", "200", "km"),
        ("识别距离", "Recognition Range", "80", "km"),
    ],
    "env": [
        ("工作温度", "Operating Temp", "-40~+55", "°C"),
        ("储存温度", "Storage Temp", "-50~+70", "°C"),
        ("防护等级", "IP Rating", "IP67", ""),
        ("抗振等级", "Vibration Grade", "MIL-STD-810G", ""),
        ("EMC等级", "EMC Level", "GJB 151B", ""),
    ],
    "reliability": [
        ("MTBF", "MTBF", "500", "h"),
        ("MTTR", "MTTR", "2", "h"),
        ("可用度", "Availability", "99.5", "%"),
        ("寿命", "Service Life", "20", "year"),
        ("大修间隔", "Overhaul Interval", "5000", "h"),
    ],
}


def _pick_specs(n=4):
    """Pick *n* random spec rows from different groups."""
    groups = random.sample(list(SPECS.keys()), min(n, len(SPECS)))
    rows = []
    for g in groups:
        rows.append(random.choice(SPECS[g]))
    return rows


def _rand_failures(n=3):
    return random.sample(FAILURE_CODES, min(n, len(FAILURE_CODES)))


def _rand_equip():
    return random.choice(EQUIPMENT_MODELS)


# ---------------------------------------------------------------------------
# Content builders for each document TYPE + domain TOPIC
# ---------------------------------------------------------------------------

def build_pdf_content(cat_key: str, topic_cn: str, topic_en: str, seq: int) -> str:
    """Build unique PDF-ready content for a given category + topic."""
    cat = CATEGORIES[cat_key]
    equip = _rand_equip()
    specs = _pick_specs(5)
    fails = _rand_failures(3)
    maint = random.sample(MAINTENANCE_LEVELS, 3)

    is_en = seq % 10 in (3, 7)  # ~30 % English
    if is_en:
        title = f"{cat['label_en'].upper()} — {topic_en} Technical Report"
        lines = [
            f"# {title}",
            "",
            f"## 1. System Overview",
            f"The {equip[2]} ({equip[0]}) is a critical asset in {cat['label_en'].lower()} operations. "
            f"This document covers {topic_en.lower()} analysis for the {equip[2]} platform.",
            "",
            "## 2. Technical Parameters",
        ]
        for s_cn, s_en, val, unit in specs:
            lines.append(f"- {s_en}: {val} {unit}")
        lines += [
            "",
            "## 3. Operational Procedures",
            "### 3.1 Pre-Operation Checks",
            "1. Verify power supply voltage is within specified range",
            "2. Run Built-In Test (BIT) and confirm all PASS",
            "3. Establish communication link with command post",
            "4. Confirm GPS lock with CEP < 10m",
            "5. Log start-up time and initial parameters",
            "",
            "### 3.2 Mission Execution",
            "Monitor system performance indicators continuously. "
            "Report any anomaly exceeding threshold within 5 minutes. "
            "Maintain log entries at minimum 15-minute intervals.",
            "",
            "## 4. Fault Diagnosis",
        ]
        for code, desc_cn, desc_en in fails:
            lines.append(f"- {code}: {desc_en} -> Check subsystem logs and replace module")
        lines += [
            "",
            "## 5. Maintenance Schedule",
        ]
        for ml_cn, ml_en, ml_per in maint:
            lines.append(f"- {ml_en}: Perform {ml_per}")
        lines += [
            "",
            "## 6. Safety Precautions",
            "- Ensure all personnel maintain safe distance during high-power transmission",
            "- Follow lock-out/tag-out procedures before maintenance",
            "- Report all near-miss incidents within 24 hours",
            "- Maintain radiation exposure records for all personnel",
            "",
            f"Document ID: {cat_key.upper()}-{topic_en.replace(' ','')}-{seq:03d}",
            f"Classification: UNCLASSIFIED // FOR OFFICIAL USE ONLY",
        ]
    else:
        title = f"{cat['label_cn']}系统{topic_cn}技术报告"
        lines = [
            f"# {title}",
            "",
            "## 第一章 系统概述",
            f"{equip[1]}（{equip[0]}）是{cat['label_cn']}领域的核心装备。"
            f"本文档重点分析{topic_cn}相关技术参数与维护规程。",
            "",
            "## 第二章 技术参数",
        ]
        for s_cn, s_en, val, unit in specs:
            lines.append(f"- {s_cn}：{val} {unit}")
        lines += [
            "",
            "## 第三章 操作规程",
            "### 3.1 开机前检查",
            "1. 确认供电电压在规定范围内",
            "2. 运行自检程序BIT，全部项目应为绿色",
            "3. 与指挥所建立通信链路",
            "4. 确认GPS定位精度 < 10米",
            "5. 记录开机时间和初始参数",
            "",
            "### 3.2 任务执行",
            "持续监控系统性能指标。任何异常超过阈值须在5分钟内上报。"
            "日志记录间隔不得大于15分钟。",
            "",
            "## 第四章 故障诊断",
        ]
        for code, desc_cn, desc_en in fails:
            lines.append(f"- {code}：{desc_cn} -> 检查对应子系统日志并更换模块")
        lines += [
            "",
            "## 第五章 维护计划",
        ]
        for ml_cn, ml_en, ml_per in maint:
            lines.append(f"- {ml_cn}：{ml_per}执行")
        lines += [
            "",
            "## 第六章 安全注意事项",
            "- 高功率发射期间确保所有人员保持安全距离",
            "- 维护前必须执行断电挂牌程序",
            "- 所有未遂事故须在24小时内报告",
            "- 做好辐射暴露剂量记录",
            "",
            f"文件编号：{cat_key.upper()}-{seq:03d}",
            "密级：内部 // 仅限公务使用",
        ]

    return "\n".join(lines)


def build_docx_content(cat_key: str, topic_cn: str, topic_en: str, seq: int) -> str:
    """Build unique DOCX-ready content — varies doc type by seq."""
    doc_types = [
        ("操作手册", "Operating Manual"),
        ("技术规格书", "Technical Specification"),
        ("标准操作程序", "SOP"),
        ("培训教材", "Training Material"),
        ("验收报告", "Acceptance Report"),
        ("改进方案", "Improvement Proposal"),
    ]
    dt = doc_types[seq % len(doc_types)]
    equip = _rand_equip()
    specs = _pick_specs(5)
    is_en = seq % 10 in (3, 7)

    if is_en:
        lines = [
            f"# {dt[1]} — {CATEGORIES[cat_key]['label_en']} {topic_en}",
            "",
            f"## Document Information",
            f"| Field | Value |",
            f"|-------|-------|",
            f"| Document Type | {dt[1]} |",
            f"| Equipment | {equip[2]} |",
            f"| Version | {random.randint(1,5)}.{random.randint(0,9)} |",
            f"| Date | 2026-06-{random.randint(1,28):02d} |",
            "",
            f"## 1. Scope",
            f"This document defines the {dt[1].lower()} for the {equip[2]} system.",
            "",
            f"## 2. Applicable Standards",
        ]
        standards = ["MIL-STD-810G", "GJB 150A-2009", "GJB 151B-2013", "ISO 9001:2015", "DEF-STAN-00-35"]
        for s in random.sample(standards, 3):
            lines.append(f"- {s}")
        lines += [
            "",
            "## 3. Technical Requirements",
        ]
        for s_cn, s_en, val, unit in specs:
            lines.append(f"- {s_en}: {val} {unit}")
        lines += [
            "",
            "## 4. Test Methods",
            "### 4.1 Environmental Test",
            "Place equipment in environmental chamber. Cycle temperature from -40C to +55C over 4 hours. Hold at each extreme for 2 hours. Verify all functional parameters within specification.",
            "",
            "### 4.2 Reliability Test",
            "Conduct continuous operation test for 500 hours. Record all failures. Calculate observed MTBF.",
            "",
            "### 4.3 EMC Test",
            "Perform radiated and conducted emissions testing per GJB 151B. Verify radiated susceptibility up to 200 V/m.",
            "",
            "## 5. Acceptance Criteria",
            "All parameters shall meet or exceed values specified in Section 3. Zero critical defects permitted.",
        ]
    else:
        lines = [
            f"# {dt[0]} — {CATEGORIES[cat_key]['label_cn']}{topic_cn}",
            "",
            f"## 文档信息",
            f"| 字段 | 内容 |",
            f"|------|------|",
            f"| 文档类型 | {dt[0]} |",
            f"| 装备型号 | {equip[0]}（{equip[1]}） |",
            f"| 版本号 | {random.randint(1,5)}.{random.randint(0,9)} |",
            f"| 编制日期 | 2026年6月{random.randint(1,28)}日 |",
            "",
            f"## 一、适用范围",
            f"本文件规定了{equip[1]}的{dt[0]}要求。",
            "",
            f"## 二、引用标准",
        ]
        standards = ["GJB 150A-2009", "GJB 151B-2013", "GJB 899A-2009", "GJB 450B-2008", "GJB 1362A-2007"]
        for s in random.sample(standards, 3):
            lines.append(f"- {s}")
        lines += [
            "",
            "## 三、技术要求",
        ]
        for s_cn, s_en, val, unit in specs:
            lines.append(f"- {s_cn}：{val} {unit}")
        lines += [
            "",
            "## 四、试验方法",
            "### 4.1 环境适应性试验",
            "将装备置于环境试验箱中，从-40°C升至+55°C循环，每个温度保持2小时。测试全部功能参数。",
            "",
            "### 4.2 可靠性试验",
            "连续运行500小时，记录所有故障，计算观测MTBF值。",
            "",
            "### 4.3 电磁兼容试验",
            "按GJB 151B进行辐射发射和传导发射测试。验证辐射敏感度不低于200 V/m。",
            "",
            "## 五、验收标准",
            "所有参数应满足第三章规定值，关键缺陷数为零。",
        ]
    return "\n".join(lines)


def build_xlsx_data(cat_key: str, topic_cn: str, topic_en: str, seq: int):
    """Return (title, headers, rows) for a unique XLSX sheet."""
    xlsx_types = [
        ("装备参数规范", "Equipment Parameters",
         ["装备型号", "参数名称", "参数值", "单位", "检验标准", "备注"]),
        ("库存台账", "Inventory Ledger",
         ["物资编号", "名称", "规格型号", "数量", "单位", "存放位置", "状态"]),
        ("维护记录", "Maintenance Log",
         ["记录编号", "装备型号", "维护级别", "维护日期", "执行人", "维护内容", "备注"]),
        ("故障统计", "Fault Statistics",
         ["故障编号", "装备型号", "故障代码", "故障现象", "发生时间", "处理措施", "恢复时间"]),
        ("训练记录", "Training Record",
         ["训练编号", "科目名称", "参训人数", "训练日期", "考核成绩", "教官", "备注"]),
        ("检验报告", "Inspection Report",
         ["检验项", "标准要求", "实测值", "判定", "检验员", "日期", "备注"]),
    ]
    idx = seq % len(xlsx_types)
    title_cn, title_en, headers = xlsx_types[idx]
    cat = CATEGORIES[cat_key]
    is_en = seq % 10 in (3, 7)
    title = title_en if is_en else f"{cat['label_cn']}{topic_cn}{title_cn}"

    rows = []
    for i in range(random.randint(8, 20)):
        equip = _rand_equip()
        if idx == 0:  # Parameters
            spec = random.choice(sum(SPECS.values(), []))
            std = random.choice(["GJB 150A-2009", "GJB 151B-2013", "GJB 899A-2009", "出厂检验", "设计指标"])
            rows.append([equip[0], spec[0], spec[2], spec[3], std, f"序号{i+1}"])
        elif idx == 1:  # Inventory
            qty = random.randint(10, 5000)
            loc = f"仓库{random.choice(['A','B','C','D'])}-{random.randint(1,20):02d}排-{random.randint(1,50):03d}号"
            status = random.choice(["正常", "待检", "维修中", "封存"])
            rows.append([f"WP-{random.randint(10000,99999)}", equip[1], equip[0], qty, "件", loc, status])
        elif idx == 2:  # Maintenance
            level = random.choice(["日常", "周", "月度", "季度", "年度"])
            month = random.randint(1, 6)
            day = random.randint(1, 28)
            content = random.choice(["清洁检查", "润滑保养", "参数校准", "部件更换", "系统升级"])
            rows.append([f"MR-{2026}{month:02d}{day:02d}-{i+1:03d}", equip[0], level,
                         f"2026-{month:02d}-{day:02d}", f"技师{random.randint(1,20):03d}号", content, ""])
        elif idx == 3:  # Fault
            fc = random.choice(FAILURE_CODES)
            month = random.randint(1, 6)
            day = random.randint(1, 28)
            fix = random.choice(["更换模块", "重新校准", "紧固连接", "软件复位", "更换部件"])
            hours = random.randint(1, 48)
            rows.append([f"FT-{2026}{month:02d}{day:02d}-{i+1:03d}", equip[0], fc[0], fc[1],
                         f"2026-{month:02d}-{day:02d} {random.randint(0,23):02d}:{random.randint(0,59):02d}",
                         fix, f"{hours}小时"])
        elif idx == 4:  # Training
            month = random.randint(1, 6)
            day = random.randint(1, 28)
            score = random.choice(["优秀", "良好", "合格", "不合格"])
            rows.append([f"TR-{2026}{month:02d}{day:02d}-{i+1:03d}",
                         f"{cat['label_cn']}{topic_cn}", random.randint(5, 50),
                         f"2026-{month:02d}-{day:02d}", score,
                         f"教官{random.randint(1,10):03d}号", ""])
        else:  # Inspection
            spec = random.choice(sum(SPECS.values(), []))
            measured = spec[2]
            judge = random.choices(["合格", "不合格"], weights=[90, 10])[0]
            rows.append([spec[0], f"{spec[2]} {spec[3]}", measured, judge,
                         f"检验员{random.randint(1,8):03d}号", f"2026-06-{random.randint(1,28):02d}", ""])

    return title, headers, rows


def build_txt_content(cat_key: str, topic_cn: str, topic_en: str, seq: int) -> str:
    """Build unique plain-text content."""
    txt_types = [
        ("野战指南", "Field Guide"),
        ("技术通报", "Technical Bulletin"),
        ("操作卡片", "Quick Reference Card"),
        ("维护规程", "Maintenance Procedure"),
        ("安全须知", "Safety Notice"),
        ("检查清单", "Checklist"),
    ]
    idx = seq % len(txt_types)
    cat = CATEGORIES[cat_key]
    equip = _rand_equip()
    is_en = seq % 10 in (3, 7)

    if is_en:
        title = f"{txt_types[idx][1]}: {cat['label_en']} {topic_en}"
        lines = [
            title,
            "=" * len(title),
            f"Equipment: {equip[2]} ({equip[0]})",
            f"Date: 2026-06-{random.randint(1,28):02d}",
            "",
            "1. PURPOSE",
            f"This {txt_types[idx][1].lower()} provides guidance for {topic_en.lower()} operations.",
            "",
            "2. PROCEDURES",
        ]
        for step in range(1, random.randint(5, 10) + 1):
            lines.append(f"  Step {step}: Operation detail line {step} for {equip[2]}.")
            lines.append(f"    - Verify parameter within range")
            lines.append(f"    - Record measurement in log")
        lines += [
            "",
            "3. PRECAUTIONS",
            "  - Always wear appropriate PPE",
            "  - Follow lock-out/tag-out procedures",
            "  - Report anomalies immediately",
            "  - Maintain clean workspace",
            "",
            "4. EMERGENCY PROCEDURES",
            "  - Shutdown: Press emergency stop button",
            "  - Evacuate: Follow posted evacuation route",
            f"  - Report: Contact maintenance at ext. {random.randint(1000,9999)}",
            "",
            "5. REFERENCES",
            "  - Technical Manual TM-{0}-{1}".format(random.randint(5, 99), random.randint(1000, 9999)),
            "  - GJB Standard GJB {0}-2009".format(random.choice([150, 151, 450, 899])),
        ]
    else:
        title = f"{txt_types[idx][0]}：{cat['label_cn']}{topic_cn}"
        lines = [
            title,
            "=" * len(title),
            f"装备型号：{equip[0]}（{equip[1]}）",
            f"编制日期：2026年6月{random.randint(1,28)}日",
            "",
            "一、目的",
            f"本{txt_types[idx][0]}用于指导{topic_cn}相关工作。",
            "",
            "二、操作步骤",
        ]
        for step in range(1, random.randint(5, 10) + 1):
            lines.append(f"  第{step}步：{equip[1]}{topic_cn}操作细节{step}。")
            lines.append(f"    - 确认参数在规定范围内")
            lines.append(f"    - 将测量数据记入日志")
        lines += [
            "",
            "三、注意事项",
            "  - 操作人员必须佩戴防护装备",
            "  - 严格执行断电挂牌制度",
            "  - 发现异常立即上报",
            "  - 保持工作区域整洁",
            "",
            "四、应急处理",
            "  - 紧急停机：按下急停按钮",
            "  - 人员撤离：按指定路线撤离",
            f"  - 事故报告：联系维修部门，分机号 {random.randint(1000,9999)}",
            "",
            "五、参考资料",
            f"  - 技术手册 TM-{random.randint(5,99)}-{random.randint(1000,9999)}",
            f"  - 国军标 GJB {random.choice([150, 151, 450, 899])}-2009",
        ]

    # Pad to ensure variety in file size
    if seq % 3 == 0:
        lines += [
            "",
            "附录A 常用数据表",
            "-" * 40,
        ]
        for i in range(10):
            spec = random.choice(sum(SPECS.values(), []))
            lines.append(f"  {spec[0]}: {spec[2]} {spec[3]}")

    return "\n".join(lines)


def build_md_content(cat_key: str, topic_cn: str, topic_en: str, seq: int) -> str:
    """Build unique Markdown content."""
    md_types = [
        ("常见问题解答", "FAQ"),
        ("操作规程", "Procedures"),
        ("检查清单", "Checklist"),
        ("技术说明", "Technical Notes"),
        ("培训大纲", "Training Outline"),
        ("变更记录", "Change Log"),
    ]
    idx = seq % len(md_types)
    cat = CATEGORIES[cat_key]
    equip = _rand_equip()
    specs = _pick_specs(4)
    is_en = seq % 10 in (3, 7)

    if is_en:
        title = f"{md_types[idx][1]}: {cat['label_en']} — {topic_en}"
        lines = [
            f"# {title}",
            "",
            f"> **Equipment:** {equip[2]} ({equip[0]})  ",
            f"> **Last Updated:** 2026-06-{random.randint(1,28):02d}",
            "",
        ]
        if idx == 0:  # FAQ
            for qi in range(1, random.randint(4, 8) + 1):
                q = f"Q{qi}: What is the {random.choice(['specification','procedure','requirement','threshold'])} for {topic_en.lower()}?"
                a = f"A{qi}: The {equip[2]} specification requires {random.choice(['periodic','continuous','manual','automated'])} monitoring with values between {random.randint(10,99)} and {random.randint(100,999)} {random.choice(['MHz','GHz','km','kW','hours'])}."
                lines += [f"## {q}", "", a, ""]
        elif idx == 2:  # Checklist
            lines += ["## Pre-Operation Checklist", ""]
            for ci in range(1, random.randint(8, 15) + 1):
                lines.append(f"- [ ] Item {ci}: Check {random.choice(['power supply','coolant level','antenna','cable','connector','display','battery','fan'])}")
            lines += ["", "## Post-Operation Checklist", ""]
            for ci in range(1, random.randint(5, 10) + 1):
                lines.append(f"- [ ] Item {ci}: Verify {random.choice(['shutdown sequence','data backup','cleaning','log entry','parameter reset'])}")
        else:
            lines += [
                "## 1. Overview",
                f"This section covers {topic_en} for the {equip[2]}.",
                "",
                "## 2. Specifications",
                "",
                "| Parameter | Value | Unit |",
                "|-----------|-------|------|",
            ]
            for s_cn, s_en, val, unit in specs:
                lines.append(f"| {s_en} | {val} | {unit} |")
            lines += [
                "",
                "## 3. Procedures",
                "```",
                f"1. Initialize {equip[2]} system",
                "2. Run self-test",
                "3. Verify all green",
                "4. Begin operation",
                "```",
                "",
                "## 4. Notes",
                "> Ensure compliance with all applicable regulations.",
            ]
    else:
        title = f"{md_types[idx][0]}：{cat['label_cn']} — {topic_cn}"
        lines = [
            f"# {title}",
            "",
            f"> **装备型号：** {equip[0]}（{equip[1]}）  ",
            f"> **更新日期：** 2026年6月{random.randint(1,28)}日",
            "",
        ]
        if idx == 0:  # FAQ
            for qi in range(1, random.randint(4, 8) + 1):
                q = f"问题{qi}：{equip[1]}的{topic_cn}{random.choice(['指标是多少？','如何操作？','有哪些注意事项？','维护周期是什么？','常见故障有哪些？'])}"
                a = (f"解答{qi}：根据技术规范，{equip[1]}的{topic_cn}要求{random.choice(['定期','连续','手动','自动'])}监测，"
                     f"数值范围{random.randint(10,99)}~{random.randint(100,999)}{random.choice(['MHz','GHz','km','kW','小时'])}。")
                lines += [f"## {q}", "", a, ""]
        elif idx == 2:  # Checklist
            lines += ["## 操作前检查清单", ""]
            for ci in range(1, random.randint(8, 15) + 1):
                lines.append(f"- [ ] 第{ci}项：检查{random.choice(['电源','冷却液','天线','电缆','接头','显示屏','电池','风扇'])}")
            lines += ["", "## 操作后检查清单", ""]
            for ci in range(1, random.randint(5, 10) + 1):
                lines.append(f"- [ ] 第{ci}项：确认{random.choice(['关机程序','数据备份','清洁','日志记录','参数复位'])}")
        else:
            lines += [
                "## 一、概述",
                f"本节介绍{equip[1]}的{topic_cn}相关内容。",
                "",
                "## 二、技术参数",
                "",
                "| 参数 | 数值 | 单位 |",
                "|------|------|------|",
            ]
            for s_cn, s_en, val, unit in specs:
                lines.append(f"| {s_cn} | {val} | {unit} |")
            lines += [
                "",
                "## 三、操作流程",
                "```",
                f"1. 初始化{equip[1]}系统",
                "2. 运行自检程序",
                "3. 确认全部通过",
                "4. 开始执行任务",
                "```",
                "",
                "## 四、备注",
                "> 请确保遵守所有适用的规章制度。",
            ]

    return "\n".join(lines)


def build_qa_content(cat_key: str, topic_cn: str, topic_en: str, seq: int) -> list:
    """Build unique QA pairs for JSON output."""
    cat = CATEGORIES[cat_key]
    equip = _rand_equip()
    specs = _pick_specs(3)
    is_en = seq % 10 in (3, 7)
    pairs = []

    if is_en:
        templates = [
            (f"What is the {spec[1]} of {equip[2]}?",
             f"{spec[2]} {spec[3]}", [spec[1].lower(), spec[2], spec[3]])
            for spec in specs
        ]
        templates += [
            (f"What category does {equip[2]} belong to?",
             cat["label_en"], [cat["label_en"].lower()]),
            (f"What is the primary function of {equip[2]}?",
             equip[2], [equip[0].lower(), equip[2].lower()]),
            (f"What maintenance interval is recommended for {equip[2]}?",
             random.choice(["daily", "weekly", "monthly", "quarterly", "annually"]),
             ["maintenance", "interval"]),
            (f"What standard applies to {equip[2]} testing?",
             random.choice(["MIL-STD-810G", "GJB 150A-2009", "GJB 151B-2013"]),
             ["standard", "GJB"]),
        ]
    else:
        templates = [
            (f"{equip[1]}的{spec[0]}是多少？",
             f"{spec[2]} {spec[3]}", [spec[0], spec[2], spec[3]])
            for spec in specs
        ]
        templates += [
            (f"{equip[0]}属于哪个类别？", cat["label_cn"], [cat["label_cn"]]),
            (f"{equip[1]}的主要用途是什么？", equip[1], [equip[0], equip[1]]),
            (f"{equip[1]}推荐的维护周期是？",
             random.choice(["每日", "每周", "每月", "每季度", "每年"]),
             ["维护", "周期"]),
            (f"{equip[1]}的检验依据什么标准？",
             random.choice(["GJB 150A-2009", "GJB 151B-2013", "GJB 899A-2009", "GJB 450B-2008"]),
             ["GJB", "标准"]),
            (f"{equip[1]}的工作温度范围是？",
             "-40°C 至 +55°C", ["温度", "-40", "+55"]),
            (f"{equip[1]}的防护等级是？",
             "IP67", ["防护", "IP67"]),
        ]

    for q, a, kw in templates:
        pairs.append({"question": q, "answer": a, "keywords": kw})

    return pairs


# ===========================================================================
# FILE WRITERS
# ===========================================================================

def write_pdf_file(output_dir: Path, cat_key: str, topic_cn: str, topic_en: str, seq: int) -> Path:
    fname = f"{cat_key}_{topic_cn}_{seq:03d}.pdf"
    content = build_pdf_content(cat_key, topic_cn, topic_en, seq)
    path = output_dir / fname
    generate_pdf(path, content)
    return path


def write_docx_file(output_dir: Path, cat_key: str, topic_cn: str, topic_en: str, seq: int) -> Path:
    fname = f"{cat_key}_{topic_cn}_{seq:03d}.docx"
    content = build_docx_content(cat_key, topic_cn, topic_en, seq)
    path = output_dir / fname
    generate_docx(path, content)
    return path


def write_xlsx_file(output_dir: Path, cat_key: str, topic_cn: str, topic_en: str, seq: int) -> Path:
    fname = f"{cat_key}_{topic_cn}_{seq:03d}.xlsx"
    title, headers, rows = build_xlsx_data(cat_key, topic_cn, topic_en, seq)
    path = output_dir / fname
    path.parent.mkdir(parents=True, exist_ok=True)

    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]  # Excel sheet name max 31 chars
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(str(path))
    return path


def write_txt_file(output_dir: Path, cat_key: str, topic_cn: str, topic_en: str, seq: int) -> Path:
    fname = f"{cat_key}_{topic_cn}_{seq:03d}.txt"
    content = build_txt_content(cat_key, topic_cn, topic_en, seq)
    path = output_dir / fname
    generate_txt(path, content)
    return path


def write_md_file(output_dir: Path, cat_key: str, topic_cn: str, topic_en: str, seq: int) -> Path:
    fname = f"{cat_key}_{topic_cn}_{seq:03d}.md"
    content = build_md_content(cat_key, topic_cn, topic_en, seq)
    path = output_dir / fname
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return path


def write_json_file(output_dir: Path, cat_key: str, topic_cn: str, topic_en: str, seq: int) -> Path:
    fname = f"{cat_key}_{topic_cn}_{seq:03d}.json"
    qa_pairs = build_qa_content(cat_key, topic_cn, topic_en, seq)
    path = output_dir / fname
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(qa_pairs, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


# ===========================================================================
# DISPATCH TABLE
# ===========================================================================

FORMAT_WRITERS = {
    "pdf":  write_pdf_file,
    "docx": write_docx_file,
    "xlsx": write_xlsx_file,
    "txt":  write_txt_file,
    "md":   write_md_file,
    "json": write_json_file,
}


# ===========================================================================
# MAIN
# ===========================================================================

def generate_all(count: int, formats: list[str]) -> dict:
    """Generate *count* files per format. Returns summary dict."""
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

    cat_keys = list(CATEGORIES.keys())
    summary = {fmt: {"count": 0, "files": []} for fmt in formats}
    total = count * len(formats)
    done = 0

    for fmt in formats:
        writer = FORMAT_WRITERS[fmt]
        fmt_dir = OUTPUT_ROOT / fmt
        fmt_dir.mkdir(parents=True, exist_ok=True)

        for i in range(1, count + 1):
            # Rotate through categories to spread domain coverage
            cat_key = cat_keys[(i - 1) % len(cat_keys)]
            cat = CATEGORIES[cat_key]
            # Pick a topic (rotate through available topics)
            topic_idx = (i - 1) % len(cat["topics"])
            topic_cn, topic_en = cat["topics"][topic_idx]

            try:
                path = writer(fmt_dir, cat_key, topic_cn, topic_en, i)
                summary[fmt]["count"] += 1
                summary[fmt]["files"].append(path.name)
                done += 1
                print(f"  [{done}/{total}] {fmt.upper():4s}  {path.name}")
            except Exception as exc:
                print(f"  [FAIL] {fmt} #{i}: {exc}")

    return summary


def main():
    parser = argparse.ArgumentParser(
        description="Generate bulk DARPA military QA test files in multiple formats."
    )
    parser.add_argument(
        "--count", type=int, default=30,
        help="Number of files per format (default: 30)",
    )
    parser.add_argument(
        "--formats", default="pdf,docx,xlsx,txt,md,json",
        help="Comma-separated list of formats (default: pdf,docx,xlsx,txt,md,json)",
    )
    args = parser.parse_args()

    formats = [f.strip().lower() for f in args.formats.split(",")]
    unknown = [f for f in formats if f not in FORMAT_WRITERS]
    if unknown:
        print(f"ERROR: Unknown format(s): {unknown}")
        print(f"Supported: {list(FORMAT_WRITERS.keys())}")
        sys.exit(1)

    print(f"=== DARPA Bulk Test Data Generator ===")
    print(f"  Output dir : {OUTPUT_ROOT}")
    print(f"  Formats    : {formats}")
    print(f"  Per format : {args.count}")
    print(f"  Total      : {args.count * len(formats)}")
    print()

    t0 = time.time()
    summary = generate_all(args.count, formats)
    elapsed = time.time() - t0

    # ---- Print summary ----
    print()
    print("=" * 60)
    print("GENERATION SUMMARY")
    print("=" * 60)
    grand_total = 0
    for fmt, info in summary.items():
        n = info["count"]
        grand_total += n
        print(f"  {fmt.upper():5s} : {n:4d} files")
    print(f"  {'TOTAL':5s} : {grand_total:4d} files")
    print(f"  Time  : {elapsed:.1f}s")
    print(f"  Output: {OUTPUT_ROOT}")
    print("=" * 60)


if __name__ == "__main__":
    main()

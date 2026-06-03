"""Generate DARPA-style test documents for E2E tests.

Creates PDF, DOCX, XLSX, TXT, MD, and JSON test files on demand.
"""

import json
import os
import tempfile
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


TEST_DATA_DIR = Path(__file__).parent.parent / "test_data"
MIXED_CORPUS_DIR = TEST_DATA_DIR / "mixed_corpus"

# ── DARPA-style Chinese content ──────────────────────────────

RADAR_REPORT_CN = """# 雷达系统技术评估报告

## 第一章 系统概述

AN/TPQ-53反炮兵雷达系统是新一代战场侦察装备，主要用于探测、定位和追踪敌方火炮、迫击炮和火箭炮发射阵地。系统探测距离可达60公里，反应时间小于10秒。

## 第二章 技术参数

- 工作频段：X波段（8-12 GHz）
- 峰值功率：45 kW
- 探测距离：60 km（火炮定位）
- 方位覆盖：360度
- 仰角覆盖：-5度至+85度
- 反应时间：< 10秒
- MTBF：> 500小时
- MTTR：< 2小时

## 第三章 维护规程

### 3.1 日常维护
每日检查天线阵列外观，确认无物理损伤。检查冷却系统液位，正常范围4.5-5.5升。运行自检程序BIT，全部项目应为绿色。

### 3.2 周维护
每周清洁天线罩表面，使用去离子水和软布。检查波导连接处密封性，扭矩值应在25-30 N·m范围内。

### 3.3 月度维护
每月校准频率合成器，频率偏差不超过±100 Hz。检查接收机灵敏度，最小可检测信号应优于-110 dBm。更新系统软件至最新版本。

### 3.4 年度维护
每年进行全面性能评估。更换冷却系统滤芯。校准天线方向图。测试全系统MTBF指标。

## 第四章 故障诊断

常见故障代码：
- ERR-001：发射机功率不足 → 检查行波管和高压电源
- ERR-002：接收机噪声过高 → 检查低噪声放大器和电缆连接
- ERR-003：目标丢失 → 检查信号处理器和跟踪算法参数
- ERR-004：通信中断 → 检查光纤链路和网络交换机
"""

EQUIPMENT_MANUAL_CN = """通信装备操作手册

第一章 设备简介
ZBD-2000数字战场通信系统是为师级以下作战单元设计的综合通信平台，支持语音、数据、图像传输。

1.1 主要功能
- 超短波通信（VHF/UHF）
- 卫星通信接入
- 数据链传输
- 加密语音通信
- 态势信息共享

1.2 技术指标
- 频率范围：30-512 MHz
- 信道数：200个预置信道
- 输出功率：5W/20W/50W可调
- 工作温度：-40°C 至 +55°C
- 防护等级：IP67
- 重量：< 3.5 kg（含电池）

第二章 操作流程

2.1 开机检查
1. 确认电池电量 > 30%
2. 连接天线，检查接头紧固
3. 按住电源键3秒启动
4. 等待自检完成（约15秒）
5. 确认屏幕显示"系统就绪"

2.2 频率设置
1. 进入菜单 → 信道管理
2. 选择目标信道号
3. 输入频率值（MHz）
4. 设置调制方式（AM/FM/数字）
5. 保存并确认

2.3 加密通信
1. 插入加密模块
2. 输入当日密钥
3. 确认加密指示灯亮起
4. 进行密钥同步测试

第三章 维护保养

3.1 日常保养
- 清洁设备表面
- 检查接口密封
- 确认电池状态

3.2 故障排除
| 故障现象 | 可能原因 | 处理方法 |
|---------|---------|---------|
| 无法开机 | 电池耗尽 | 更换电池 |
| 接收信号弱 | 天线松动 | 紧固天线接头 |
| 通信中断 | 频率漂移 | 重新校准频率 |
| 加密失败 | 密钥过期 | 更新当日密钥 |
"""

FIELD_GUIDE_CN = """野战维护指南

一、装备检查清单

每日检查项目：
1. 外观检查：确认无破损、无锈蚀、无渗漏
2. 电源检查：电池电压 > 24V，备用电池 > 22V
3. 通信检查：与指挥所建立通信链路，信号强度 > 3级
4. 定位检查：GPS定位精度 < 10米
5. 记录检查：确认上一班次交接记录完整

二、常见故障应急处理

1. 发电机故障
   - 现象：输出电压不稳定
   - 步骤：检查燃油 → 检查机油 → 检查空滤 → 测量输出电压
   - 应急：切换备用发电机，记录故障时间

2. 天线系统故障
   - 现象：信号中断或质量下降
   - 步骤：检查馈线连接 → 检查天线方向 → 检查避雷器
   - 应急：启用备用天线方案

三、弹药储存规范

温湿度要求：
- 温度：-20°C 至 +40°C
- 湿度：45% - 75%
- 通风：每小时换气次数 ≥ 3次

堆放要求：
- 底部垫高 ≥ 15cm
- 垛间距 ≥ 0.5m
- 垛与墙距 ≥ 0.3m
- 单垛高度不超过2m

四、车辆保养计划

每500公里或每周：
- 检查机油液面
- 检查冷却液
- 检查轮胎气压（标准值：前轮 3.5bar，后轮 4.0bar）
- 检查制动系统
"""

POLICY_LAWS_CN = """# 军事装备管理条例

## 第一章 总则

第一条 为加强军事装备管理，保障装备完好率和战备水平，根据《中华人民共和国国防法》制定本条例。

第二条 本条例适用于全军各级装备管理部门及装备使用单位。

## 第二章 装备分类与编配

第四条 装备按用途分为以下类别：
（一）作战装备：直接用于作战行动的武器、弹药及配套设备
（二）保障装备：用于作战保障的通信、工程、防化、运输等装备
（三）训练装备：用于部队训练的模拟器和教练设备
（四）储备装备：战略储备和应急储备的装备物资

## 第三章 日常管理

第六条 装备日常管理实行"三定"制度：
（一）定人管理：每件装备明确管理责任人
（二）定位存放：装备在固定位置存放，标识清晰
（三）定期检查：按规定的周期和内容进行检查

第七条 装备检查分为：
（一）日常检查：每日由使用人员完成
（二）周检查：每周由班组长组织
（三）月检查：每月由连队主官组织
（四）季度检查：每季度由营级单位组织

第八条 装备维护保养分为：
（一）日常保养：使用前后进行
（二）一级保养：每月进行，以清洁、润滑、紧固为主
（三）二级保养：每季度进行，以检测、调整、更换易损件为主
（四）三级保养：每年进行，由专业修理机构实施

## 第四章 装备维修

第九条 装备维修实行三级维修体制：
（一）基层级维修：由使用单位完成简单故障排除
（二）中继级维修：由旅团修理所完成部件更换和调整
（三）基地级维修：由修理工厂完成大修和翻新
"""

# ── English content ──────────────────────────────────────────

TECH_REPORT_EN = """DARPA Advanced Radar Signal Processing Technical Report

Executive Summary
This report presents findings from the Phase II evaluation of adaptive radar signal processing algorithms under the DARPA MTO program.

1. Introduction
Modern battlefield radar systems face increasingly complex electromagnetic environments.

2. Technical Approach
We developed a novel Space-Time Adaptive Processing (STAP) architecture.

3. Key Findings
3.1 Detection Performance
The adaptive algorithm achieved a 12 dB improvement in SINR compared to conventional matched filtering.

3.2 Computational Requirements
Processing latency: 45 ms per CPI (128 pulses, 16 channels)
Memory footprint: 2.3 GB for full covariance estimation
Power consumption: 85W average on Xilinx VU9P FPGA

3.3 Clutter Rejection
Main beam clutter rejection: > 55 dB
Sidelobe clutter rejection: > 70 dB

4. Conclusions
Recommendation: proceed to Phase III field trials with operational units.
"""

# ── QA pairs for accuracy evaluation ────────────────────────

QA_PAIRS = [
    {"question": "AN/TPQ-53雷达的探测距离是多少？", "answer": "60公里", "keywords": ["探测距离", "60", "公里"]},
    {"question": "雷达的日常维护中冷却系统液位正常范围是多少？", "answer": "4.5-5.5升", "keywords": ["冷却系统", "液位", "4.5", "5.5"]},
    {"question": "ZBD-2000通信系统的频率范围？", "answer": "30-512 MHz", "keywords": ["频率", "30", "512", "MHz"]},
    {"question": "通信装备加密通信需要插入什么？", "answer": "加密模块", "keywords": ["加密", "模块"]},
    {"question": "弹药储存温度要求是什么？", "answer": "-20°C至+40°C", "keywords": ["温度", "-20", "+40"]},
    {"question": "弹药储存湿度要求？", "answer": "45%-75%", "keywords": ["湿度", "45", "75"]},
    {"question": "车辆保养中前轮标准胎压是多少？", "answer": "3.5bar", "keywords": ["前轮", "胎压", "3.5"]},
    {"question": "装备日常管理的三定制度是什么？", "answer": "定人管理、定位存放、定期检查", "keywords": ["三定", "定人", "定位", "定期"]},
    {"question": "装备维修三级体制是什么？", "answer": "基层级维修、中继级维修、基地级维修", "keywords": ["基层", "中继", "基地", "维修"]},
    {"question": "雷达ERR-001故障代码表示什么？", "answer": "发射机功率不足", "keywords": ["ERR-001", "发射机", "功率"]},
    {"question": "雷达周维护需要检查波导连接处的扭矩值范围？", "answer": "25-30 N·m", "keywords": ["扭矩", "25", "30"]},
    {"question": "通信装备开机后自检大约需要多少秒？", "answer": "15秒", "keywords": ["自检", "15秒"]},
    {"question": "DARPA报告中的自适应算法SINR提升了多少？", "answer": "12 dB", "keywords": ["SINR", "12", "dB"]},
    {"question": "雷达信号处理延迟是多少毫秒？", "answer": "45ms", "keywords": ["延迟", "45", "ms"]},
    {"question": "装备一级保养的周期是？", "answer": "每月", "keywords": ["一级保养", "每月"]},
    {"question": "ZBD-2000防护等级是多少？", "answer": "IP67", "keywords": ["防护", "IP67"]},
    {"question": "弹药堆放底部垫高要求？", "answer": "大于等于15cm", "keywords": ["垫高", "15cm"]},
    {"question": "通信装备重量限制？", "answer": "小于3.5kg含电池", "keywords": ["重量", "3.5", "kg"]},
    {"question": "DARPA雷达主瓣杂波抑制指标？", "answer": "大于55dB", "keywords": ["主瓣", "杂波", "55", "dB"]},
    {"question": "雷达接收机灵敏度指标？", "answer": "优于-110 dBm", "keywords": ["灵敏度", "-110", "dBm"]},
]


def ensure_dirs():
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    MIXED_CORPUS_DIR.mkdir(parents=True, exist_ok=True)


def generate_txt(path: Path, content: str) -> Path:
    ensure_dirs()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return path


def generate_pdf(path: Path, content: str) -> Path:
    ensure_dirs()
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=A4)
    w, h = A4
    margin = 50
    y = h - margin
    line_height = 18
    for line in content.split("\n"):
        if y < margin:
            c.showPage()
            y = h - margin
        line = line.strip()
        if not line:
            y -= line_height
            continue
        while len(line) > 80:
            c.drawString(margin, y, line[:80])
            y -= line_height
            if y < margin:
                c.showPage()
                y = h - margin
            line = line[80:]
        c.drawString(margin, y, line)
        y -= line_height
    c.save()
    return path


def generate_docx(path: Path, content: str) -> Path:
    ensure_dirs()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    doc.save(str(path))
    return path


def generate_xlsx(path: Path) -> Path:
    ensure_dirs()
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "装备参数规范"
    headers = ["装备型号", "参数名称", "参数值", "单位", "检验标准", "备注"]
    ws.append(headers)
    data = [
        ["AN/TPQ-53", "工作频段", "8-12", "GHz", "GJB XXXX-2020", "X波段"],
        ["AN/TPQ-53", "峰值功率", "45", "kW", "出厂检验", "行波管输出"],
        ["AN/TPQ-53", "探测距离", "60", "km", "GJB XXXX-2020", "火炮定位"],
        ["AN/TPQ-53", "方位覆盖", "360", "度", "设计指标", "全向扫描"],
        ["AN/TPQ-53", "MTBF", "500", "小时", "可靠性试验", "平均无故障时间"],
        ["ZBD-2000", "频率范围", "30-512", "MHz", "GJB XXXX-2020", "全频段覆盖"],
        ["ZBD-2000", "输出功率", "5/20/50", "W", "出厂检验", "三档可调"],
        ["ZBD-2000", "工作温度", "-40~+55", "°C", "环境试验", "全温度范围"],
        ["ZBD-2000", "防护等级", "IP67", "", "GJB XXXX-2020", "防尘防水"],
        ["ZBD-2000", "重量", "3.5", "kg", "称重", "含电池"],
    ]
    for row in data:
        ws.append(row)
    wb.save(str(path))
    return path


def generate_qa_json(path: Path) -> Path:
    ensure_dirs()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(QA_PAIRS, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def generate_all_test_files(output_dir: Path = None) -> dict:
    """Generate all test data files. Returns dict of {name: Path}."""
    base = output_dir or TEST_DATA_DIR
    ensure_dirs()
    files = {
        "radar_report": generate_txt(base / "radar_report.txt", RADAR_REPORT_CN),
        "equipment_manual": generate_docx(base / "equipment_manual.docx", EQUIPMENT_MANUAL_CN),
        "spec_table": generate_xlsx(base / "spec_table.xlsx"),
        "field_guide": generate_txt(base / "field_guide.txt", FIELD_GUIDE_CN),
        "policy_laws": generate_txt(base / "policy_laws.md", POLICY_LAWS_CN),
        "cn_military_doc": generate_txt(MIXED_CORPUS_DIR / "cn_military_doc.txt", RADAR_REPORT_CN),
        "en_tech_report": generate_txt(MIXED_CORPUS_DIR / "en_tech_report.txt", TECH_REPORT_EN),
        "qa_pairs": generate_qa_json(MIXED_CORPUS_DIR / "structured_qa.json"),
    }
    return files


def upload_test_documents(client, dataset_id: str, output_dir: Path = None) -> list:
    """Generate files and upload them to ragflow dataset. Returns list of document data."""
    files = generate_all_test_files(output_dir)
    all_docs = []
    for name, path in files.items():
        if path.suffix == ".json":
            continue  # Skip JSON QA pairs (used for evaluation only)
        docs = client.upload_document(dataset_id, str(path))
        if isinstance(docs, list):
            all_docs.extend(docs)
        else:
            all_docs.append(docs)
    return all_docs


def wait_for_parsing(client, dataset_id: str, timeout: int = 300, interval: int = 5) -> bool:
    """Delegate to client's wait_for_parsing."""
    return client.wait_for_parsing(dataset_id, timeout=timeout, interval=interval)


def get_qa_pairs() -> list:
    """Return QA evaluation pairs."""
    return QA_PAIRS

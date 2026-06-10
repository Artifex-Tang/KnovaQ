"""
基于 软件设计模板.docx 填充 DARPA智能问答服务工具-软件设计说明书
保留模板的封面页、修改记录页、目录页样式，替换正文内容
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

TPL  = Path(r'E:\ccode\KnovaQ\docs\软件设计模板.docx')
MD   = Path(r'E:\ccode\KnovaQ\docs\output\DARPA智能问答服务工具-软件设计说明书.md')
import os, uuid
_f = f'sdd_{uuid.uuid4().hex[:8]}.docx'
# remove if exists from previous run
_p = os.path.join(r'E:\ccode\KnovaQ\docs\output', _f)
if os.path.exists(_p):
    os.remove(_p)
OUT  = Path(_p)

# ── helper: style names (Chinese, must match template) ────────────
S_BODY     = 'A 正文'            # A 正文
S_BODY_L1  = 'A 正文 1级'    # A 正文 1级
S_BODY_L2  = 'A 正文 2级'    # A 正文 2级
S_FIG_CAP  = 'A 图题'            # A 图题
S_FIG      = 'A 图'                  # A 图
S_NOTE     = 'A 注解'            # A 注解
S_TBL_GRID = 'Table Grid'

def heading_style(level):
    return f'Heading {level}'

# ── helper: remove element from XML body ─────────────────────────
def remove_elem(elem):
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)

def clear_from_heading1(doc):
    """Remove everything from first Heading 1 onwards, return insert point."""
    body = doc.element.body
    found = False
    to_remove = []
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if not found and tag == 'p':
            # check if this paragraph uses Heading 1
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
                    found = True
        if found:
            to_remove.append(child)
    # keep last sectPr
    last_sp = body.find(qn('w:sectPr'))
    for elem in to_remove:
        if elem is not last_sp:
            remove_elem(elem)
    return found

# ── helper: add content using template styles ────────────────────
def add_heading(doc, text, level):
    p = doc.add_paragraph(text, style=heading_style(level))
    return p

def add_body(doc, text, indent_level=0):
    """Add body text paragraph. Handles **bold** markers."""
    style = S_BODY if indent_level == 0 else (S_BODY_L1 if indent_level == 1 else S_BODY_L2)
    try:
        p = doc.add_paragraph(style=style)
    except:
        p = doc.add_paragraph(style='Normal')
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
    return p

def add_figure_caption(doc, text):
    """Add figure caption (centered, bold)."""
    try:
        p = doc.add_paragraph(style=S_FIG_CAP)
    except:
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    return p

def add_code_block(doc, code_text):
    """Add ASCII art / code block in monospace."""
    for line in code_text.split('\n'):
        p = doc.add_paragraph(style='Normal')
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.left_indent = Cm(1)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # gray background
        pPr = p._element.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F5F5F5"/>')
        pPr.append(shd)

def add_table(doc, rows):
    """Add a formatted table."""
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols, style=S_TBL_GRID)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci < ncols:
                cell = tbl.rows[ri].cells[ci]
                cell.text = ''
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(1)
                cp.paragraph_format.space_after = Pt(1)
                run = cp.add_run(cell_text)
                run.font.size = Pt(9)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    # blue header
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="4472C4"/>')
                    tcPr.append(shd)
    doc.add_paragraph()  # spacer
    return tbl

# ── main: fill template ─────────────────────────────────────────
def fill_cover(doc):
    """Fill cover page tables with our info."""
    # table 0: identification (rows x 2 cols)
    # row 2 col 0: title
    t0 = doc.tables[0]
    # title cell
    title_cell = t0.rows[2].cells[0]
    title_cell.text = ''
    for p in title_cell.paragraphs:
        for r in p.runs:
            r.clear()
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n\nDARPA\n智能问答服务工具\n软件设计说明书\n\n')
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

    # row 0 col 0: identification
    id_cell = t0.rows[0].cells[0]
    id_cell.text = ''
    p = id_cell.paragraphs[0]
    run = p.add_run('标识号：DARPA-IQAS/SDD\n版本号：V1.0')
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # row 0 col 1: numbering
    num_cell = t0.rows[0].cells[1]
    num_cell.text = ''
    p = num_cell.paragraphs[0]
    run = p.add_run('编号：DARPA-IQAS-SDD-001\n密级：内部')
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # table 1: signature page
    t1 = doc.tables[1]
    sig_cell = t1.rows[1].cells[0]
    sig_cell.text = ''
    p = sig_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DARPA智能问答服务工具\n软件设计说明书')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

def fill_revision(doc):
    """Fill revision history table."""
    t2 = doc.tables[2]
    # row 1 (0 is header)
    revisions = [
        ('V1.0', '2026-06-04', '编写组', '初始版本', ''),
    ]
    for ri, rev in enumerate(revisions):
        row = t2.rows[ri + 1]
        for ci, val in enumerate(rev):
            cell = row.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def fill_references_table(doc):
    """Fill the 引用文档 table (table 6 in template)."""
    t6 = doc.tables[6]
    refs = [
        ('1', 'GJB 438B-2009', '军用软件开发文档通用要求', '—', '2009', '国委装备发展部'),
        ('2', 'GJB 450A-2004', '装备可靠性工作通用要求', '—', '2004', '—'),
        ('3', 'GJB 368B-2009', '装备维修性工作通用要求', '—', '2009', '—'),
        ('4', 'GJB 2547A-2012', '装备测试性工作通用要求', '—', '2012', '—'),
        ('5', 'GJB 900-90', '系统安全性通用大纲', '—', '1990', '—'),
        ('6', 'GJB 4239-2001', '装备环境适应性通用要求', '—', '2001', '—'),
        ('7', '—', 'DARPA智能问答服务工具-软件需求规格说明书', 'V1.0', '2026', '本项目'),
        ('8', '—', 'ragflow v0.18.0 官方文档', 'v0.18.0', '2025', 'infiniflow'),
    ]
    # Ensure enough rows
    while len(t6.rows) < len(refs) + 1:
        t6.add_row()
    for ri, ref in enumerate(refs):
        row = t6.rows[ri + 1]
        for ci, val in enumerate(ref):
            cell = row.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ── parse markdown and build content ────────────────────────────
def parse_md_and_build(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buf = []

    # skip everything before first chapter heading "# 1 范围"
    start = 0
    for idx, line in enumerate(lines):
        if re.match(r'^# 1 范围', line):
            start = idx
            break

    i = start
    while i < len(lines):
        line = lines[i]

        # code block
        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, '\n'.join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # skip empty / hr
        if line.strip() == '' or re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # figure/table caption: "图 X ..." or "表 X ..."
        cap_m = re.match(r'^(图|表)\s+(\d+)\s+(.+)$', line)
        if cap_m and not line.startswith('|'):
            add_figure_caption(doc, line)
            i += 1
            continue

        # heading
        hm = re.match(r'^(#{1,7})\s+(.+)$', line)
        if hm:
            level = min(len(hm.group(1)), 6)
            text = hm.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # table
        if line.strip().startswith('|'):
            tbl_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rl = lines[i].strip()
                if re.match(r'^\|[\s\-:|]+\|$', rl):
                    i += 1
                    continue
                cells = [c.strip() for c in rl.strip('|').split('|')]
                tbl_rows.append(cells)
                i += 1
            if tbl_rows:
                add_table(doc, tbl_rows)
            continue

        # normal paragraph - detect indentation level
        stripped = line.strip()
        indent = 0
        # count leading spaces for indent level
        if line.startswith('    '):
            spaces = len(line) - len(line.lstrip())
            indent = min(spaces // 4, 2)

        add_body(doc, stripped, indent_level=indent)
        i += 1


# ── main ────────────────────────────────────────────────────────
def main():
    # open template directly, save to different file
    doc = Document(str(TPL))

    # 1. fill cover
    fill_cover(doc)

    # 2. fill revision
    fill_revision(doc)

    # 3. fill references
    fill_references_table(doc)

    # 4. clear template content from first Heading 1
    clear_from_heading1(doc)

    # 5. read markdown content
    md_text = MD.read_text(encoding='utf-8')

    # 6. build content
    parse_md_and_build(doc, md_text)

    # 7. save
    doc.save(str(OUT))
    sz = OUT.stat().st_size / 1024
    print(f'Saved: {OUT} ({sz:.0f} KB)')

if __name__ == '__main__':
    main()

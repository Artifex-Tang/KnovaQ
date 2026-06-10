"""Generate DARPA系统测试大纲.docx following GJB template structure.

Merges 系统测试计划模板 + 系统测试说明模板 into a single document.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT as WDTableAlignment
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent

# ── style helpers ────────────────────────────────────────────

def set_run(run, font_name='宋体', font_name_ascii='Times New Roman', size=Pt(10.5), bold=False, color=None):
    run.font.size = size
    run.bold = bold
    run.font.name = font_name_ascii
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color


def set_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(0), space_after=Pt(6), line_spacing=1.5):
    pf = p.paragraph_format
    pf.alignment = alignment
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing


def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    el = OxmlElement('w:shd')
    el.set(qn('w:fill'), color_hex)
    el.set(qn('w:val'), 'clear')
    tcPr.append(el)


def set_cell_vertical_alignment(cell, align='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    val_map = {'center': 'center', 'top': 'top', 'bottom': 'bottom'}
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val_map.get(align, 'center'))
    tcPr.append(vAlign)


def make_table(doc, headers, rows, col_widths=None, header_color='D9E2F3'):
    """Create a styled table with header row."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WDTableAlignment.CENTER

    # Set column widths if provided
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=Pt(9), bold=True)
        set_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
        set_cell_shading(cell, header_color)
        set_cell_vertical_alignment(cell)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run(run, size=Pt(9))
            set_paragraph(p, space_after=Pt(2))

    return table


def add_heading(doc, text, level=1):
    """Add heading with proper formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return h


def add_para(doc, text, bold=False, indent=False, bullet=False):
    """Add a paragraph with standard formatting."""
    if bullet:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
    set_run(run, bold=bold)
    set_paragraph(p)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_test_case_table(doc, tc_id, tc_name, purpose, precondition, steps, expected, trace):
    """Add a test case in template table format (系统测试说明模板 Table 4)."""
    # TC title
    p = doc.add_paragraph()
    run = p.add_run(f'{tc_id}  {tc_name}')
    set_run(run, font_name='黑体', size=Pt(11), bold=True)
    set_paragraph(p, space_before=Pt(6), space_after=Pt(4))

    # Test case table following template structure
    rows_data = [
        ['测试目的', purpose],
        ['前置条件', precondition],
        ['测试步骤', steps],
        ['预期结果', expected],
        ['需求追踪', trace],
    ]
    make_table(doc, ['项目', '说明'], rows_data, col_widths=[Cm(3), Cm(13)])
    doc.add_paragraph()  # spacer


# ── build document ───────────────────────────────────────────

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ═══════════════════════════════════════════════════════════════
# SIGNATURE PAGE (模板封面)
# ═══════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DARPA智能问答服务工具')
set_run(run, font_name='黑体', size=Pt(22), bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('系统测试大纲')
set_run(run, font_name='黑体', size=Pt(22), bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('（测试计划 + 测试说明 合并版）')
set_run(run, size=Pt(14))

for _ in range(2):
    doc.add_paragraph()

# Signature table
sig_table = doc.add_table(rows=6, cols=4)
sig_table.style = 'Table Grid'
sig_table.alignment = WDTableAlignment.CENTER
sig_labels = [
    ('编制', '日期'), ('校对', '日期'), ('审核', '日期'),
    ('会签', '日期'), ('签发', '日期'), ('批准', '日期'),
]
for ri, (label, date_label) in enumerate(sig_labels):
    for ci in range(4):
        cell = sig_table.rows[ri].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        if ci == 0:
            run = p.add_run(label)
            set_run(run, size=Pt(11))
        elif ci == 2:
            run = p.add_run(date_label)
            set_run(run, size=Pt(11))

# ═══════════════════════════════════════════════════════════════
# REVISION HISTORY (文档修改记录表)
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
run = p.add_run('文档修改记录')
set_run(run, font_name='黑体', size=Pt(14), bold=True)

make_table(doc,
    ['版本号', '修改日期', '修改人', '摘要', '备注'],
    [['V1.0', '2026-06-04', '', '初始版本，合并测试计划与测试说明', '']],
    col_widths=[Cm(2), Cm(2.5), Cm(2), Cm(6), Cm(3)])

# ═══════════════════════════════════════════════════════════════
# 1 范围
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '1  范围', 1)

add_heading(doc, '1.1  标识', 2)
add_para(doc, '本文档适用的系统为：')
for item in [
    'a)  名称：DARPA智能问答服务工具',
    'b)  标识：DARPA-IQAS',
    'c)  简称：DARPA问答工具',
    'd)  版本号：V1.0',
]:
    add_para(doc, item, indent=True)

add_para(doc, '本系统属于研究内容四"DARPA智能问答服务工具开发"，联合军事科学院军事科学信息研究中心共同开展。系统基于"外挂知识库—RAG检索增强—交互式提示"三级架构设计，融合结构化知识管理与检索增强生成技术，打造具备高精度领域适应能力的离线智能问答系统。', indent=True)

add_heading(doc, '1.2  系统概述', 2)
add_para(doc, 'DARPA智能问答服务工具是面向军事科研人员的离线智能问答系统，核心目标为突破多源异构数据整合瓶颈，实现军事文档的深度加工、精准检索和智能问答。系统采用三级架构设计：')
for item in [
    '第一级——外挂知识库模块（M1）：对非结构化军事文档进行解析、智能分块、元数据标注，支持PDF/Word/Excel/TXT/图片等多格式文档管理；',
    '第二级——RAG文档检索增强模块（M2）：基于ragflow v0.18.0成熟框架领域适配，构建向量搜索、混合检索、语义重排序、知识图谱等多维度检索能力；',
    '第三级——交互式提示词工程模块（M3）：通过动态模板引擎与结构化约束机制，实现聊天助手管理、多轮上下文对话、流式响应、引用溯源等能力。',
]:
    add_para(doc, item, bullet=True)

add_para(doc, '技术栈包括：Spring Boot后端、Vue 3前端、ragflow v0.18.0 RAG引擎、智谱GLM-9B本地部署LLM、MySQL 8.0、Elasticsearch 8.11.3、MinIO对象存储、Redis/Valkey缓存。部署模式为Docker Compose容器化离线部署，无外网依赖。', indent=True)

add_heading(doc, '1.3  文档概述', 2)
add_para(doc, '本文档是测试计划与测试说明的合并版，包含以下内容：')
for item in [
    'a) 测试环境、测试标识、测试准备等计划性内容（第2-5章）；',
    'b) 按三级架构组织的48条测试用例详细说明（第6章，核心章节）；',
    'c) 测试进度、需求可追踪性等管理性内容（第7-9章）。',
]:
    add_para(doc, item, indent=True)
add_para(doc, '本文档用于指导系统测试的执行，为《系统测试报告》提供依据。', indent=True)

# ═══════════════════════════════════════════════════════════════
# 2 引用文档 (模板表1)
# ═══════════════════════════════════════════════════════════════
add_heading(doc, '2  引用文档', 1)
add_para(doc, '表 1  引用文档')

make_table(doc,
    ['序号', '文档标识', '文档名称', '版本', '来源'],
    [
        ['1', '', 'DARPA智能问答服务工具-软件需求规格说明书', 'V1.0', '本项目'],
        ['2', '', 'DARPA智能问答服务工具-软件设计说明书', 'V1.0', '本项目'],
        ['3', '', 'ragflow v0.18.0 官方API文档', 'v0.18.0', 'infiniflow'],
        ['4', 'GJB 438B', '军用软件文档通用要求', '', '国军标'],
        ['5', '', 'pytest官方文档', '8.x', 'pytest.org'],
        ['6', '', 'Playwright官方文档', '1.x', 'playwright.dev'],
    ],
    col_widths=[Cm(1), Cm(2.5), Cm(6.5), Cm(1.5), Cm(2.5)])

# ═══════════════════════════════════════════════════════════════
# 3 软件测试环境 (测试计划模板 Ch3)
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '3  软件测试环境', 1)

# 3.1 软件项 (模板表2)
add_heading(doc, '3.1  软件项', 2)
add_para(doc, '表 2  软件项')

make_table(doc,
    ['序号', '名称', '版本标识', '用途', '备注'],
    [
        ['1', 'ragflow', 'v0.18.0', 'RAG引擎核心', 'Docker镜像'],
        ['2', 'Spring Boot (gaisoft-mes)', '2.x', '应用服务后端', 'Docker镜像'],
        ['3', 'Vue 3 (gaisoft-ui)', '3.x', '前端展示', 'Docker镜像'],
        ['4', 'MySQL', '8.0.39', '关系数据库', 'Docker镜像'],
        ['5', 'Elasticsearch', '8.11.3', '向量索引+全文检索', 'Docker镜像'],
        ['6', 'MinIO', 'latest', '文档对象存储', 'Docker镜像'],
        ['7', 'Redis/Valkey', '8', '缓存服务', 'Docker镜像'],
        ['8', '智谱GLM-9B', '-', '本地部署LLM', '离线模型'],
        ['9', 'Docker Engine', '24+', '容器运行时', '宿主机安装'],
        ['10', 'Docker Compose', 'v2', '服务编排', '宿主机安装'],
        ['11', 'pytest', '8.x', '测试执行框架', 'test-runner容器'],
        ['12', 'Playwright', '1.x', 'UI自动化测试', 'test-runner容器'],
    ],
    col_widths=[Cm(1), Cm(4), Cm(2), Cm(4), Cm(2.5)])

# 3.2 硬件及固件项 (模板表3)
add_heading(doc, '3.2  硬件及固件项', 2)
add_para(doc, '表 3  硬件及固件项')

make_table(doc,
    ['序号', '设备类型', '名称', '型号', '用途', '使用时间', '备注'],
    [
        ['1', '服务器', '测试服务器', '8核CPU/16GB/500GB', '承载全部Docker服务', '测试全程', ''],
        ['2', '客户端', '测试PC', '4核/8GB', '浏览器访问前端', 'UI测试', ''],
        ['3', '通信设备', '局域网交换机', '千兆', '服务互联', '测试全程', ''],
    ],
    col_widths=[Cm(1), Cm(1.5), Cm(2), Cm(3), Cm(3), Cm(2), Cm(2)])

# 3.3 其它材料 (模板表4)
add_heading(doc, '3.3  其它材料', 2)
add_para(doc, '表 4  其它材料')

make_table(doc,
    ['其他材料', '名称', '用途', '备注'],
    [
        ['军事文档样本', '雷达系统技术评估报告等', '知识库测试数据', '中文军事技术文档'],
        ['英文技术报告', 'DARPA雷达信号处理技术报告', '跨语言检索测试数据', '英文DARPA文档'],
        ['装备参数规范表', 'XLSX格式装备参数', '表格解析测试数据', '含10条装备参数'],
        ['测试数据生成器', 'test_data_factory.py', '自动生成多格式测试文档', 'PDF/DOCX/XLSX/TXT/MD'],
        ['QA评估对', '20组标准问答对', '检索准确率评估', '覆盖多领域'],
    ],
    col_widths=[Cm(3), Cm(4.5), Cm(4), Cm(3)])

# 3.4
add_heading(doc, '3.4  特定性质、需方权利与许可证', 2)
add_para(doc, '表 5  专有权性质、需方权利与许可证')

make_table(doc,
    ['元素名称', '特定性质', '需方权利', '许可证'],
    [
        ['ragflow v0.18.0', '开源RAG框架', 'Apache 2.0', '开源许可'],
        ['智谱GLM-9B', '本地部署LLM', '授权使用', '模型许可'],
    ],
    col_widths=[Cm(3), Cm(4), Cm(4), Cm(3)])

# 3.5
add_heading(doc, '3.5  安装、测试与控制', 2)
add_para(doc, '测试环境的安装、测试与控制说明如下：', indent=True)
for item in [
    'a) 通过 start.sh/ps1 <project> 启动测试环境全部Docker服务；',
    'b) test-runner容器内置pytest框架，执行 pytest --alluredir=allure-results 自动运行全部48条测试用例；',
    'c) 测试结果通过Allure报告展示，支持按模块、级别、标记筛选查看。',
]:
    add_para(doc, item, indent=True)

# 3.6 参与组织 (模板表6)
add_heading(doc, '3.6  参与组织', 2)
add_para(doc, '表 6  参与组织')

make_table(doc,
    ['组织名', '角色', '职责'],
    [
        ['开发方', '开发与测试执行', '系统开发、测试用例设计、测试执行、缺陷修复'],
        ['军事科学院军事科学信息研究中心', '需方/验收方', '需求确认、验收评审'],
    ],
    col_widths=[Cm(5), Cm(3), Cm(6)])

# 3.7 人员 (模板表7)
add_heading(doc, '3.7  人员', 2)
add_para(doc, '表 7  人员')

make_table(doc,
    ['人员数量', '数量', '技能水平', '要求培训和持续时间', '特殊要求'],
    [
        ['测试负责人', '1', '高级', '1天', '熟悉三级架构与RAG技术'],
        ['自动化测试工程师', '1', '中高级', '1天', '熟悉pytest/Playwright'],
        ['系统测试人员', '2', '中级', '2天', '熟悉Docker与Web测试'],
    ],
    col_widths=[Cm(3), Cm(1.5), Cm(2), Cm(4), Cm(4)])

# 3.8
add_heading(doc, '3.8  定向计划', 2)
add_para(doc, '测试前需完成以下培训：测试环境搭建培训、测试用例执行培训、Allure报告解读培训。', indent=True)

# ═══════════════════════════════════════════════════════════════
# 4 测试标识 (测试计划模板 Ch4)
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '4  测试标识', 1)

# 4.1 一般信息
add_heading(doc, '4.1  一般信息', 2)

# 4.1.1 测试级 (模板表8)
add_heading(doc, '4.1.1  测试级', 3)
add_para(doc, '表 8  测试级')

make_table(doc,
    ['测试级名称', '测试级描述'],
    [['系统测试', '对DARPA智能问答服务工具进行全面的功能测试、接口测试和端到端测试，验证系统是否满足需求规格说明书规定的全部要求']],
    col_widths=[Cm(3), Cm(11)])

# 4.1.2 测试类别 (模板表9)
add_heading(doc, '4.1.2  测试类别', 3)
add_para(doc, '表 9  测试类别')

make_table(doc,
    ['测试类别名称', '测试类别描述'],
    [
        ['功能测试', '验证三大模块（M1/M2/M3）的功能正确性，共38条用例'],
        ['性能测试', '验证检索响应时间、并发能力等性能指标，含基准测试'],
        ['接口测试', '验证ragflow API及应用服务API的接口正确性'],
        ['UI测试', '验证前端界面的功能可用性和交互正确性，共5条用例'],
        ['端到端测试', '验证完整业务流程的正确性，共4条用例'],
    ],
    col_widths=[Cm(3), Cm(11)])

# 4.1.3 一般测试条件 (模板表10)
add_heading(doc, '4.1.3  一般测试条件', 3)
add_para(doc, '表 10  一般测试条件')

make_table(doc,
    ['条件项', '描述'],
    [
        ['测试环境', '离线Docker Compose部署，所有服务通过ragflow桥接网络互联'],
        ['测试数据', 'DARPA军事文档样本（中文雷达报告、通信装备手册、英文技术报告、装备参数规范表）'],
        ['LLM模型', '智谱GLM-9B模型本地部署运行，无外网访问'],
        ['前置检查', '测试执行前需确保所有Docker服务健康就绪（健康检查通过）'],
        ['数据隔离', '每条测试用例使用独立的测试数据集，互不干扰'],
    ],
    col_widths=[Cm(3), Cm(11)])

# 4.1.4 测试进程
add_heading(doc, '4.1.4  测试进程', 3)
add_para(doc, '测试按pytest自动化流程执行：', indent=True)
for item in [
    'a) 环境初始化：启动所有Docker服务，等待健康检查通过（ragflow 120秒超时）；',
    'b) 会话级Fixture准备：创建全局测试数据集、上传并解析测试文档、创建测试聊天助手和会话；',
    'c) 按模块执行：pytest按module1→module2→module3→module4→ui→e2e顺序执行；',
    'd) 结果收集：Allure框架自动收集测试结果和度量数据。',
]:
    add_para(doc, item, indent=True)

# 4.1.5 数据记录、归约和分析
add_heading(doc, '4.1.5  数据记录、归约和分析', 3)
for item in [
    'a) 测试结果通过Allure框架自动记录，包括通过/失败/跳过状态；',
    'b) 性能基准测试记录响应延迟数据，计算平均值和P95值；',
    'c) 检索准确率测试记录命中率，计算Recall指标；',
    'd) 所有度量数据自动归档至Allure报告。',
]:
    add_para(doc, item, indent=True)

# 4.2 计划执行的测试
add_heading(doc, '4.2  计划执行的测试', 2)

# 4.2.1 功能测试 (模板表11)
add_heading(doc, '4.2.1  功能测试', 3)
add_para(doc, '表 11  功能测试信息')

make_table(doc,
    ['测试范围', '测试级', '测试类型', '对应需求', '数据记录'],
    [
        ['M1外挂知识库', '系统测试', '功能测试', 'M1-REQ-001~008', '文档ID、分块数、解析状态'],
        ['M2 RAG检索增强', '系统测试', '功能测试', 'M2-REQ-001~008', '检索结果、相似度、延迟'],
        ['M3交互式提示', '系统测试', '功能测试', 'M3-REQ-001~008', '回答内容、引用信息、流式数据'],
        ['集成验证', '系统测试', '功能+接口测试', 'INT-REQ-001~004', '会话数据、SSE流、Token'],
        ['UI界面', '系统测试', '功能测试', 'M1~M3 UI需求', '页面元素、交互结果'],
        ['端到端', '系统测试', '功能+性能测试', 'INT-REQ-005~010', '全流程数据'],
    ],
    col_widths=[Cm(2.5), Cm(1.5), Cm(2.5), Cm(3), Cm(4)])

# 4.2.2 性能测试 (模板表12)
add_heading(doc, '4.2.2  性能测试', 3)
add_para(doc, '表 12  性能测试信息')

make_table(doc,
    ['测试名称', '测试目的', '对应需求'],
    [
        ['检索性能基准', '100次检索请求，P95延迟<5s', '3.10.4'],
        ['阈值调节性能', '不同相似度阈值的检索效率', 'M2-REQ-003'],
    ],
    col_widths=[Cm(3.5), Cm(6), Cm(3)])

# 4.2.3 接口测试 (模板表13)
add_heading(doc, '4.2.3  接口测试', 3)
add_para(doc, '表 13  接口测试信息')

make_table(doc,
    ['测试名称', '测试目的', '对应需求'],
    [
        ['ragflow REST API', '知识库/文档/对话接口正确性', '3.3.2'],
        ['应用服务 REST API', '知识管理/聊天/认证接口', '3.3.3'],
        ['OpenAI兼容接口', '/chats_openai端点兼容性', 'M3-REQ-007'],
        ['Stream Proxy', 'SSE流式代理转发正确性', 'INT-REQ-003'],
    ],
    col_widths=[Cm(3.5), Cm(6), Cm(3)])

# ═══════════════════════════════════════════════════════════════
# 5 测试准备 (测试说明模板 Ch3)
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '5  测试准备', 1)

add_heading(doc, '5.1  硬件准备', 2)
add_para(doc, '表 14  硬件设备要求')

make_table(doc,
    ['序号', '硬件项目', '硬件要求', '说明'],
    [
        ['1', '测试服务器', '8核CPU/16GB内存/500GB硬盘', '离线环境Linux服务器'],
        ['2', '网络环境', '局域网TCP/IP', '无需外网接入'],
    ],
    col_widths=[Cm(1), Cm(3), Cm(5.5), Cm(4)])

add_heading(doc, '5.2  软件准备', 2)
add_para(doc, '表 15  软件准备要求')

make_table(doc,
    ['序号', '软件项目', '软件要求', '说明'],
    [
        ['1', 'Docker Engine', '24+', '容器运行时'],
        ['2', 'Docker Compose', 'v2', '服务编排'],
        ['3', '离线镜像包', '全部服务镜像', '通过offline-load加载'],
        ['4', '智谱GLM-9B模型', '本地部署', 'LLM推理服务'],
    ],
    col_widths=[Cm(1), Cm(3.5), Cm(4.5), Cm(4.5)])

add_heading(doc, '5.3  其它测试前准备', 2)
for item in [
    'a) 配置测试环境变量：RAGFLOW_BASE_URL、RAGFLOW_API_KEY、GAISOFT_API_URL、GAISOFT_FRONTEND_URL；',
    'b) 准备测试账号：gaisoft登录用户名/密码；',
    'c) 准备军事文档测试样本：雷达报告、通信手册、装备规范表等；',
    'd) 确认智谱GLM-9B模型加载完成。',
]:
    add_para(doc, item, indent=True)

add_heading(doc, '5.4  测试环境组成', 2)
add_para(doc, '测试环境架构：全部服务运行在Docker Compose编排的离线容器环境中，包括ragflow引擎、Spring Boot后端、Vue 3前端、MySQL、Elasticsearch、MinIO、Redis、智谱GLM-9B。test-runner容器内置pytest和Playwright，通过ragflow桥接网络访问各服务。', indent=True)

# ═══════════════════════════════════════════════════════════════
# 6 测试说明 (核心章节，测试说明模板 Ch4)
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '6  测试说明', 1)
add_para(doc, '本章为核心章节，按三级架构组织48条测试用例，每条用例描述测试目的、前置条件、测试步骤和预期结果。用例逻辑全部从pytest自动化测试代码提取。', indent=True)

# ──────────────────────────────────────────────────────────────
# 6.1 M1
# ──────────────────────────────────────────────────────────────
add_heading(doc, '6.1  第一级——外挂知识库功能测试（M1模块）', 2)
add_para(doc, '本节测试第一级"外挂知识库模块"的功能，涵盖知识库CRUD、文档上传、文档解析、分块管理和元数据过滤，共21条测试用例。', indent=True)

add_heading(doc, '6.1.1  知识库CRUD测试', 3)

add_test_case_table(doc, 'TC-M1-001', '创建知识库',
    '验证通过ragflow API创建知识库，返回有效的知识库ID和名称',
    'ragflow服务运行正常，API可访问',
    '1) 调用POST /api/v1/datasets，参数：name="darpa_test_xxx"，chunk_method="naive"\n2) 检查返回数据中包含有效id字段\n3) 验证返回的name与请求一致',
    '创建成功，返回dataset对象包含有效id，name与请求一致',
    'M1-REQ-001')

add_test_case_table(doc, 'TC-M1-002', '指定嵌入模型创建知识库',
    '验证创建知识库时指定嵌入模型（BAAI/bge-large-zh-v1.5）',
    'ragflow服务运行正常',
    '1) 调用POST /api/v1/datasets，参数：name="darpa_emb_xxx"，embedding_model="BAAI/bge-large-zh-v1.5"\n2) 检查返回dataset id',
    '创建成功，返回有效id',
    'M1-REQ-001')

add_test_case_table(doc, 'TC-M1-003', '列出知识库',
    '验证列出所有知识库返回有效的列表结构',
    'ragflow服务运行正常',
    '1) 调用GET /api/v1/datasets\n2) 验证返回值为list类型',
    '返回有效的知识库列表（list类型）',
    'M1-REQ-001')

add_test_case_table(doc, 'TC-M1-004', '级联删除知识库',
    '验证删除知识库时级联删除其下所有文档和分块',
    '知识库已创建，已上传文档并完成解析',
    '1) 创建知识库cascade_test_xxx\n2) 上传并解析文档，等待完成\n3) 验证文档存在\n4) 调用DELETE /api/v1/datasets删除\n5) 列出知识库验证已删除',
    '删除后知识库不再出现，关联文档和分块被级联删除',
    'M1-REQ-001')

add_heading(doc, '6.1.2  文档上传测试', 3)

add_test_case_table(doc, 'TC-M1-005', '上传PDF文档',
    '验证上传PDF格式文档到知识库，返回有效文档ID',
    '知识库已创建',
    '1) 生成PDF格式的雷达系统技术评估报告\n2) 调用上传接口上传PDF\n3) 检查返回文档id',
    '上传成功，返回有效文档ID',
    'M1-REQ-002')

add_test_case_table(doc, 'TC-M1-006', '上传DOCX文档',
    '验证上传Word格式（.docx）通信装备操作手册',
    '知识库已创建',
    '1) 生成DOCX格式的通信装备操作手册\n2) 上传到知识库',
    '上传成功，返回有效文档ID',
    'M1-REQ-002')

add_test_case_table(doc, 'TC-M1-007', '上传XLSX文档',
    '验证上传Excel格式（.xlsx）装备参数规范表',
    '知识库已创建',
    '1) 生成XLSX格式装备参数规范表（含10条参数）\n2) 上传到知识库',
    '上传成功，返回有效文档ID',
    'M1-REQ-002')

add_test_case_table(doc, 'TC-M1-008', '上传TXT文档',
    '验证上传纯文本格式（.txt）野战维护指南',
    '知识库已创建',
    '1) 生成TXT格式野战维护指南\n2) 上传到知识库',
    '上传成功，返回有效文档ID',
    'M1-REQ-002')

add_test_case_table(doc, 'TC-M1-009', '上传MD文档',
    '验证上传Markdown格式（.md）军事装备管理条例',
    '知识库已创建',
    '1) 生成MD格式军事装备管理条例\n2) 上传到知识库',
    '上传成功，返回有效文档ID',
    'M1-REQ-002')

add_test_case_table(doc, 'TC-M1-010', '大文件上传',
    '验证上传大文件（约5MB军事装备测试数据）可正常处理',
    '知识库已创建',
    '1) 生成约5MB大文本文件\n2) 上传到知识库',
    '大文件上传成功，返回有效文档ID',
    'M1-REQ-002')

add_test_case_table(doc, 'TC-M1-011', '批量文档上传',
    '验证一次批量上传20个文档',
    '知识库已创建',
    '1) 生成20个测试文档\n2) 调用批量上传接口\n3) 验证返回文档数量',
    '批量上传成功，返回20个文档记录',
    'M1-REQ-002')

add_heading(doc, '6.1.3  文档解析测试', 3)

add_test_case_table(doc, 'TC-M1-012', '通用分块解析（naive）',
    '验证naive分块方法按固定token数分块（chunk_token_num=512）',
    '知识库已创建（chunk_method="naive"），文档已上传',
    '1) 创建naive知识库\n2) 上传野战维护指南TXT\n3) 触发解析等待完成（300秒超时）\n4) 验证chunk_num > 0',
    '解析完成，生成有效分块（chunk_num > 0）',
    'M1-REQ-003, M1-REQ-004')

add_test_case_table(doc, 'TC-M1-013', '书籍分块解析（book）',
    '验证book分块方法保留章节结构',
    '知识库已创建（chunk_method="book"）',
    '1) 创建book知识库\n2) 上传通信装备操作手册DOCX\n3) 触发解析等待完成',
    '书籍解析完成，保留章节结构',
    'M1-REQ-003, M1-REQ-004')

add_test_case_table(doc, 'TC-M1-014', '表格分块解析（table）',
    '验证table分块方法保留表格结构',
    '知识库已创建（chunk_method="table"）',
    '1) 创建table知识库\n2) 上传装备参数规范表XLSX\n3) 触发解析等待完成',
    '表格解析完成，保留表格行列结构',
    'M1-REQ-003, M1-REQ-004')

add_test_case_table(doc, 'TC-M1-015', '论文分块解析（paper）',
    '验证paper分块方法分离摘要/正文/参考文献',
    '知识库已创建（chunk_method="paper"）',
    '1) 创建paper知识库\n2) 上传DARPA雷达论文TXT\n3) 触发解析等待完成',
    '论文解析完成，正确分离论文结构',
    'M1-REQ-003, M1-REQ-004')

add_test_case_table(doc, 'TC-M1-016', '多源异构数据整合',
    '验证同一知识库同时导入中英文文档和表格无冲突',
    '知识库已创建（chunk_method="naive"）',
    '1) 创建知识库\n2) 同时上传中文雷达报告、英文技术报告、装备参数规范表\n3) 触发全部解析\n4) 等待完成',
    '全部文档解析成功，无冲突',
    'M1-REQ-008')

add_heading(doc, '6.1.4  分块管理测试', 3)

add_test_case_table(doc, 'TC-M1-017', '手动添加分块',
    '验证手动添加新分块，含关键词标注',
    '知识库已创建，文档已解析完成',
    '1) 调用POST .../chunks，content含关键词\n2) 验证返回包含chunk或id字段',
    '分块创建成功，返回包含chunk_id或id',
    'M1-REQ-005')

add_test_case_table(doc, 'TC-M1-018', '列出文档分块',
    '验证列出文档中的所有分块',
    '文档已完成解析',
    '1) 调用GET .../chunks\n2) 获取列表验证非空',
    '返回分块列表，数量大于0',
    'M1-REQ-005')

add_test_case_table(doc, 'TC-M1-019', '更新分块内容',
    '验证修改已有分块内容',
    '分块已存在',
    '1) 添加分块"原始内容待更新"\n2) PUT更新为"装备维护规程v2.0"\n3) 验证返回非空',
    '更新成功，返回确认',
    'M1-REQ-005')

add_test_case_table(doc, 'TC-M1-020', '删除分块',
    '验证删除指定分块',
    '分块已存在',
    '1) 添加分块"待删除的测试分块"\n2) DELETE删除\n3) 验证返回非空',
    '删除成功，返回确认',
    'M1-REQ-005')

add_heading(doc, '6.1.5  元数据过滤测试', 3)

add_test_case_table(doc, 'TC-M1-021', '元数据过滤检索',
    '验证设置文档元数据后可通过元数据过滤检索',
    '知识库中已上传并解析多个文档',
    '1) 创建知识库上传两个文档\n2) 解析完成\n3) 检索"雷达维护周期"\n4) 验证返回分块 > 0',
    '检索返回雷达相关分块，数量>0',
    'M1-REQ-006')

# ──────────────────────────────────────────────────────────────
# 6.2 M2
# ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6.2  第二级——RAG检索增强功能测试（M2模块）', 2)
add_para(doc, '本节测试第二级"RAG文档检索增强模块"的功能，涵盖向量搜索、混合检索、相似度阈值、重排序、知识图谱、跨语言检索和检索准确率，共15条测试用例。', indent=True)

add_heading(doc, '6.2.1  向量搜索测试', 3)

add_test_case_table(doc, 'TC-M2-001', '基础语义检索',
    '验证中文语义查询"雷达维护周期"返回相关分块，相似度>0.1',
    '测试数据集已创建，文档已上传并解析完成',
    '1) 调用POST /api/v1/retrieval，question="雷达维护周期"\n2) 获取chunks\n3) 验证chunks>0且首块similarity>=0.1',
    '返回至少1个分块，最高相似度>=0.1',
    'M2-REQ-001')

add_test_case_table(doc, 'TC-M2-002', '装备参数查询',
    '验证查询装备参数"ZBD-2000通信系统频率范围"返回相关结果',
    '测试数据集已准备',
    '1) 检索"ZBD-2000通信系统频率范围"\n2) 验证返回分块>0',
    '返回与通信系统频率相关的分块',
    'M2-REQ-001')

add_test_case_table(doc, 'TC-M2-003', '无关查询低相似度',
    '验证与军事文档无关的查询返回低相似度或空结果',
    '测试数据集已准备',
    '1) 检索"今天天气怎么样"，similarity_threshold=0.5\n2) 若返回分块，验证首块similarity<0.5',
    '返回空结果或低相似度（<0.5）',
    'M2-REQ-007')

add_test_case_table(doc, 'TC-M2-004', '长查询检索',
    '验证200+字符的长查询仍能正确检索',
    '测试数据集已准备',
    '1) 构造200+字符长查询（雷达发射机故障排查）\n2) 执行检索\n3) 验证返回分块>0',
    '长查询返回相关分块',
    'M2-REQ-001')

add_heading(doc, '6.2.2  混合搜索测试', 3)

add_test_case_table(doc, 'TC-M2-005', '向量+关键词混合检索',
    '验证混合检索模式（vector_similarity_weight=0.5, keyword=true）',
    '测试数据集已准备',
    '1) 检索"雷达ERR-001故障"，vector_similarity_weight=0.5，keyword=true\n2) 验证分块>0',
    '混合检索返回故障相关分块',
    'M2-REQ-002')

add_test_case_table(doc, 'TC-M2-006', 'top_k参数验证',
    '验证不同top_k参数影响返回结果数量',
    '测试数据集已准备',
    '1) 分别以top_k=10和top_k=50检索"装备维护"\n2) 比较分块数',
    'top_k=50返回数>=top_k=10',
    'M2-REQ-002')

add_test_case_table(doc, 'TC-M2-007', '跨知识库检索',
    '验证跨多个知识库检索返回不同知识库的结果',
    '两个独立知识库，分别含不同文档',
    '1) 创建ds1、ds2两个知识库\n2) 分别上传和解析\n3) 以dataset_ids=[ds1,ds2]检索\n4) 验证分块>0',
    '跨库检索返回来自多个知识库的分块',
    'M2-REQ-002')

add_heading(doc, '6.2.3  相似度阈值测试', 3)

add_test_case_table(doc, 'TC-M2-008', '阈值递增调节',
    '验证提高相似度阈值（0.1→0.7）后返回结果数量递减',
    '测试数据集已准备',
    '1) 分别设0.1、0.3、0.5、0.7\n2) 每个阈值检索"雷达维护"\n3) 记录分块数',
    '阈值0.1结果数>=阈值0.7结果数',
    'M2-REQ-003')

add_test_case_table(doc, 'TC-M2-009', '极高阈值验证',
    '验证极高阈值（0.9）仅返回极少量结果',
    '测试数据集已准备',
    '1) 设similarity_threshold=0.9\n2) 检索"装备检查"',
    '返回分块数<=3',
    'M2-REQ-003')

add_heading(doc, '6.2.4  重排序测试', 3)

add_test_case_table(doc, 'TC-M2-010', '重排序效果验证',
    '验证重排序参数可正常传递',
    '测试数据集已准备',
    '1) 不启用重排序检索\n2) 启用重排序检索\n3) 验证两次均返回有效结果',
    '两次检索均返回有效分块',
    'M2-REQ-004')

add_heading(doc, '6.2.5  知识图谱测试', 3)

add_test_case_table(doc, 'TC-M2-011', 'GraphRAG检索',
    '验证GraphRAG构建知识图谱后可通过图谱检索',
    '知识库已创建并解析完成',
    '1) 上传雷达报告并解析\n2) 调用run_graphrag\n3) 等待构建\n4) 以use_kg=true检索\n5) 验证分块>0',
    '知识图谱检索返回关联分块（未配置则跳过）',
    'M2-REQ-005')

add_heading(doc, '6.2.6  跨语言检索测试', 3)

add_test_case_table(doc, 'TC-M2-012', '中文查询英文文档',
    '验证中文查询通过cross_languages找到英文文档',
    '双语知识库已创建并解析',
    '1) 以cross_languages=true检索"雷达信号处理算法"\n2) 验证分块>0',
    '中文查询返回英文文档相关分块',
    'M2-REQ-006')

add_test_case_table(doc, 'TC-M2-013', '英文查询中文文档',
    '验证英文查询可找到中文文档',
    '双语知识库已准备',
    '1) 以cross_languages=true检索"radar maintenance interval hours"\n2) 验证分块>0',
    '英文查询返回中文文档分块',
    'M2-REQ-006')

add_heading(doc, '6.2.7  检索准确率测试', 3)

add_test_case_table(doc, 'TC-M2-014', '检索精确率/召回率评估',
    '使用20组标准QA对评估检索Recall>=60%',
    '测试数据集已准备，20组QA对已定义',
    '1) 对每组QA对检索top-5\n2) 检查是否含answer或keywords\n3) 计算命中率',
    'Recall>=60%',
    'M2-REQ-007')

add_test_case_table(doc, 'TC-M2-015', '检索性能基准测试',
    '100次检索请求，P95延迟<5秒',
    '测试数据集已准备',
    '1) 循环100次检索\n2) 记录每次延迟\n3) 计算P95',
    'P95延迟<5秒',
    '3.10.4')

# ──────────────────────────────────────────────────────────────
# 6.3 M3
# ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6.3  第三级——交互式提示功能测试（M3模块）', 2)
add_para(doc, '本节测试第三级"交互式提示词工程模块"的功能，涵盖聊天助手管理、系统提示词、多轮对话、流式响应、引用溯源和提示模板，共12条测试用例。', indent=True)

add_heading(doc, '6.3.1  聊天助手测试', 3)

add_test_case_table(doc, 'TC-M3-001', '创建聊天助手并绑定知识库',
    '验证创建聊天助手并绑定数据集，尝试配置系统提示词',
    '测试数据集已创建',
    '1) POST /api/v1/chats创建助手，绑定dataset_ids\n2) 验证返回有效id和name\n3) 尝试PUT设置prompt_config',
    '聊天助手创建成功，id有效',
    'M3-REQ-001')

add_test_case_table(doc, 'TC-M3-002', '单轮问答',
    '验证单轮问答返回非空答案和引用信息',
    '测试聊天助手和会话已创建',
    '1) 提问"AN/TPQ-53雷达的探测距离是多少？"\n2) 验证answer>0\n3) 验证reference非空',
    '返回非空答案，包含引用信息',
    'M3-REQ-001')

add_test_case_table(doc, 'TC-M3-003', '会话CRUD管理',
    '验证会话的创建、列表查询和删除',
    '聊天助手已创建',
    '1) 创建会话验证id有效\n2) 列出会话验证list类型\n3) 删除会话',
    '会话CRUD全部成功',
    'M3-REQ-001')

add_heading(doc, '6.3.2  系统提示词测试', 3)

add_test_case_table(doc, 'TC-M3-004', '领域约束提示词',
    '验证军事装备领域约束提示词拒绝无关问题',
    '测试数据集已准备',
    '1) 设置提示词"你是DARPA装备分析专家"\n2) 提问"推荐一部好看的电影"\n3) 检查拒绝关键词',
    '对无关问题给出拒绝回复或极短回复',
    'M3-REQ-002')

add_test_case_table(doc, 'TC-M3-005', '知识变量注入',
    '验证{knowledge}变量被替换为检索结果',
    '测试聊天助手和会话已创建',
    '1) 提问"雷达系统技术参数有哪些？"\n2) 验证回答>20字符',
    '回答含知识库雷达参数信息',
    'M3-REQ-006')

add_test_case_table(doc, 'TC-M3-006', '结构化JSON输出',
    '验证通过提示词约束LLM输出JSON格式',
    '测试数据集已准备',
    '1) 设置提示词要求JSON输出\n2) 提问"雷达探测距离是多少？请用JSON格式回答。"\n3) 提取JSON片段',
    '回答中包含可解析的JSON片段',
    'M3-REQ-007')

add_heading(doc, '6.3.3  多轮对话测试', 3)

add_test_case_table(doc, 'TC-M3-007', '多轮上下文对话',
    '验证4轮对话第4轮引用第1轮上下文',
    '测试数据集已准备',
    '1) 第1轮："AN/TPQ-53雷达的工作频段是什么？"\n2) 第2轮："它的探测距离呢？"\n3) 第3轮："故障代码ERR-001怎么处理？"\n4) 第4轮："刚才说的那个频段的雷达，维护周期是多久？"\n5) 验证4轮均非空',
    '4轮均有有效回答，第4轮引用第1轮上下文',
    'M3-REQ-003')

add_test_case_table(doc, 'TC-M3-008', '温度参数效果',
    '验证不同temperature（0.0 vs 1.0）产生不同输出',
    '测试数据集已准备',
    '1) 创建temperature=0.0的chat_low\n2) 创建temperature=1.0的chat_high\n3) 分别提问验证均非空',
    '两个温度的助手均返回有效回答',
    'M3-REQ-007')

add_heading(doc, '6.3.4  流式响应测试', 3)

add_test_case_table(doc, 'TC-M3-009', 'SSE流式输出',
    '验证stream=true返回SSE数据块且流正确终止',
    '测试聊天助手和会话已创建',
    '1) 调用completions，stream=true\n2) 收集SSE chunks\n3) 验证chunks>0\n4) 验证包含终止标记',
    '返回多个SSE数据块，流正确终止',
    'M3-REQ-004')

add_heading(doc, '6.3.5  引用溯源测试', 3)

add_test_case_table(doc, 'TC-M3-010', '引用信息验证',
    '验证引用信息包含chunk_id、similarity等字段',
    '测试聊天助手和会话已创建',
    '1) 提问"雷达系统技术参数有哪些？"\n2) 获取reference\n3) 验证含chunk_id和similarity',
    '引用信息含chunk_id和similarity字段',
    'M3-REQ-005')

add_test_case_table(doc, 'TC-M3-011', '空知识库响应',
    '验证对完全无关问题返回有效响应结构',
    '至少一个知识库存在并已解析',
    '1) 提问"量子纠缠在星际旅行中的应用前景如何？"\n2) 验证answer为string类型',
    '返回有效的字符串类型回答',
    'M3-REQ-002')

add_heading(doc, '6.3.6  提示模板测试', 3)

add_test_case_table(doc, 'TC-M3-012', 'OpenAI兼容接口',
    '验证OpenAI兼容API端点返回OpenAI格式响应',
    '测试聊天助手已创建',
    '1) 调用POST /api/v1/chats_openai/{id}/chat/completions\n2) 验证返回含"choices"或"data"字段',
    '返回符合OpenAI格式的响应',
    'M3-REQ-006')

# ──────────────────────────────────────────────────────────────
# 6.4 集成
# ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6.4  集成验证测试', 2)
add_para(doc, '本节测试跨模块集成功能，包括知识库-会话绑定、知识库聊天、流式代理和认证集成，共6条测试用例。', indent=True)

add_heading(doc, '6.4.1  知识库会话绑定', 3)

add_test_case_table(doc, 'TC-INT-001', '会话持久化',
    '验证通过gaisoft-mes API创建的会话可持久保存',
    'gaisoft-mes服务运行正常',
    '1) GET /aftersales/session/list获取列表\n2) POST创建新会话\n3) 再次list验证出现',
    '会话创建后可在列表中查到',
    'INT-REQ-001')

add_test_case_table(doc, 'TC-INT-002', '多用户会话隔离',
    '验证两个聊天助手会话相互独立',
    '测试数据集已准备',
    '1) 创建chat1和chat2绑定同一数据集\n2) 分别创建会话\n3) 分别提问不同问题\n4) 验证回答互不干扰',
    '两个会话回答互不干扰',
    'INT-REQ-001')

add_heading(doc, '6.4.2  知识库聊天', 3)

add_test_case_table(doc, 'TC-INT-003', '知识库聊天记录',
    '验证gaisoft-mes API查询聊天和会话记录',
    'gaisoft-mes服务运行正常',
    '1) GET /aftersales/chat/list\n2) 验证返回有效数据\n3) GET /aftersales/session/list\n4) 验证非空',
    '聊天和会话列表均返回有效数据',
    'INT-REQ-002')

add_heading(doc, '6.4.3  流式代理', 3)

add_test_case_table(doc, 'TC-INT-004', 'SSE流代理转发',
    '验证gaisoft StreamProxyController正确转发ragflow SSE流',
    'ragflow和gaisoft-mes均运行正常',
    '1) POST /proxy/stream转发ragflow completions\n2) 收集SSE chunks\n3) 验证chunks>0且含data事件',
    '流代理正确转发SSE数据块',
    'INT-REQ-003')

add_heading(doc, '6.4.4  认证集成', 3)

add_test_case_table(doc, 'TC-INT-005', '认证Token缓存',
    '验证gaisoft-mes缓存ragflow认证Token',
    'gaisoft-mes服务运行正常',
    '1) 调用/ragflow/common代理ragflow healthz\n2) 再次调用\n3) 验证两次均成功',
    '两次请求均成功（Token复用）',
    'INT-REQ-004')

add_test_case_table(doc, 'TC-INT-006', '用户登录验证',
    '验证gaisoft-mes登录返回有效Token和用户信息',
    'gaisoft-mes服务运行正常',
    '1) 调用GET /getInfo\n2) 验证返回code=200或含user字段',
    '返回有效用户数据',
    'INT-REQ-004')

# ──────────────────────────────────────────────────────────────
# 6.5 UI
# ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6.5  UI界面测试', 2)
add_para(doc, '本节测试基于Playwright的前端界面功能，共5条测试用例。', indent=True)

add_heading(doc, '6.5.1  知识管理界面', 3)

add_test_case_table(doc, 'TC-UI-001', '知识库页面导航与渲染',
    '验证登录后知识库管理页面正常渲染',
    '前端服务运行正常，已通过Playwright登录',
    '1) 查找知识库导航链接\n2) 点击或直接访问"/#/kb"\n3) 等待页面加载\n4) 验证body可见',
    '知识库管理页面正常渲染',
    'M1-REQ-001')

add_test_case_table(doc, 'TC-UI-002', '知识库页面元素验证',
    '验证知识库页面包含预期UI元素',
    '知识库页面已加载',
    '1) 获取页面HTML\n2) 验证内容长度>100字符',
    '页面包含渲染内容',
    'M1-REQ-001')

add_heading(doc, '6.5.2  文档上传界面', 3)

add_test_case_table(doc, 'TC-UI-003', '文档上传UI元素',
    '验证知识库页面存在上传按钮或上传区域',
    '已登录并导航到知识库页面',
    '1) 导航到"/#/kb"\n2) 查找上传按钮、file input等元素\n3) 验证至少存在一种',
    '页面上存在上传相关UI元素',
    'M1-REQ-002')

add_heading(doc, '6.5.3  聊天交互界面', 3)

add_test_case_table(doc, 'TC-UI-004', '聊天交互与流式响应',
    '验证聊天界面可输入问题并发送，系统返回流式响应',
    '已登录',
    '1) 查找并点击对话导航\n2) 输入"雷达探测距离是多少？"\n3) 点击发送\n4) 等待10秒\n5) 验证页面未崩溃',
    '聊天界面正常交互，页面稳定',
    'M3-REQ-001, M3-REQ-004')

add_heading(doc, '6.5.4  RAG测试界面', 3)

add_test_case_table(doc, 'TC-UI-005', 'RAG检索测试页面',
    '验证RAG检索测试页面可正常渲染',
    '已登录',
    '1) 查找并点击"检索测试"等导航\n2) 或直接访问"/#/ragTest"\n3) 验证页面内容>100字符',
    'RAG测试页面正常渲染',
    'M2-REQ-007')

# ──────────────────────────────────────────────────────────────
# 6.6 E2E
# ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6.6  端到端验证测试', 2)
add_para(doc, '本节测试完整的端到端业务流程，验证从知识入库到智能问答的全链路，共4条测试用例。', indent=True)

add_heading(doc, '6.6.1  知识→答案全流程', 3)

add_test_case_table(doc, 'TC-E2E-001', '完整知识到答案流程',
    '验证完整流程：创建知识库→上传→解析→创建助手→提问→验证答案',
    'ragflow服务运行正常',
    '1) 创建知识库\n2) 上传雷达报告\n3) 解析等待完成\n4) 创建聊天助手\n5) 提问"AN/TPQ-53雷达的探测距离是多少？"\n6) 验证回答非空且含"60"或"公里"或"探测"',
    '全流程成功，回答含雷达探测距离（60公里）',
    'INT-REQ-007')

add_heading(doc, '6.6.2  多文档推理', 3)

add_test_case_table(doc, 'TC-E2E-002', '跨文档联合推理',
    '验证需要多个文档信息的跨文档推理',
    'ragflow服务运行正常',
    '1) 创建知识库，上传雷达报告+通信手册+装备规范表\n2) 解析全部\n3) 创建助手\n4) 提问"请对比AN/TPQ-53雷达和ZBD-2000通信系统的维护要求"\n5) 验证回答>30字符',
    '回答充实（>30字符），含两种装备维护信息',
    'INT-REQ-008')

add_heading(doc, '6.6.3  离线部署验证', 3)

add_test_case_table(doc, 'TC-E2E-003', '离线无外网依赖验证',
    '验证所有核心服务在离线环境下正常运行',
    '系统在离线Docker Compose中运行',
    '1) 调用ragflow健康检查\n2) 调用gaisoft /getInfo\n3) 调用ragflow dataset listing\n4) 验证全部正常',
    '全部服务离线可达且正常响应',
    'INT-REQ-005')

add_heading(doc, '6.6.4  数据安全验证', 3)

add_test_case_table(doc, 'TC-E2E-004', '数据隔离验证',
    '验证知识库A的文档不出现在知识库B中',
    'ragflow服务运行正常',
    '1) 创建知识库A和B\n2) 仅向A上传文档\n3) 查询B文档数=0\n4) 查询A文档数>0',
    'A有文档，B无文档，数据严格隔离',
    'INT-REQ-009')

# ═══════════════════════════════════════════════════════════════
# 7 测试进度表 (测试计划模板 Ch5)
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '7  测试进度表', 1)
add_para(doc, '表 16  测试进度')

make_table(doc,
    ['阶段', '活动', '预计耗时', '说明'],
    [
        ['1', '环境搭建与验证', '0.5天', 'Docker服务启动、健康检查、Fixture初始化'],
        ['2', 'M1外挂知识库测试（21条）', '1天', '含文档上传和解析等待时间'],
        ['3', 'M2 RAG检索增强测试（15条）', '1天', '含性能基准测试（100次检索）'],
        ['4', 'M3交互式提示测试（12条）', '1天', '含LLM响应等待时间'],
        ['5', '集成验证测试（6条）', '0.5天', '跨模块集成'],
        ['6', 'UI界面测试（5条）', '0.5天', 'Playwright自动化'],
        ['7', '端到端测试（4条）', '0.5天', '含全链路等待时间'],
        ['8', '测试报告编写', '1天', '整理Allure结果、编写报告'],
        ['总计', '', '6天', ''],
    ],
    col_widths=[Cm(1.5), Cm(5.5), Cm(2), Cm(6.5)])

# ═══════════════════════════════════════════════════════════════
# 8 需求可追踪性 (测试计划模板 Ch6)
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '8  需求的可追踪性', 1)

add_heading(doc, '8.1  正向追踪：需求→测试用例', 2)
add_para(doc, '表 17  正向追踪')

make_table(doc,
    ['需求编号', '需求名称', '对应测试用例'],
    [
        ['M1-REQ-001', '知识库创建/删除/修改/查询', 'TC-M1-001~004'],
        ['M1-REQ-002', '非结构化军事文档上传', 'TC-M1-005~011'],
        ['M1-REQ-003', '文档深度解析', 'TC-M1-012~016'],
        ['M1-REQ-004', '智能语义化分块', 'TC-M1-012~015'],
        ['M1-REQ-005', '分块管理与预览', 'TC-M1-017~020'],
        ['M1-REQ-006', '元数据标注与过滤', 'TC-M1-021'],
        ['M1-REQ-008', '多源异构数据整合', 'TC-M1-016'],
        ['M2-REQ-001', '向量语义检索', 'TC-M2-001~004'],
        ['M2-REQ-002', '混合检索', 'TC-M2-005~007'],
        ['M2-REQ-003', '相似度阈值动态调节', 'TC-M2-008~009'],
        ['M2-REQ-004', '检索结果重排序', 'TC-M2-010'],
        ['M2-REQ-005', '知识图谱辅助检索', 'TC-M2-011'],
        ['M2-REQ-006', '跨语言检索', 'TC-M2-012~013'],
        ['M2-REQ-007', '检索准确率度量', 'TC-M2-003, TC-M2-014'],
        ['M3-REQ-001', '聊天助手创建与管理', 'TC-M3-001~003'],
        ['M3-REQ-002', '系统提示词配置', 'TC-M3-004, TC-M3-011'],
        ['M3-REQ-003', '多轮上下文对话', 'TC-M3-007'],
        ['M3-REQ-004', '流式响应输出', 'TC-M3-009'],
        ['M3-REQ-005', '引用溯源', 'TC-M3-010'],
        ['M3-REQ-006', '动态提示模板引擎', 'TC-M3-005, TC-M3-012'],
        ['M3-REQ-007', '结构化约束机制', 'TC-M3-006, TC-M3-008'],
        ['INT-REQ-001', '知识库-会话绑定', 'TC-INT-001~002'],
        ['INT-REQ-002', '统一聊天接口', 'TC-INT-003'],
        ['INT-REQ-003', '流式代理转发', 'TC-INT-004'],
        ['INT-REQ-004', '用户认证集成', 'TC-INT-005~006'],
        ['INT-REQ-005', '离线容器化部署', 'TC-E2E-003'],
        ['INT-REQ-007', '端到端知识问答流程', 'TC-E2E-001'],
        ['INT-REQ-008', '多文档联合推理', 'TC-E2E-002'],
        ['INT-REQ-009', '数据安全隔离', 'TC-E2E-004'],
        ['3.10.4', '有效性（检索响应时间）', 'TC-M2-015'],
    ],
    col_widths=[Cm(2.5), Cm(5), Cm(6)])

add_heading(doc, '8.2  逆向追踪：测试用例→需求', 2)
add_para(doc, '表 18  逆向追踪')

make_table(doc,
    ['测试用例', '测试名称', '对应需求'],
    [
        ['TC-M1-001~004', '知识库CRUD', 'M1-REQ-001'],
        ['TC-M1-005~011', '文档上传', 'M1-REQ-002'],
        ['TC-M1-012~016', '文档解析', 'M1-REQ-003, M1-REQ-004, M1-REQ-008'],
        ['TC-M1-017~020', '分块CRUD', 'M1-REQ-005'],
        ['TC-M1-021', '元数据过滤', 'M1-REQ-006'],
        ['TC-M2-001~004', '向量搜索', 'M2-REQ-001, M2-REQ-007'],
        ['TC-M2-005~007', '混合检索', 'M2-REQ-002'],
        ['TC-M2-008~009', '相似度阈值', 'M2-REQ-003'],
        ['TC-M2-010', '重排序', 'M2-REQ-004'],
        ['TC-M2-011', 'GraphRAG', 'M2-REQ-005'],
        ['TC-M2-012~013', '跨语言检索', 'M2-REQ-006'],
        ['TC-M2-014~015', '准确率与性能', 'M2-REQ-007, 3.10.4'],
        ['TC-M3-001~003', '聊天助手', 'M3-REQ-001'],
        ['TC-M3-004, 011', '提示词约束', 'M3-REQ-002'],
        ['TC-M3-005, 012', '提示模板', 'M3-REQ-006'],
        ['TC-M3-006, 008', '结构化约束', 'M3-REQ-007'],
        ['TC-M3-007', '多轮对话', 'M3-REQ-003'],
        ['TC-M3-009', '流式响应', 'M3-REQ-004'],
        ['TC-M3-010', '引用溯源', 'M3-REQ-005'],
        ['TC-INT-001~002', '会话绑定与隔离', 'INT-REQ-001'],
        ['TC-INT-003', '聊天记录', 'INT-REQ-002'],
        ['TC-INT-004', '流式代理', 'INT-REQ-003'],
        ['TC-INT-005~006', '认证集成', 'INT-REQ-004'],
        ['TC-UI-001~002', '知识库界面', 'M1-REQ-001'],
        ['TC-UI-003', '上传界面', 'M1-REQ-002'],
        ['TC-UI-004', '聊天界面', 'M3-REQ-001, M3-REQ-004'],
        ['TC-UI-005', 'RAG测试界面', 'M2-REQ-007'],
        ['TC-E2E-001', '知识→答案', 'INT-REQ-007'],
        ['TC-E2E-002', '多文档推理', 'INT-REQ-008'],
        ['TC-E2E-003', '离线部署', 'INT-REQ-005'],
        ['TC-E2E-004', '数据隔离', 'INT-REQ-009'],
    ],
    col_widths=[Cm(3), Cm(4.5), Cm(5.5)])

# ═══════════════════════════════════════════════════════════════
# 9 注解
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '9  注解', 1)

add_heading(doc, '9.1  术语表', 2)
add_para(doc, '表 19  术语表')

make_table(doc,
    ['术语', '定义'],
    [
        ['RAG', '检索增强生成（Retrieval-Augmented Generation）'],
        ['LLM', '大语言模型（Large Language Model），本系统使用智谱GLM-9B'],
        ['SSE', '服务器发送事件（Server-Sent Events），用于流式响应'],
        ['CRUD', '创建、读取、更新、删除'],
        ['GraphRAG', '基于知识图谱的检索增强生成'],
        ['STAP', '空时自适应处理（Space-Time Adaptive Processing）'],
        ['MTBF', '平均无故障时间（Mean Time Between Failures）'],
        ['MTTR', '平均修复时间（Mean Time To Repair）'],
        ['BIT', '内置测试（Built-In Test）'],
        ['Fixture', 'pytest测试夹具，用于测试前准备和测试后清理'],
        ['Playwright', '微软浏览器自动化测试框架'],
        ['Allure', '测试报告框架'],
    ],
    col_widths=[Cm(3), Cm(11)])

# ═══════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════

out_path = OUT_DIR / 'DARPA智能问答服务工具-系统测试大纲.docx'
doc.save(str(out_path))
print(f'Saved: {out_path}')
print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')

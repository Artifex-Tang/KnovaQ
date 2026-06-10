"""
Convert DARPA智能问答服务工具-软件设计说明书.md to formatted .docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── helpers ──────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=9, font_name='宋体'):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.bold = True

def add_header_row(table, texts, bg_color="4472C4"):
    """Format first row as header."""
    for i, text in enumerate(texts):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, text, bold=True, size=9, font_name='黑体')
        set_cell_shading(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def format_table(table):
    """Apply basic table formatting."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


# ── main converter ───────────────────────────────────────────────────

def convert(md_path: Path, docx_path: Path):
    md_text = md_path.read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()

    # ── page setup ──
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── styles ──
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    heading_styles = {
        1: ('黑体', 22, True),
        2: ('黑体', 16, True),
        3: ('黑体', 14, True),
        4: ('黑体', 12, True),
        5: ('黑体', 10.5, True),
        6: ('黑体', 10.5, True),
    }
    for level, (fname, fsize, fbold) in heading_styles.items():
        sname = f'Heading {level}'
        if sname in doc.styles:
            s = doc.styles[sname]
            s.font.name = fname
            s._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
            s.font.size = Pt(fsize)
            s.font.bold = fbold
            s.font.color.rgb = RGBColor(0, 0, 0)
            s.paragraph_format.space_before = Pt(12)
            s.paragraph_format.space_after = Pt(6)

    # ── state machine ──
    i = 0
    in_code = False
    code_lines = []
    figure_counter = 0
    table_counter = 0

    def add_para(text, style_name='Normal', bold=False, font_size=None, font_name=None, alignment=None):
        p = doc.add_paragraph(style=style_name)
        if alignment:
            p.alignment = alignment
        # handle bold markers **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            if font_size:
                run.font.size = Pt(font_size)
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if bold and not (part.startswith('**') and part.endswith('**')):
                run.bold = True
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        # shading
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F5F5F5"/>')
        pPr.append(shading)

    def parse_table_line(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    def add_md_table(rows):
        nonlocal table_counter
        table_counter += 1
        ncols = len(rows[0])
        tbl = doc.add_table(rows=len(rows), cols=ncols)
        format_table(tbl)
        # header
        for j, text in enumerate(rows[0]):
            set_cell_text(tbl.rows[0].cells[j], text, bold=True, font_name='黑体')
            set_cell_shading(tbl.rows[0].cells[j], "4472C4")
            for p in tbl.rows[0].cells[j].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # data
        for ri in range(1, len(rows)):
            for j, text in enumerate(rows[ri]):
                if j < ncols:
                    set_cell_text(tbl.rows[ri].cells[j], text)
            # zebra stripe
            if ri % 2 == 0:
                for j in range(ncols):
                    set_cell_shading(tbl.rows[ri].cells[j], "F2F2F2")
        doc.add_paragraph()  # spacer

    # ── skip frontmatter before first chapter heading ──
    # find first real heading
    start_idx = 0
    for idx, line in enumerate(lines):
        if re.match(r'^# 1 范围', line):
            start_idx = idx
            break

    # ── add cover info ──
    for _ in range(6):
        doc.add_paragraph()
    add_para('软件设计说明书', font_size=26, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para('DARPA智能问答服务工具', font_size=18, font_name='黑体', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    cover_items = [
        ('系统标识', 'DARPA-IQAS'),
        ('版本号', 'V1.0'),
        ('密级', '内部'),
        ('研究内容', '研究内容四——DARPA智能问答服务工具开发'),
        ('合作单位', '军事科学院军事科学信息研究中心'),
    ]
    for label, value in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ── process lines ──
    i = start_idx
    while i < len(lines):
        line = lines[i]

        # code block start/end
        if line.strip().startswith('```'):
            if in_code:
                # end code block
                code_text = '\n'.join(code_lines)
                add_code_block(code_text)
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # empty line
        if line.strip() == '':
            i += 1
            continue

        # heading
        hm = re.match(r'^(#{1,7})\s+(.+)$', line)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            # strip leading numbering like "5.1.3.1 "
            add_para(text, style_name=f'Heading {level}')
            i += 1
            continue

        # figure reference like "图 1 ..."
        figm = re.match(r'^图\s+(\d+)\s+(.+)$', line)
        if figm and not line.startswith('|'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'图 {figm.group(1)}  {figm.group(2)}')
            run.font.size = Pt(9)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.bold = True
            i += 1
            continue

        # table like "表 8 ..."
        tblm = re.match(r'^表\s+(\d+)\s+(.+)$', line)
        if tblm and not line.startswith('|'):
            p = doc.add_paragraph()
            run = p.add_run(f'表 {tblm.group(1)}  {tblm.group(2)}')
            run.font.size = Pt(9)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.bold = True
            i += 1
            continue

        # table rows
        if line.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                # skip separator row
                if re.match(r'^\|[\s\-:|]+\|$', row_line):
                    i += 1
                    continue
                cells = parse_table_line(row_line)
                table_rows.append(cells)
                i += 1
            if table_rows:
                add_md_table(table_rows)
            continue

        # horizontal rule
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # normal paragraph
        text = line.strip()
        if text:
            add_para(text)
        i += 1

    # ── save ──
    doc.save(str(docx_path))
    print(f'Saved: {docx_path}')
    print(f'  Figures: {figure_counter}, Tables: {table_counter}')


if __name__ == '__main__':
    base = Path(r'E:\ccode\KnovaQ\docs\output')
    md_file = base / 'DARPA智能问答服务工具-软件设计说明书.md'
    docx_file = base / 'DARPA智能问答服务工具-软件设计说明书.docx'
    convert(md_file, docx_file)

# -*- coding: utf-8 -*-
"""
Insert ALL UI screenshots into DARPA智能问答服务工具-软件用户手册.docx
Matches 47 screenshots from screenshot_manual.py + supplement scripts.
"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = Path(r'E:\ccode\KnovaQ\docs\output\DARPA智能问答服务工具-软件用户手册.docx')
IMG = Path(r'E:\ccode\KnovaQ\docs\output\images\screenshots')

# (anchor_text, screenshot_file, caption, description)
SCREENSHOTS = [
    # Login & Homepage
    ('登录凭据', '01-登录页.png', '系统登录页面'),
    ('安装验证通过', '02-首页主界面.png', '系统首页主界面'),
    # Q&A
    ('智能问答主界面', '03-智能问答页.png', '智能问答主界面'),
    ('配置助理列表', '09-配置助理列表.png', '配置助理列表页面'),
    ('新增助理配置弹窗', '10-配置助理-新增弹窗.png', '新增助理配置弹窗'),
    # File views
    ('文件查看列表', '11-文件查看列表.png', '文件查看列表页面'),
    ('文件详情', '12-文件查看-文件详情.png', '文件详情页面'),
    ('文件管理列表', '13-文件管理列表.png', '文件管理列表页面'),
    ('新建文件夹弹窗', '15-文件管理-新建文件夹.png', '新建文件夹弹窗'),
    ('文件分类管理', '16-文件分类管理.png', '文件分类管理页面'),
    ('新增分类弹窗', '17-文件分类-新增弹窗.png', '新增分类弹窗'),
    # Model & KB
    ('模型管理列表', '18-模型管理列表.png', '模型管理列表页面'),
    ('设置默认模型弹窗', '19-模型管理-设置默认模型.png', '设置默认模型弹窗'),
    ('知识库管理列表', '21-知识库管理列表.png', '知识库管理列表页面'),
    ('新增知识库弹窗', '22-知识库管理-新增弹窗.png', '新增知识库弹窗'),
    ('知识库配置详情', '23-知识库管理-配置页.png', '知识库配置详情页面'),
    # System mgmt
    ('用户管理列表', '24-用户管理列表.png', '系统用户管理列表页面'),
    ('新增用户弹窗', '25-用户管理-新增弹窗.png', '新增用户弹窗'),
    ('角色管理列表', '26-角色管理列表.png', '角色管理列表页面'),
    ('新增角色弹窗', '27-角色管理-新增弹窗.png', '新增角色弹窗'),
    ('菜单管理列表', '28-菜单管理列表.png', '菜单管理列表页面'),
    ('部门管理', '37-部门管理页.png', '部门管理页面'),
    ('参数设置列表', '29-参数设置列表.png', '系统参数设置列表页面'),
    ('在线用户监控', '30-在线用户列表.png', '在线用户监控页面'),
    ('服务器监控', '31-服务器监控.png', '服务器监控页面'),
    ('操作日志列表', '32-操作日志列表.png', '操作日志列表页面'),
    ('登录日志', '41-登录日志页.png', '登录日志页面'),
    ('定时任务', '42-定时任务页.png', '定时任务页面'),
    # Additional
    ('知识检索界面', '04-知识检索页.png', '知识检索界面'),
    ('字典管理', '38-字典管理页.png', '字典管理页面'),
    ('通知公告', '39-通知公告页.png', '通知公告页面'),
    ('缓存监控', '43-缓存监控页.png', '缓存监控页面'),
]


def find_para_index(doc, text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i
    return None


def insert_after(doc, para_index, image_path, caption_text, width=4.8):
    para = doc.paragraphs[para_index]
    element = para._element
    body = doc.element.body
    idx = list(body).index(element)

    # Caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_para.add_run(caption_text)
    run.font.size = Pt(9)
    run.font.bold = True
    cap_element = cap_para._element
    body.remove(cap_element)
    body.insert(idx + 1, cap_element)

    # Image
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    img_element = img_para._element
    body.remove(img_element)
    body.insert(idx + 1, img_element)

    print(f'  OK {image_path.name}')


def main():
    print(f'DOCX: {DOCX.name}')
    print(f'IMG: {IMG}')

    # Check existing screenshots
    available = []
    for anchor, fname, caption in SCREENSHOTS:
        path = IMG / fname
        if path.exists():
            available.append((anchor, fname, caption, path))
        else:
            print(f'  SKIP missing: {fname}')

    print(f'\nFound {len(available)} screenshots to insert')

    doc = Document(str(DOCX))
    inserted = 0
    skipped = 0

    for anchor, fname, caption, path in available:
        idx = find_para_index(doc, anchor)
        if idx is None:
            print(f'  NO ANCHOR: "{anchor}" for {fname}')
            skipped += 1
            continue

        try:
            insert_after(doc, idx, path, caption)
            inserted += 1
            doc.save(str(DOCX))
            doc = Document(str(DOCX))
        except Exception as e:
            print(f'  ERR {fname}: {e}')

    doc.save(str(DOCX))
    print(f'\nDone! Inserted={inserted}, Skipped={skipped}')


if __name__ == '__main__':
    main()

"""
Post-process all 5 docx files to insert PNG images and descriptive text.
Finds figure caption paragraphs, inserts image + description before them.

Usage: python insert_images_to_docx.py
"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

IMG_DIR = Path(r'E:\ccode\KnovaQ\docs\output\images')
OUT_DIR = Path(r'E:\ccode\KnovaQ\docs\output')

def find_paragraph_by_text(doc, search_text):
    """Find paragraph index containing search_text."""
    for i, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            return i
    return None

def insert_paragraph_after(paragraph, text='', style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = deepcopy(paragraph._element)
    # Clear content
    for child in list(new_p):
        if child.tag.endswith('}r'):
            new_p.remove(child)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

def insert_image_before_element(body, element, image_path, width_inches=5.2):
    """Insert an image paragraph (centered) before the given XML element."""
    # Create a temporary document to generate the image paragraph
    tmp_doc = Document()
    p = tmp_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    # Insert the paragraph element before the target element
    img_element = p._element
    body.insert(list(body).index(element), img_element)
    return img_element

def insert_body_text_before(body, element, text, style_name=None):
    """Insert a body text paragraph before the given XML element."""
    tmp_doc = Document()
    p = tmp_doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if style_name:
        try:
            p.style = tmp_doc.styles[style_name]
        except:
            pass
    for run in p.runs:
        run.font.size = Pt(10.5)

    text_element = p._element
    body.insert(list(body).index(element), text_element)
    return text_element

def process_docx(docx_path, image_mappings):
    """
    Process a single docx file.
    image_mappings: list of (search_caption_text, image_filename, description_text)
    """
    print(f'\nProcessing: {docx_path.name}')
    doc = Document(str(docx_path))
    body = doc.element.body

    for caption_text, img_file, description in image_mappings:
        img_path = IMG_DIR / img_file
        if not img_path.exists():
            print(f'  WARNING: Image not found: {img_path}')
            continue

        # Find the caption paragraph
        found = False
        for para in doc.paragraphs:
            if caption_text in para.text:
                element = para._element
                # Insert description text before caption
                if description:
                    insert_body_text_before(body, element, description)
                    # Re-find index after insertion
                # Insert image before caption (or before description if added)
                insert_image_before_element(body, element, img_path, width_inches=5.2)
                print(f'  ✓ Inserted: {img_file} before "{caption_text}"')
                found = True
                break

        if not found:
            print(f'  ✗ Caption not found: "{caption_text}"')

    # Save
    doc.save(str(docx_path))
    print(f'  Saved: {docx_path.name}')


# ═══════════════════════════════════════════════════════════════
# Doc1: 软件需求规格说明书
# ═══════════════════════════════════════════════════════════════
doc1_mappings = [
    ('图 1  功能组成图', '功能组成图.png',
     '图 1 展示了DARPA-IQAS系统的功能组成结构。系统采用三级架构设计：顶层为用户交互层，提供Web前端界面；中间层为应用服务层，包含gaisoft-server（Spring Boot）和ragflow-server（RAG引擎）；底层为数据存储层，包括MySQL、Redis、Elasticsearch和MinIO。各层级之间通过标准接口解耦通信。'),
    ('图 2  DARPA-IQAS系统部署关系图', '系统部署关系图.png',
     '图 2 展示了DARPA-IQAS系统的部署架构。系统采用Docker Compose单机编排方式部署在Linux服务器上，所有服务容器共享ragflow桥接网络。Nginx作为统一入口对外暴露80端口，内部通过反向代理分发请求至ragflow-server（:9384）和gaisoft-server（:8088）。数据持久化通过Docker Volume实现。'),
    ('图 3  DARPA-IQAS应用模式图', '三级架构图.png',
     '图 3 展示了DARPA-IQAS系统的B/S应用模式。用户通过Chrome浏览器访问系统，所有请求经Nginx反向代理路由至后端服务。系统支持完全离线运行，LLM推理引擎（智谱GLM-9B）本地部署，无需外网连接即可完成智能问答。'),
    ('图 4  外部接口关系图', '外部接口关系图.png',
     '图 4 展示了DARPA-IQAS系统的外部接口关系。系统对外提供Web管理界面（HTTP :80），通过Nginx统一入口与用户浏览器交互。内部接口包括：gaisoft-server与MySQL的JDBC连接（:3306）、与Redis的缓存连接（:6379）、与ragflow-server的HTTP API调用（:9384），以及ragflow-server与Elasticsearch（:9200）、MinIO（:9000）、LLM推理服务的连接。'),
]

# ═══════════════════════════════════════════════════════════════
# Doc2: 软件设计说明书 (filename: DARPA-IQAS-软件设计说明书.docx)
# ═══════════════════════════════════════════════════════════════
doc2_mappings = [
    ('图 1 DARPA智能问答服务工具输入/输出数据流示意图', '数据流图.png',
     '图 1 展示了系统的主要数据流。用户提问经由gaisoft-server转发至ragflow-server，ragflow执行文档检索和向量相似度匹配后组装Prompt，调用本地LLM生成回答，最终通过SSE流式返回给用户。文档上传流程将文件存储至MinIO，经解析分块后写入Elasticsearch向量索引。'),
    ('图 2 DARPA智能问答服务工具ER图', 'ER图.png',
     '图 2 展示了equipment_iqas数据库的核心实体关系。数据库分为三个业务域：知识库域（kb_session、kb_chat、kb_source_file、kb_source_type）管理会话和文件分类；系统管理域（sys_user、sys_dept、sys_role）提供用户权限支撑；调度任务域（QRTZ_*系列表）提供定时任务调度。'),
    ('图 3 DARPA智能问答服务工具四层技术体系结构', '四层技术架构图.png',
     '图 3 展示了DARPA智能问答服务工具的四层技术体系结构。第一层为表现层（Vue 3 + Nginx），提供SPA前端和反向代理；第二层为应用层（Spring Boot + ragflow API），承载MES业务逻辑和RAG检索增强生成；第三层为引擎层（LLM推理 + Elasticsearch向量检索 + 文档解析），提供AI核心能力；第四层为基础设施层（MySQL + Redis + MinIO），提供持久化存储服务。'),
    ('图 4 DARPA智能问答服务工具部署关系图', '系统部署关系图.png',
     '图 4 展示了DARPA智能问答服务工具的部署架构。所有服务通过Docker Compose编排运行在单台Linux服务器上，共享ragflow桥接网络。Nginx（ragflow内置）作为统一入口对外暴露HTTP :80端口，内部反向代理至ragflow-server（:9384）和gaisoft-server（:8088）。数据层包括MySQL（:3306）、Redis（:6379）、Elasticsearch（:9200）和MinIO（:9000）。'),
    ('图 5 知识入库流程图', '问答流程活动图.png',
     '图 5 展示了知识入库的完整流程。文档上传后进入解析队列，经过版面分析、表格提取、图文识别等深度解析步骤；随后进行智能分块，将文档切分为语义完整的知识片段；最后向量化写入Elasticsearch索引，供后续检索使用。'),
    ('图 6 问答交互流程图', '问答流程活动图.png',
     '图 6 展示了智能问答的完整交互流程。用户发起提问后，gaisoft-server创建会话并转发至ragflow-server；ragflow执行知识库检索，通过Elasticsearch向量相似度匹配获取TopK结果；随后组装Prompt并调用本地LLM生成回答；最终通过SSE流式或非流式方式返回给用户，同时gaisoft-server将对话记录持久化至MySQL。'),
    ('图 10 外部接口示意图', '外部接口关系图.png',
     '图 10 展示了DARPA智能问答服务工具的外部接口关系。系统对外暴露HTTP :80端口，通过Nginx反向代理统一接入用户浏览器请求。内部服务间接口包括：gaisoft-server通过HTTP调用ragflow-server REST API（:9384）进行知识库操作和问答；ragflow-server内部连接Elasticsearch（:9200）进行向量检索、MinIO（:9000）存储文件、本地LLM服务进行推理生成。'),
    ('图 11 内部接口示意图', '内部接口示意图.png',
     '图 11 展示了DARPA智能问答服务工具的内部接口架构。gaisoft-frontend（Vue 3 SPA）通过HTTP/REST调用gaisoft-server（Spring Boot）的各类业务API，包括登录认证、会话管理、对话发送、文件上传和系统管理等接口。gaisoft-server内部通过Spring Bean依赖注入实现各Service层模块间调用，对外通过HTTP REST API调用ragflow-server的知识库、文档、检索和对话补全接口。数据访问层分别连接MySQL（JDBC）、Redis（Lettuce客户端）和Elasticsearch（REST Client）。'),
]

# ═══════════════════════════════════════════════════════════════
# Doc3: 软件用户手册
# ═══════════════════════════════════════════════════════════════
doc3_mappings = [
    ('图 1  软件功能组织图', '功能组成图.png',
     '图 1 展示了DARPA智能问答服务工具的功能组成。系统分为两大功能域：智能问答域（知识库管理、文档上传、检索配置、对话问答、提示词配置）提供核心RAG问答能力；系统管理域（用户权限、角色菜单、部门管理、系统配置）提供基础运维支撑。'),
]

# ═══════════════════════════════════════════════════════════════
# Doc4: 系统测试大纲
# ═══════════════════════════════════════════════════════════════
doc4_mappings = [
    ('软件项', '测试环境架构图.png',
     '图 1 展示了DARPA智能问答服务工具的测试环境架构。测试客户端为Windows测试机，运行Pytest测试框架、Playwright浏览器自动化和requests HTTP客户端，通过HTTP :80端口访问目标系统。目标系统部署在Linux服务器的Docker Compose容器集群中，包含Nginx（反向代理）、ragflow-server（RAG引擎核心）、gaisoft-server（Spring Boot业务后端）、MySQL 8.0（双库存储）、Redis（缓存）、Elasticsearch 8.x（向量检索）、MinIO（对象存储）等容器服务。所有容器共享ragflow桥接网络，通过容器名互相访问。'),
]


if __name__ == '__main__':
    # Process each document
    docs = [
        (OUT_DIR / 'DARPA智能问答服务工具-软件需求规格说明书.docx', doc1_mappings),
        (OUT_DIR / 'DARPA智能问答服务工具-软件设计说明书.docx', doc2_mappings),
        (OUT_DIR / 'DARPA智能问答服务工具-软件用户手册.docx', doc3_mappings),
        (OUT_DIR / 'DARPA智能问答服务工具-系统测试大纲.docx', doc4_mappings),
        # Doc5 测试报告无需图片
    ]

    for docx_path, mappings in docs:
        if docx_path.exists():
            process_docx(docx_path, mappings)
        else:
            print(f'\nSKIP (not found): {docx_path.name}')

    print('\n═══ All docx image insertion complete! ═══')

# -*- coding: utf-8 -*-
"""
Fix DARPA智能问答服务工具-软件用户手册 — v5
基于 v2 修复，输出 v3：
1. 标题编号：段落级 numId=1 + ilvl
2. 图片：删重复居中图注，图片移到"图 X"上面，无图标题的图片补加
3. 表格：Table 4 起格式化（前4个封面/签署表格不动）
4. 新增内容：将安装部署/升级合并到第3章"软件入门"内，删除尾部独立章节
"""
import sys, shutil, copy
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

V2 = Path(r'E:\ccode\KnovaQ\docs\output\DARPA智能问答服务工具-软件用户手册-v2.docx')
V3 = Path(r'E:\ccode\KnovaQ\docs\output\DARPA智能问答服务工具-软件用户手册-v3.docx')
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

LEVEL_MAP = {
    'Heading 1': '0', 'Heading 2': '1', 'Heading 3': '2', 'Heading 4': '3',
    'Heading 5': '4', 'Heading 6': '5',
}


def fix_heading_numbering(doc):
    print('\n=== Fix 1: 标题编号 ===')
    count = 0
    for p in doc.paragraphs:
        if p.style.name not in LEVEL_MAP:
            continue
        ilvl_val = LEVEL_MAP[p.style.name]
        pPr = p._element.find('w:pPr', ns)
        if pPr is None:
            pPr = etree.SubElement(p._element, qn('w:pPr'))
            p._element.insert(0, pPr)
        old = pPr.find('w:numPr', ns)
        if old is not None:
            pPr.remove(old)
        numPr = etree.SubElement(pPr, qn('w:numPr'))
        etree.SubElement(numPr, qn('w:ilvl')).set(qn('w:val'), ilvl_val)
        etree.SubElement(numPr, qn('w:numId')).set(qn('w:val'), '1')
        count += 1
    print(f'  Fixed {count} headings')


def fix_images(doc):
    print('\n=== Fix 2: 图片位置 + 补图标题 ===')
    body = doc.element.body

    # --- Phase A: Remove centered captions after images ---
    children = list(body)
    imgs, ctrs, figs = {}, {}, {}
    for i, child in enumerate(children):
        tag = etree.QName(child).localname
        if tag != 'p': continue
        has_d = len(child.findall('.//w:drawing', ns)) > 0
        texts = ''.join([t.text for t in child.findall('.//w:t', ns) if t.text]).strip()
        jc = child.find('.//w:jc', ns)
        centered = jc is not None and jc.get(qn('w:val')) == 'center'
        if has_d: imgs[i] = True
        elif centered and texts and not texts.startswith('图 '): ctrs[i] = texts
        elif texts.startswith('图 ') and len(texts) > 3: figs[i] = texts

    to_remove = set()
    for img_idx in imgs:
        for off in [1, 2]:
            if img_idx + off in ctrs:
                to_remove.add(img_idx + off)
                break
    print(f'  Remove {len(to_remove)} duplicate captions')
    for idx in sorted(to_remove, reverse=True):
        body.remove(children[idx])

    # --- Phase B: Move images before "图 X" ---
    children = list(body)
    imgs, figs = {}, {}
    for i, child in enumerate(children):
        tag = etree.QName(child).localname
        if tag != 'p': continue
        has_d = len(child.findall('.//w:drawing', ns)) > 0
        texts = ''.join([t.text for t in child.findall('.//w:t', ns) if t.text]).strip()
        if has_d: imgs[i] = True
        elif texts.startswith('图 ') and len(texts) > 3: figs[i] = texts

    moved = 0
    for fig_idx in sorted(figs.keys()):
        fig_text = figs[fig_idx]
        # Already correct?
        ok = False
        for j in range(max(0, fig_idx-2), fig_idx):
            if j in imgs: ok = True; break
        if ok: continue
        # Find nearest image before fig
        target = None
        for j in range(fig_idx-1, max(0, fig_idx-8), -1):
            if j in imgs: target = j; break
        if target is None: continue
        img_el = children[target]
        body.remove(img_el)
        for c in list(body):
            texts = ''.join([t.text for t in c.findall('.//w:t', ns) if t.text]).strip()
            if texts == fig_text:
                c.addprevious(img_el)
                moved += 1
                break
        children = list(body)
        imgs, figs = {}, {}
        for i, child in enumerate(children):
            tag = etree.QName(child).localname
            if tag != 'p': continue
            if len(child.findall('.//w:drawing', ns)) > 0: imgs[i] = True
            texts = ''.join([t.text for t in child.findall('.//w:t', ns) if t.text]).strip()
            if texts.startswith('图 ') and len(texts) > 3: figs[i] = texts

    print(f'  Moved {moved} images before "图 X"')

    # --- Phase C: Add caption for images without "图 X" ---
    children = list(body)
    # Count existing figure numbers
    max_fig = 0
    for i, child in enumerate(children):
        texts = ''.join([t.text for t in child.findall('.//w:t', ns) if t.text]).strip()
        if texts.startswith('图 ') and len(texts) > 3:
            try:
                num = int(texts.split()[1].strip())
                if num > max_fig: max_fig = num
            except: pass

    added_captions = 0
    # Find images without "图 X" after them
    for i, child in enumerate(children):
        tag = etree.QName(child).localname
        if tag != 'p': continue
        if len(child.findall('.//w:drawing', ns)) == 0: continue
        # Check next elements for "图 X"
        has_fig = False
        for j in range(i+1, min(i+4, len(children))):
            c = children[j]
            texts = ''.join([t.text for t in c.findall('.//w:t', ns) if t.text]).strip()
            if texts.startswith('图 ') and len(texts) > 3:
                has_fig = True; break
            style_el = c.find('.//w:pStyle', ns)
            if style_el is not None and 'Heading' in style_el.get(qn('w:val'), ''):
                break
        if has_fig: continue

        # This image needs a caption - get context from previous paragraph
        prev_texts = ''
        if i > 0:
            prev_texts = ''.join([t.text for t in children[i-1].findall('.//w:t', ns) if t.text]).strip()
        # Extract a short description
        caption_text = '系统界面'
        if '登录' in prev_texts or '账号' in prev_texts:
            caption_text = '系统登录页面'
        elif '主界面' in prev_texts or '安装验证' in prev_texts:
            caption_text = '系统首页主界面'
        elif '首页' in prev_texts:
            caption_text = '系统首页'

        max_fig += 1
        full_caption = f'图 {max_fig}  {caption_text}'

        # Create caption paragraph and insert after image
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(full_caption)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # Move to after image
        cap_el = cap_para._element
        body.remove(cap_el)
        child.addnext(cap_el)
        added_captions += 1

    print(f'  Added {added_captions} figure captions')


def fix_tables(doc):
    print('\n=== Fix 3: 表格格式（Table 4 起） ===')
    total = len(doc.tables)
    # Table 0-3 are cover/signing/TOC, Table 4+ are content
    skip = 4
    for t_idx in range(skip, total):
        table = doc.tables[t_idx]
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = etree.SubElement(tbl, qn('w:tblPr'))

        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None: tblPr.remove(existing)

        borders = etree.SubElement(tblPr, qn('w:tblBorders'))
        for name in ['top','left','bottom','right']:
            b = etree.SubElement(borders, qn(f'w:{name}'))
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '12')
            b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')
        for name in ['insideH','insideV']:
            b = etree.SubElement(borders, qn(f'w:{name}'))
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._element.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = etree.SubElement(cell._element, qn('w:tcPr'))
                    cell._element.remove(tcPr); cell._element.insert(0, tcPr)
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is None: vAlign = etree.SubElement(tcPr, qn('w:vAlign'))
                vAlign.set(qn('w:val'), 'center')
                for para in cell.paragraphs:
                    if r_idx == 0 or c_idx == 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if len(table.rows) > 0:
            trPr = table.rows[0]._element.find(qn('w:trPr'))
            if trPr is None:
                trPr = etree.SubElement(table.rows[0]._element, qn('w:trPr'))
                table.rows[0]._element.insert(0, trPr)
            if trPr.find(qn('w:tblHeader')) is None:
                etree.SubElement(trPr, qn('w:tblHeader'))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    print(f'  Formatted Table {skip}..{total-1} ({total-skip} tables)')


def fix_deployment_section(doc):
    """将尾部独立的'系统安装部署'和'系统升级'章节内容合并到第3章'软件入门'"""
    print('\n=== Fix 4: 合并安装部署/升级到第3章 ===')
    body = doc.element.body

    # Find indices of the standalone H1 headings
    install_h1 = None
    upgrade_h1 = None
    paras = doc.paragraphs

    for i, p in enumerate(paras):
        if p.style.name == 'Heading 1':
            if '系统安装部署' in p.text:
                install_h1 = i
            elif '系统升级' in p.text:
                upgrade_h1 = i

    print(f'  系统安装部署 H1 at para {install_h1}')
    print(f'  系统升级 H1 at para {upgrade_h1}')

    if install_h1 is None or upgrade_h1 is None:
        print('  WARNING: standalone sections not found, skipping')
        return

    # Collect content elements between install_h1 and upgrade_h1, and upgrade_h1 to end
    # We need the actual XML elements, not paragraph objects
    body_children = list(body)

    # Find the body elements for these sections
    install_el = paras[install_h1]._element
    upgrade_el = paras[upgrade_h1]._element

    # Collect all elements from install_el to upgrade_el (exclusive), and upgrade_el to end
    install_section = []
    upgrade_section = []
    phase = 'before'
    for child in body_children:
        if child is install_el:
            phase = 'install'
            continue
        if child is upgrade_el:
            phase = 'upgrade'
            continue
        if phase == 'install':
            install_section.append(child)
        elif phase == 'upgrade':
            upgrade_section.append(child)

    print(f'  Install section elements: {len(install_section)}')
    print(f'  Upgrade section elements: {len(upgrade_section)}')

    # Find chapter 3 "软件入门" and its "安装和设置" H3
    # We want to insert after "安装后验证" (H4) and before "设置" (H3)
    # Change the standalone H1/H2 to H3/H4 to fit inside chapter 3
    install_h1_el = install_el
    upgrade_h1_el = upgrade_el

    # Change install H1 to H3 (under 软件入门)
    for pPr in install_h1_el.findall('w:pPr', ns):
        for pStyle in pPr.findall('w:pStyle', ns):
            pStyle.set(qn('w:val'), '3')  # Heading 3
    # Change text
    for t in install_h1_el.findall('.//w:t', ns):
        if '系统安装部署' in (t.text or ''):
            t.text = t.text.replace('系统安装部署', '系统安装部署（完整流程）')

    # Change upgrade H1 to H3
    for pPr in upgrade_h1_el.findall('w:pPr', ns):
        for pStyle in pPr.findall('w:pStyle', ns):
            pStyle.set(qn('w:val'), '3')

    # Change H2 sub-headings to H4
    for el in install_section + upgrade_section:
        for pStyle in el.findall('.//w:pStyle', ns):
            val = pStyle.get(qn('w:val'), '')
            if val == '2':  # Heading 2
                pStyle.set(qn('w:val'), '4')

    # Now move install_el + install_section + upgrade_el + upgrade_section
    # to after "安装后验证" H4 paragraph
    target_para = None
    for p in paras:
        if p.style.name == 'Heading 4' and '安装后验证' in p.text:
            target_para = p
            break

    if target_para is None:
        print('  WARNING: "安装后验证" not found, appending after chapter 3')
        # Find chapter 3 H1 instead
        for p in paras:
            if p.style.name == 'Heading 1' and '软件入门' in p.text:
                target_para = p
                break

    if target_para is None:
        print('  ERROR: No target found, aborting section move')
        return

    target_el = target_para._element

    # Remove all elements from their current position, then insert after target
    all_to_move = [install_h1_el] + install_section + [upgrade_h1_el] + upgrade_section
    for el in all_to_move:
        body.remove(el)

    # Insert in reverse order after target_el (so order is preserved)
    insert_ref = target_el
    for el in all_to_move:
        insert_ref.addnext(el)
        insert_ref = el

    print('  Moved sections into chapter 3 (软件入门)')


def main():
    print(f'Input:  {V2.name}')
    doc = Document(str(V2))
    fix_heading_numbering(doc)
    fix_images(doc)
    fix_tables(doc)
    fix_deployment_section(doc)
    doc.save(str(V3))
    print(f'\nSaved: {V3.name} ({V3.stat().st_size / 1024:.1f} KB)')


if __name__ == '__main__':
    main()

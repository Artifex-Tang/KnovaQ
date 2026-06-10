"""Convert DARPA-IQAS SRS markdown to formatted docx."""
import re, sys
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD_PATH = Path(__file__).parent / 'DARPA智能问答服务工具-软件需求规格说明书.md'
OUT_PATH = Path(__file__).parent / 'DARPA智能问答服务工具-软件需求规格说明书.docx'

# ── helpers ──
def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    el = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(el)

def add_run(para, text, bold=False, size=10.5, font_name='宋体'):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run

def make_table(doc, headers, rows, col_widths=None):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h.strip(), bold=True, size=10)
        set_cell_shading(cell, 'D9E2F3')
    # data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            add_run(p, val.strip(), size=10)
    return table

# ── parse markdown ──
lines = MD_PATH.read_text(encoding='utf-8').split('\n')

doc = Document()
# page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# helper: parse markdown table → (headers, rows)
def parse_md_table(table_lines):
    headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
    rows = []
    for line in table_lines[2:]:  # skip separator
        if not line.strip() or line.strip().startswith('|--'):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return headers, rows

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # skip empty / hr
    if not stripped or stripped == '---':
        i += 1
        continue

    # skip TOC lines
    if stripped.startswith('- ') and i < 55:
        i += 1
        continue

    # code block (diagrams) → monospace paragraph
    if stripped.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1  # skip closing ```
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_lines))
        run.font.size = Pt(7)
        run.font.name = 'Consolas'
        continue

    # table
    if stripped.startswith('|') and '|' in stripped[1:]:
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        # check if it's actually a table (has separator line)
        if len(table_lines) >= 2 and '---' in table_lines[1]:
            headers, rows = parse_md_table(table_lines)
            make_table(doc, headers, rows)
        else:
            for tl in table_lines:
                p = doc.add_paragraph(tl)
        continue

    # heading 1
    if stripped.startswith('# ') and not stripped.startswith('## '):
        text = stripped[2:].strip()
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(16)
        i += 1
        continue

    # heading 2
    if stripped.startswith('## '):
        text = stripped[3:].strip()
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(14)
        i += 1
        continue

    # heading 3
    if stripped.startswith('### '):
        text = stripped[4:].strip()
        p = doc.add_heading(text, level=3)
        for run in p.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(12)
        i += 1
        continue

    # heading 4
    if stripped.startswith('#### '):
        text = stripped[5:].strip()
        p = doc.add_heading(text, level=4)
        for run in p.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(11)
        i += 1
        continue

    # bullet / numbered list
    if stripped.startswith('- **') or stripped.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        text = stripped[2:].strip()
        # handle bold segments inline
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for idx, part in enumerate(parts):
            is_bold = idx % 2 == 1
            add_run(p, part, bold=is_bold, size=10.5)
        i += 1
        continue

    if re.match(r'^\d+\.\s', stripped):
        p = doc.add_paragraph(style='List Number')
        text = re.sub(r'^\d+\.\s+', '', stripped)
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for idx, part in enumerate(parts):
            is_bold = idx % 2 == 1
            add_run(p, part, bold=is_bold, size=10.5)
        i += 1
        continue

    # sub-bullet (indented list item under numbered)
    if stripped.startswith('  - ') or stripped.startswith('  * '):
        p = doc.add_paragraph(style='List Bullet 2')
        text = stripped.strip().lstrip('-*').strip()
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for idx, part in enumerate(parts):
            is_bold = idx % 2 == 1
            add_run(p, part, bold=is_bold, size=10.5)
        i += 1
        continue

    # normal paragraph with inline bold
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', stripped)
    for idx, part in enumerate(parts):
        if not part:
            continue
        is_bold = idx % 2 == 1
        add_run(p, part, bold=is_bold, size=10.5)
    i += 1

doc.save(str(OUT_PATH))
print(f'Saved: {OUT_PATH}')
print(f'Size: {OUT_PATH.stat().st_size / 1024:.1f} KB')

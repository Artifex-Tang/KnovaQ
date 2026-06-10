"""
基于 需求规格说明模板.docx 模板生成 DARPA智能问答服务工具-软件需求规格说明书.docx
保留模板的样式定义，清除模板占位内容，填入实际内容。
"""
import sys, re, copy
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

TPL = Path(r'E:\ccode\KnovaQ\docs\需求规格说明模板.docx')
OUT = Path(r'E:\ccode\KnovaQ\docs\output\DARPA智能问答服务工具-软件需求规格说明书.docx')

doc = Document(str(TPL))

# ── Step 1: Remove all body paragraphs (from index 79 onward) ──
# Template body starts at paragraph 79 (after cover + TOC + figure list)
body = doc.element.body
# Collect XML elements to remove (paragraphs 79+)
paras_to_remove = []
for i, p in enumerate(doc.paragraphs):
    if i >= 79:
        paras_to_remove.append(p._element)

# Also remove tables that are part of template body
tables_to_remove = list(doc.tables)

for elem in paras_to_remove:
    body.remove(elem)
for tbl in tables_to_remove:
    body.remove(tbl._element)

# ── Step 2: Helper functions using template styles ──

def add_h1(text):
    p = doc.add_paragraph(text, style='Heading 1')
    return p

def add_h2(text):
    p = doc.add_paragraph(text, style='Heading 2')
    return p

def add_h3(text):
    p = doc.add_paragraph(text, style='Heading 3')
    return p

def add_h4(text):
    p = doc.add_paragraph(text, style='Heading 4')
    return p

def add_body(text):
    """Add body text paragraph with inline bold support."""
    p = doc.add_paragraph(style='A 正文')
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if idx % 2 == 1:
            run.bold = True
    return p

def add_list1(text):
    p = doc.add_paragraph(style='A 并列项 1级')
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if idx % 2 == 1:
            run.bold = True
    return p

def add_list2(text):
    p = doc.add_paragraph(style='A 并列项 2级')
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if idx % 2 == 1:
            run.bold = True
    return p

def add_list3(text):
    p = doc.add_paragraph(style='A 并列项 3级')
    run = p.add_run(text)
    return p

def add_table_caption(text):
    p = doc.add_paragraph(text, style='A 表名')
    return p

def add_figure_caption(text):
    p = doc.add_paragraph(text, style='A 图名')
    return p

def add_code_block(text):
    """Add a code/diagram block with monospace font."""
    p = doc.add_paragraph(style='A 正文')
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(7.5)
    return p

def make_table(headers, rows):
    """Create table using template's custom table style."""
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.style = doc.styles['A 自定义表格样式']
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h.strip())
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val.strip())
    return t


# ══════════════════════════════════════════════════════
# Step 3: Write actual document content
# ══════════════════════════════════════════════════════

# ── Chapter 1: 范围 ──
add_h1('范围')

add_h2('标识')
add_body('本规格说明书适用的系统名称、标识和版本号如下：')
add_body('名称：DARPA智能问答服务工具；')
add_body('标识：DARPA-IQAS；')
add_body('简称：DARPA问答工具；')
add_body('版本号：V1.0；')
add_body('研究内容：研究内容四——DARPA智能问答服务工具开发；')
add_body('合作单位：军事科学院军事科学信息研究中心。')

add_h2('系统概述')
add_body('DARPA智能问答服务工具（DARPA-IQAS）是研究内容四的核心成果，联合军事科学院军事科学信息研究中心共同开展。')
add_body('本系统面向DARPA国防高级研究计划局相关领域的军事文档智能问答场景，突破多源异构数据整合瓶颈，通过融合结构化知识管理与检索增强生成技术（RAG），打造具备高精度领域适应能力的离线智能问答系统。')
add_body('系统采用"外挂知识库—RAG检索增强—交互式提示"三级架构设计，包含三大核心模块：')
add_list1('M1 外挂知识库模块：对非结构化军事文档进行深度加工与语义化重构，支持PDF/Word/Excel/TXT/Markdown/图片等多格式文档的解析、智能分块与元数据标注。')
add_list1('M2 RAG文档检索增强模块：基于成熟开源框架（ragflow v0.18.0）进行领域适配，构建向量检索、混合检索（向量+关键词）、相似度阈值调节、重排序、知识图谱、跨语言检索等多维度检索能力。')
add_list1('M3 交互式提示词工程模块：通过动态模板引擎与结构化约束机制，实现聊天助手管理、系统提示词配置、多轮上下文对话、流式响应、引用溯源等能力，将用户意图与DARPA文档知识精准对齐。')
add_body('系统具备以下关键特性：')
add_list1('完全离线运行：采用Docker Compose容器化部署，所有服务（含LLM）均在本地运行，无外网依赖；')
add_list1('本地LLM：集成智谱GLM-9B大语言模型，本地部署，保障数据安全；')
add_list1('军事文档适配：针对DARPA相关军事文档进行深度解析与语义化处理；')
add_list1('多文本特征融合：向量语义检索与关键词检索混合，提升检索精度；')
add_list1('三级架构解耦：知识库、检索增强、提示工程各层级独立可扩展。')

add_h2('文档概述')
add_body('本规格说明书定义DARPA智能问答服务工具（DARPA-IQAS V1.0）的功能需求、性能需求、接口需求及质量特性。本说明书是编写《DARPA智能问答服务工具-软件设计说明书》《DARPA智能问答服务工具-系统测试大纲》《DARPA智能问答服务工具-软件用户手册》等文档的依据。')
add_body('本文档主要包含以下内容：')
add_list1('第一章，范围。包括标识、系统概述、文档概述等。')
add_list1('第二章，引用文档。')
add_list1('第三章，需求。包括要求的状态和方式、CSCI能力需求（按三级架构展开）、外部接口需求、内部接口需求、内部数据需求、适应性需求、保密性需求、安全性需求、环境适应性需求、其他质量特性、计算机资源需求、设计和实现约束、人员相关需求、训练相关需求、软件保障需求、包装需求、其他需求及需求的优先顺序和关键性。')
add_list1('第四章，合格性规定。定义每项需求的合格性检验方法。')
add_list1('第五章，需求可追踪性。建立需求到测试用例的正向和逆向追踪关系。')
add_list1('第六章，注释。术语表和缩略语。')
add_body('本文档密级为内部。')

# ── Chapter 2: 引用文档 ──
add_h1('引用文档')
add_body('DARPA智能问答服务工具软件需求规格说明书的引用文档见表 1。')
add_table_caption('表 1  引用文档一览表')
make_table(
    ['序号', '文档编号', '文档名称', '版本', '编写单位'],
    [
        ['1', '—', 'ragflow v0.18.0 官方技术文档', 'v0.18.0', 'Infiniflow'],
        ['2', '—', '智谱GLM-9B 大语言模型技术文档', 'V1.0', '智谱AI'],
        ['3', '—', 'Spring Boot 2.x 参考文档', '2.x', 'Pivotal'],
        ['4', '—', 'Vue 3 官方文档', '3.x', 'Vue.js'],
        ['5', '—', 'Elasticsearch 8.11 官方文档', '8.11', 'Elastic'],
        ['6', '—', 'Docker 官方文档', '24+', 'Docker Inc.'],
        ['7', '—', 'MySQL 8.0 参考手册', '8.0', 'Oracle'],
        ['8', 'GJB 438B', '军用软件开发文档通用要求', '—', '国防科学技术工业委员会'],
        ['9', '—', 'DARPA相关技术规范与文档', '—', 'DARPA'],
    ]
)

# ── Chapter 3: 需求 ──
add_h1('需求')
add_body('本章规定DARPA智能问答服务工具（DARPA-IQAS）的CSCI需求，即作为CSCI验收条件的CSCI特征。每条需求指定项目唯一的标识符以便测试和追踪，并以一种能为验证和确认方法所接受的方式描述。')

add_h2('要求的状态和方式')

add_h3('能力需求')
add_body('DARPA-IQAS系统具备三大核心能力，对应三级架构的三个层级，见表 3-1。')
add_table_caption('表 3-1  三级核心能力一览表')
make_table(
    ['层级', '能力名称', '核心能力描述'],
    [
        ['第一级', '外挂知识库能力', '对非结构化军事文档进行深度加工与语义化重构，实现文档解析、智能分块、元数据标注与知识库管理'],
        ['第二级', 'RAG检索增强能力', '基于成熟框架领域适配，构建向量检索、混合检索、重排序、知识图谱、跨语言等多维度检索能力'],
        ['第三级', '交互式提示能力', '通过动态模板引擎与结构化约束机制，实现用户意图与DARPA文档知识的精准对齐'],
    ]
)
add_body('三级能力自底向上协同工作：第一级为第二级提供结构化知识片段输入，第二级为第三级提供精准的检索结果支撑，第三级将检索结果与用户意图结合，通过LLM生成最终答案。')

add_h3('功能组成')
add_body('DARPA-IQAS系统的功能组成如图 1所示。')
add_code_block(
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│                    DARPA智能问答服务工具                          │\n'
    '│  ┌───────────────────────────────────────────────────────────┐  │\n'
    '│  │            M3 交互式提示词工程模块                          │  │\n'
    '│  │  聊天助手管理 │ 系统提示词配置 │ 多轮对话 │ 引用溯源       │  │\n'
    '│  └──────────────────────────┬────────────────────────────────┘  │\n'
    '│  ┌──────────────────────────▼────────────────────────────────┐  │\n'
    '│  │            M2 RAG文档检索增强模块                           │  │\n'
    '│  │  向量检索 │ 混合检索 │ 相似度阈值 │ 重排序 │ 知识图谱     │  │\n'
    '│  │  跨语言检索 │ 检索准确率评估                                │  │\n'
    '│  └──────────────────────────┬────────────────────────────────┘  │\n'
    '│  ┌──────────────────────────▼────────────────────────────────┐  │\n'
    '│  │            M1 外挂知识库模块                                │  │\n'
    '│  │  知识库CRUD │ 文档上传 │ 文档解析 │ 智能分块              │  │\n'
    '│  │  分块管理 │ 元数据标注 │ 多源异构数据整合                  │  │\n'
    '│  └───────────────────────────────────────────────────────────┘  │\n'
    '│  ┌───────────────────────────────────────────────────────────┐  │\n'
    '│  │  集成与部署：用户认证(JWT) │ 离线Docker部署 │ 数据安全    │  │\n'
    '│  └───────────────────────────────────────────────────────────┘  │\n'
    '└─────────────────────────────────────────────────────────────────┘'
)
add_figure_caption('图 1  功能组成图')

add_h3('软件实体')
add_body('DARPA-IQAS系统包含以下软件实体，见表 3-2。')
add_table_caption('表 3-2  软件实体一览表')
make_table(
    ['序号', '软件实体', '类型', '版本', '用途说明'],
    [
        ['1', 'ragflow引擎', 'RAG核心服务', 'v0.18.0', '提供文档解析、向量索引、混合检索、LLM对话等RAG核心能力'],
        ['2', 'gaisoft-server', '应用服务', 'V1.0', 'Spring Boot后端，提供用户认证、知识库管理、对话管理等业务服务'],
        ['3', 'gaisoft-frontend', '前端界面', 'V1.0', 'Vue 3 + nginx前端，提供Web用户界面'],
        ['4', 'Elasticsearch', '搜索引擎', '8.11.3', '提供向量索引和全文检索能力'],
        ['5', 'MySQL', '关系数据库', '8.0.39', '存储ragflow引擎元数据（rag_flow库）和业务数据（darpa_iqas库）'],
        ['6', 'MinIO', '对象存储', 'RELEASE.2023-12-20', '存储上传的文档文件'],
        ['7', 'Valkey/Redis', '缓存服务', '8', '提供会话缓存和检索缓存'],
        ['8', 'test-runner', '测试框架', 'V1.0', 'pytest + Playwright自动化测试框架，48条测试用例'],
    ]
)

add_h3('部署模式')
add_body('DARPA-IQAS系统采用离线Docker Compose容器化部署模式。所有服务通过docker-compose.yml编排，以Docker容器方式运行。支持以下部署特性：')
add_list1('离线交付：通过offline-save.sh/ps1导出Docker镜像为离线包，目标环境通过offline-load.sh/ps1加载，无需外网连接；')
add_list1('多客户环境隔离：通过docker/projects/<customer>/目录实现每客户独立配置（端口、密码、nginx路由）；')
add_list1('脚本化管理：提供start.sh/ps1、stop.sh/ps1、build-mes.sh/ps1、build-ui.sh/ps1等脚本简化运维操作；')
add_list1('服务自动恢复：所有服务配置restart: unless-stopped，主机重启后自动恢复。')
add_body('DARPA-IQAS系统部署关系如图 2所示。')
add_code_block(
    '┌──────────────────────────────────────────────────────────┐\n'
    '│              Docker Compose 容器编排                       │\n'
    '│  gaisoft-frontend(:8899)  gaisoft-server(:8088)           │\n'
    '│  ragflow-server(:9380/:8070)                              │\n'
    '│  MySQL 8.0(:5455)  Valkey/Redis 8(:6380)                 │\n'
    '│  Elasticsearch 8.11.3(:1200)  MinIO(:9100)               │\n'
    '│  test-runner(pytest+Playwright, profile:test)            │\n'
    '└──────────────────────────────────────────────────────────┘'
)
add_figure_caption('图 2  DARPA-IQAS系统部署关系图')

add_h3('应用模式')
add_body('DARPA-IQAS系统采用B/S（浏览器/服务器）架构，用户通过Web浏览器离线访问系统。应用模式如图 3所示。')
add_body('用户浏览器(离线局域网，http://IP:8899) → gaisoft-frontend(Vue 3 + nginx) → gaisoft-server(Spring Boot :8088) → ragflow引擎(RAG + GLM-9B :9380)。')
add_figure_caption('图 3  DARPA-IQAS应用模式图')

# ── 3.2 CSCI能力需求 ──
add_h2('CSCI能力需求')
add_body('本章按三级架构逐一列出与DARPA-IQAS各能力相关的需求。每条需求带唯一标识符、优先级和合格性方法。')

add_h3('能力标识')
add_body('DARPA-IQAS系统CSCI能力标识见表 3-3。')
add_table_caption('表 3-3  CSCI能力标识表')
make_table(
    ['能力标识', '能力名称', '所属模块', '优先级'],
    [
        ['M1-REQ-001', '知识库创建/删除/修改/查询', 'M1 外挂知识库', '1'],
        ['M1-REQ-002', '非结构化军事文档上传', 'M1 外挂知识库', '1'],
        ['M1-REQ-003', '文档深度解析', 'M1 外挂知识库', '1'],
        ['M1-REQ-004', '智能语义化分块', 'M1 外挂知识库', '1'],
        ['M1-REQ-005', '分块管理与预览', 'M1 外挂知识库', '1'],
        ['M1-REQ-006', '元数据标注与过滤', 'M1 外挂知识库', '1'],
        ['M1-REQ-007', '文档状态监控', 'M1 外挂知识库', '2'],
        ['M1-REQ-008', '多源异构数据整合', 'M1 外挂知识库', '2'],
        ['M2-REQ-001', '向量语义检索', 'M2 RAG检索增强', '1'],
        ['M2-REQ-002', '混合检索', 'M2 RAG检索增强', '1'],
        ['M2-REQ-003', '相似度阈值动态调节', 'M2 RAG检索增强', '1'],
        ['M2-REQ-004', '检索结果重排序', 'M2 RAG检索增强', '2'],
        ['M2-REQ-005', '知识图谱辅助检索', 'M2 RAG检索增强', '2'],
        ['M2-REQ-006', '跨语言检索', 'M2 RAG检索增强', '2'],
        ['M2-REQ-007', '检索准确率度量与评估', 'M2 RAG检索增强', '1'],
        ['M2-REQ-008', '基于成熟框架的领域适配', 'M2 RAG检索增强', '1'],
        ['M3-REQ-001', '聊天助手创建与管理', 'M3 交互式提示', '1'],
        ['M3-REQ-002', '系统提示词配置', 'M3 交互式提示', '1'],
        ['M3-REQ-003', '多轮上下文对话', 'M3 交互式提示', '1'],
        ['M3-REQ-004', '流式响应输出', 'M3 交互式提示', '1'],
        ['M3-REQ-005', '引用溯源', 'M3 交互式提示', '1'],
        ['M3-REQ-006', '动态提示模板引擎', 'M3 交互式提示', '2'],
        ['M3-REQ-007', '结构化约束机制', 'M3 交互式提示', '2'],
        ['M3-REQ-008', '用户意图与知识精准对齐', 'M3 交互式提示', '1'],
        ['INT-REQ-001', '知识库-会话绑定', '集成与部署', '1'],
        ['INT-REQ-002', '统一聊天接口', '集成与部署', '1'],
        ['INT-REQ-003', '流式代理转发', '集成与部署', '1'],
        ['INT-REQ-004', '用户认证集成', '集成与部署', '1'],
        ['INT-REQ-005', '离线容器化部署', '集成与部署', '1'],
        ['INT-REQ-006', '多客户环境隔离', '集成与部署', '2'],
        ['INT-REQ-007', '端到端知识问答流程', '集成与部署', '1'],
        ['INT-REQ-008', '多文档联合推理', '集成与部署', '1'],
        ['INT-REQ-009', '数据安全隔离', '集成与部署', '1'],
        ['INT-REQ-010', '离线交付包制作', '集成与部署', '2'],
    ]
)

# ── 3.2.2 M1 ──
add_h3('第一级——外挂知识库模块（M1）')

# M1-REQ-001
add_h4('M1-REQ-001：知识库创建/删除/修改/查询')
add_body('标识：M1-REQ-001。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应提供知识库的全生命周期管理能力，包括创建、删除、修改和查询操作。')
add_list1('创建知识库：用户可创建知识库并指定名称、分块方法（naive/book/table/paper/law/manual/qna等），创建时返回唯一的知识库ID；')
add_list1('指定嵌入模型：创建知识库时可指定嵌入模型（如BAAI/bge-large-zh-v1.5），用于后续文档向量化；')
add_list1('查询知识库列表：系统应支持列出所有知识库及其基本信息（ID、名称、文档数量、分块方法等）；')
add_list1('删除知识库：支持删除指定知识库，删除时应级联清理关联的文档、分块数据及向量索引；')
add_list1('修改知识库配置：支持修改知识库的名称、分块方法等配置参数。')
add_body('合格性标准：创建知识库返回非空ID且名称与输入一致；列表接口返回有效知识库列表；级联删除后知识库不再出现在列表中。')

# M1-REQ-002
add_h4('M1-REQ-002：非结构化军事文档上传')
add_body('标识：M1-REQ-002。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应支持多种格式的非结构化军事文档上传至指定知识库。')
add_list1('多格式支持：支持PDF、Word(.docx)、Excel(.xlsx)、TXT、Markdown(.md)、图片（PNG/JPG等，支持OCR识别）等文档格式；')
add_list1('大文件上传：支持5MB以上的军事文档上传，上传过程稳定可靠；')
add_list1('批量上传：支持一次性批量上传多份文档（不少于20份），并返回每份文档的上传状态；')
add_list1('上传状态反馈：每份文档上传后应返回唯一的文档ID，用于后续的解析状态追踪。')
add_body('合格性标准：各格式文档上传均返回有效文档ID；5MB以上文件上传成功；批量上传20份文档返回数量一致。')

# M1-REQ-003
add_h4('M1-REQ-003：文档深度解析')
add_body('标识：M1-REQ-003。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应对上传的军事文档进行深度解析，支持多种解析策略以适应不同文档类型。')
add_list1('通用解析（naive）：对普通军事文档按语义边界进行智能分块，解析完成后生成的分块数量大于0；')
add_list1('书籍解析（book）：按章节/段落结构解析书籍类文档，保留层次结构；')
add_list1('表格解析（table）：对Excel等表格文档进行结构化解析，提取单元格内容和行列关系；')
add_list1('论文解析（paper）：对学术论文类文档进行解析，提取摘要、章节、参考文献等结构；')
add_list1('多源异构解析：支持同一知识库内中英文文档、表格文档等不同来源混合解析，解析过程无冲突。')
add_body('合格性标准：各解析策略均能完成解析并返回成功状态；通用解析后分块数量大于0；多源文档混合解析全部成功。')

# M1-REQ-004
add_h4('M1-REQ-004：智能语义化分块')
add_body('标识：M1-REQ-004。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应对解析后的文档内容按语义边界进行智能分块，生成结构化的知识片段。')
add_list1('语义边界分块：按段落、章节、语义转折点等自然边界进行分块，保持语义完整性；')
add_list1('分块内容保留：每个分块保留原文内容、所属文档信息、位置信息；')
add_list1('关键词提取：支持为分块自动提取或手动指定关键词，用于后续检索增强；')
add_list1('自定义分块：支持用户手动添加分块，指定内容和关键词。')
add_body('合格性标准：分块内容完整保留原文信息；手动添加的分块可被正常检索到。')

# M1-REQ-005
add_h4('M1-REQ-005：分块管理与预览')
add_body('标识：M1-REQ-005。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应提供文档分块的增删改查管理能力。')
add_list1('列出分块：支持按文档列出所有分块，展示分块内容、关键词、来源信息；')
add_list1('更新分块：支持修改已有分块的内容和关键词；')
add_list1('删除分块：支持删除指定分块；')
add_list1('分块预览：提供分块内容的在线预览功能，方便知识工程师验证分块质量。')
add_body('合格性标准：分块列表非空且包含正确信息；更新后内容与修改一致；删除后分块不再出现。')

# M1-REQ-006
add_h4('M1-REQ-006：元数据标注与过滤')
add_body('标识：M1-REQ-006。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应支持为文档和分块添加元数据标注，并基于元数据进行检索过滤。')
add_list1('元数据标注：支持为文档添加元数据标签（如文档类型、密级、来源、日期等）；')
add_list1('元数据过滤检索：在检索时支持基于元数据条件过滤，返回符合特定条件的知识片段；')
add_list1('多文档元数据区分：不同文档可设置不同元数据，检索结果能精确匹配过滤条件。')
add_body('合格性标准：元数据标注后可正确存储；基于元数据过滤的检索返回符合条件的分块。')

# M1-REQ-007
add_h4('M1-REQ-007：文档状态监控')
add_body('标识：M1-REQ-007。优先级：2（应该实现）。合格性方法：B-测试。')
add_body('系统应提供文档解析状态的实时监控能力。')
add_list1('解析进度跟踪：展示文档的解析进度状态（等待中、解析中、已完成、失败）；')
add_list1('解析结果统计：展示解析完成后生成的分块数量；')
add_list1('失败原因展示：文档解析失败时展示具体失败原因；')
add_list1('批量状态查询：支持查询知识库内所有文档的解析状态。')
add_body('合格性标准：文档解析过程中状态正确更新；解析完成后展示正确分块数量。')

# M1-REQ-008
add_h4('M1-REQ-008：多源异构数据整合能力')
add_body('标识：M1-REQ-008。优先级：2（应该实现）。合格性方法：B-测试。')
add_body('系统应支持将不同来源、不同格式的军事文档整合到统一的知识库中。')
add_list1('多格式共存：同一知识库中支持PDF、Word、Excel、TXT、Markdown等多种格式文档共存；')
add_list1('多语言共存：同一知识库支持中文和英文文档共存，各自独立解析；')
add_list1('跨格式检索：整合后的多格式文档可被统一检索访问。')
add_body('合格性标准：多格式文档上传至同一知识库后均可正常解析；中英文文档混合解析无冲突。')

# ── 3.2.3 M2 ──
add_h3('第二级——RAG文档检索增强模块（M2）')

# M2-REQ-001
add_h4('M2-REQ-001：向量语义检索')
add_body('标识：M2-REQ-001。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应提供基于向量语义相似度的检索能力，将用户查询与知识库中的文档分块进行语义匹配。')
add_list1('语义相似度计算：对用户查询进行向量化，与知识库分块向量计算相似度，返回语义最相关的结果；')
add_list1('相似度评分：每个检索结果应附带相似度评分，评分范围0~1，越高表示越相关；')
add_list1('领域查询支持：支持DARPA军事领域专业术语的语义检索（如"雷达维护周期""ZBD-2000通信系统频率范围"等）；')
add_list1('无关查询过滤：对与知识库内容无关的查询（如"今天天气怎么样"），应返回空结果或低相似度结果（相似度低于阈值0.5）；')
add_list1('长查询支持：支持200字符以上的长查询文本，仍能返回有效检索结果；')
add_list1('top_k参数：支持配置返回结果数量（top_k），返回相似度最高的k个结果。')
add_body('合格性标准：领域相关查询返回非空结果且相似度大于0.1；无关查询返回空结果或最高相似度低于0.5；长查询返回有效结果。')

# M2-REQ-002
add_h4('M2-REQ-002：混合检索（向量+关键词多特征融合）')
add_body('标识：M2-REQ-002。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应支持向量语义检索与关键词检索的混合模式，融合多种文本特征提升检索精度。')
add_list1('混合检索模式：同时启用向量语义检索和关键词检索，对结果进行融合排序；')
add_list1('权重可配：支持通过vector_similarity_weight参数调节向量检索与关键词检索的权重比例；')
add_list1('关键词增强：支持启用关键词匹配（keyword）增强检索，确保包含关键术语的文档不被遗漏；')
add_list1('多知识库跨库检索：支持在多个知识库中同时检索，汇总结果后统一排序；')
add_list1('top_k参数：支持配置返回结果数量，top_k=50的结果数量应不少于top_k=10的结果数量。')
add_body('合格性标准：混合检索模式返回有效结果；跨知识库检索返回来自多个知识库的结果；top_k参数单调性正确。')

# M2-REQ-003
add_h4('M2-REQ-003：相似度阈值动态调节')
add_body('标识：M2-REQ-003。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应支持通过相似度阈值参数动态控制检索结果的质量。')
add_list1('阈值参数：支持设置相似度阈值（0~1范围），仅返回相似度大于等于阈值的检索结果；')
add_list1('阈值单调性：提高阈值应减少或保持返回结果数量（阈值0.7的结果数量≤阈值0.1的结果数量）；')
add_list1('极端阈值处理：设置高阈值（如0.9）时，仅返回极少量高相关性结果（不超过3条）；')
add_list1('阈值精度：支持0.01粒度的阈值设置。')
add_body('合格性标准：阈值单调性成立；阈值0.9返回结果数量≤3；所有返回结果的相似度均不低于设定阈值。')

# M2-REQ-004
add_h4('M2-REQ-004：检索结果重排序（Reranking）')
add_body('标识：M2-REQ-004。优先级：2（应该实现）。合格性方法：B-测试。')
add_body('系统应支持对检索结果进行二次重排序，提升结果相关性。')
add_list1('重排序模型：支持配置重排序模型（rerank_id），对初次检索结果进行二次精排；')
add_list1('重排序开关：支持启用/禁用重排序功能，对比重排序前后的结果差异。')
add_body('合格性标准：不启用重排序时检索结果正常返回；启用重排序后结果顺序发生变化且相关性提升。')

# M2-REQ-005
add_h4('M2-REQ-005：知识图谱辅助检索')
add_body('标识：M2-REQ-005。优先级：2（应该实现）。合格性方法：B-测试。')
add_body('系统应支持基于知识图谱（GraphRAG）的辅助检索能力。')
add_list1('知识图谱构建：支持对知识库中的文档运行GraphRAG流程，自动构建知识图谱；')
add_list1('图谱增强检索：在检索时启用知识图谱辅助（use_kg=True），利用实体关系提升检索精度；')
add_list1('优雅降级：当知识图谱未构建或不可用时，自动降级为普通检索模式，不影响用户使用。')
add_body('合格性标准：GraphRAG构建流程可正常执行；图谱不可用时系统正常运行，不报错。')

# M2-REQ-006
add_h4('M2-REQ-006：跨语言检索能力')
add_body('标识：M2-REQ-006。优先级：2（应该实现）。合格性方法：B-测试。')
add_body('系统应支持中英文DARPA文档的跨语言检索能力。')
add_list1('中文查英文：使用中文查询能够检索到英文文档中的相关内容（cross_languages=True）；')
add_list1('英文查中文：使用英文查询能够检索到中文文档中的相关内容（cross_languages=True）；')
add_list1('双语知识库：支持在同一知识库中存储中英文文档并进行跨语言检索。')
add_body('合格性标准：中文查询英文文档返回非空结果；英文查询中文文档返回非空结果。')

# M2-REQ-007
add_h4('M2-REQ-007：检索准确率度量与评估')
add_body('标识：M2-REQ-007。优先级：1（必须实现）。合格性方法：B-测试。')
add_body('系统应提供检索准确率的度量与评估能力，确保检索质量满足军事问答需求。')
add_list1('准确率/召回率评估：支持基于标准问答对（QA pairs）评估检索的召回率，top-5结果的召回率不低于60%；')
add_list1('性能基准测试：支持对检索接口进行性能基准测试，100次请求的P95延迟应低于5秒；')
add_list1('检索结果质量：检索返回的分块应包含查询相关的关键词或答案片段。')
add_body('合格性标准：基于标准QA对的top-5召回率≥60%；100次检索请求P95延迟<5秒。')

# M2-REQ-008
add_h4('M2-REQ-008：基于成熟框架的领域适配')
add_body('标识：M2-REQ-008。优先级：1（必须实现）。合格性方法：D-审查。')
add_body('M2模块的检索增强能力基于成熟开源框架ragflow v0.18.0进行领域适配。')
add_list1('框架版本：使用ragflow v0.18.0作为RAG核心引擎，利用其成熟的文档解析、向量索引、混合检索能力；')
add_list1('领域适配：针对DARPA军事文档特点进行分块策略、检索参数、嵌入模型等方面的适配优化；')
add_list1('API兼容：通过ragflow REST API进行交互，保持与框架版本的兼容性；')
add_list1('持续可用：框架提供的核心API（知识库CRUD、文档上传解析、检索、对话）均正常可用。')
add_body('合格性标准：ragflow v0.18.0所有核心API功能正常；针对军事文档的适配参数可配置。')

# ── 3.2.4 M3 ──
add_h3('第三级——交互式提示词工程模块（M3）')

# M3-REQ-001~008
for req_id, req_name, priority, method, desc_brief, bullets, criteria in [
    ('M3-REQ-001', '聊天助手创建与管理', '1', 'B-测试',
     '系统应提供聊天助手的创建、配置和管理能力。',
     ['创建聊天助手：支持创建聊天助手并绑定指定的知识库，创建后返回唯一的助手ID；',
      '助手配置：支持配置助手的名称、关联知识库、系统提示词、温度参数（temperature）等；',
      '更新助手：支持修改已创建助手的配置参数（如系统提示词、温度等）；',
      '删除助手：支持删除不再使用的聊天助手；',
      '会话管理：支持创建、列出和删除与助手的对话会话，每个会话独立维护上下文。'],
     '创建助手返回有效ID且名称一致；更新后配置在后续对话中生效；会话CRUD操作正常。'),
    ('M3-REQ-002', '系统提示词配置', '1', 'B-测试',
     '系统应支持通过系统提示词（System Prompt）对LLM行为进行领域约束。',
     ['领域角色设定：支持设置系统提示词，将LLM行为约束在DARPA军事文档领域；',
      '领域拒绝：当系统提示词设定为军事领域专用时，对非领域问题（如电影推荐）应拒绝回答或引导回领域话题；',
      '知识变量注入：支持在提示词中引用知识库变量，使回答基于检索到的文档知识；',
      '结构化输出：支持通过系统提示词约束LLM输出格式（如JSON结构化输出）。'],
     '军事领域提示词设置后，非领域问题回答包含拒绝或引导性表述；JSON格式约束后输出可被解析为合法JSON。'),
    ('M3-REQ-003', '多轮上下文对话', '1', 'B-测试',
     '系统应支持多轮上下文对话，保持对话历史并实现上下文连贯。',
     ['上下文保持：在同一会话中，后续对话应能引用前文内容，实现上下文连贯；',
      '多轮引用：第N轮对话可引用第1轮对话的内容，实现跨轮次上下文传递；',
      '温度参数控制：支持设置不同的温度参数（temperature 0.0~1.0），控制回答的确定性与多样性；',
      '所有轮次有效：多轮对话中每一轮均应返回非空回答。'],
     '4轮对话中每轮回答非空；第4轮能基于第1轮上下文回答；不同温度参数均能返回有效回答。'),
    ('M3-REQ-004', '流式响应输出', '1', 'B-测试',
     '系统应支持对话响应的流式输出，提升用户交互体验。',
     ['SSE流式输出：支持通过Server-Sent Events（SSE）方式逐token输出回答内容；',
      '流式终止：流式输出应正确终止，发送结束标记；',
      '内容完整性：流式输出拼接后的完整内容与非流式模式的内容一致。'],
     '流式模式返回至少1个数据块；流式输出正确终止；拼接后内容完整有意义。'),
    ('M3-REQ-005', '引用溯源', '1', 'B-测试',
     '系统应支持对回答内容进行引用溯源，将答案定位到源文档的具体位置。',
     ['引用信息：每个回答应附带引用信息，包含chunk_id（源分块ID）、similarity（相似度评分）、doc_name（源文档名称）；',
      '空知识处理：当查询内容不在知识库范围内时，系统应返回合理的提示而非编造答案；',
      '引用准确性：引用的源文档和分块应与回答内容实际相关。'],
     '回答包含有效引用信息；空知识查询返回有效字符串响应，不编造答案。'),
    ('M3-REQ-006', '动态提示模板引擎', '2', 'B-测试',
     '系统应提供动态提示模板引擎，支持OpenAI兼容的API调用方式。',
     ['OpenAI兼容接口：提供与OpenAI API兼容的对话接口（/api/v1/chat/completions），接受messages数组格式的输入；',
      '标准响应格式：响应包含choices或data字段，符合OpenAI API响应格式规范；',
      '多消息支持：支持在messages数组中包含system、user、assistant多种角色的消息。'],
     'OpenAI兼容接口返回包含choices或data的响应；响应格式符合OpenAI API规范。'),
    ('M3-REQ-007', '结构化约束机制', '2', 'B-测试',
     '系统应通过结构化约束机制确保LLM输出的可控性和规范性。',
     ['输出格式约束：支持通过提示词约束LLM以特定格式输出（如JSON、表格等）；',
      '内容范围约束：支持通过系统提示词将回答范围限定在知识库内容内，减少幻觉；',
      '温度参数约束：支持通过temperature参数控制输出的随机性（0.0=确定性，1.0=高随机性）。'],
     '格式约束后输出符合指定格式；知识库外的内容不作为事实回答。'),
    ('M3-REQ-008', '用户意图/提问方式/行为习惯与DARPA文档知识精准对齐', '1', 'A-演示',
     '系统应通过交互式提示工程实现用户意图与DARPA文档知识的精准对齐。',
     ['意图识别：系统应正确理解用户的提问意图，包括专业术语、缩写、模糊描述等；',
      '知识检索匹配：将用户意图转化为精准的检索查询，从知识库获取相关知识片段；',
      '回答生成：基于检索到的知识片段，结合用户提问方式和行为习惯，生成精准的回答；',
      '端到端验证：完整的知识入库→检索增强→回答生成流程应产出正确结果。'],
     '对"雷达探测距离"类问题，回答包含正确事实数据（如"60公里"）；回答内容来源于知识库文档。'),
]:
    add_h4(f'{req_id}：{req_name}')
    add_body(f'标识：{req_id}。优先级：{priority}（{"必须" if priority=="1" else "应该"}实现）。合格性方法：{method}。')
    add_body(desc_brief)
    for b in bullets:
        add_list1(b)
    add_body(f'合格性标准：{criteria}')

# ── 3.2.5 集成与部署需求 ──
add_h3('集成与部署需求')

for req_id, req_name, priority, method, desc_brief, bullets, criteria in [
    ('INT-REQ-001', '知识库-会话绑定（M1↔M3联动）', '1', 'B-测试',
     '系统应支持将知识库与聊天会话绑定，使对话基于指定知识库内容进行回答。',
     ['会话持久化：知识库会话应持久保存，支持列出历史会话；',
      '会话创建：支持为指定聊天助手创建新会话，会话与知识库关联；',
      '多用户隔离：不同用户/助手的会话数据相互隔离，互不干扰。'],
     '会话列表接口返回有效数据；创建新会话后列表更新；不同助手的会话回答内容独立。'),
    ('INT-REQ-002', '统一聊天接口（M2↔M3联动）', '1', 'B-测试',
     '系统应提供统一的聊天接口，将M2检索能力与M3对话能力整合。',
     ['聊天记录查询：支持通过应用服务API查询知识库聊天记录列表；',
      '会话列表查询：支持查询知识库会话列表；',
      '统一响应格式：聊天接口返回统一的JSON格式（包含code、rows/data等字段）。'],
     '聊天记录接口返回code=200或包含有效数据；会话列表接口返回有效列表。'),
    ('INT-REQ-003', '流式代理转发', '1', 'B-测试',
     '应用服务层应支持将ragflow引擎的SSE流式响应代理转发给前端。',
     ['SSE代理：应用服务接收前端的流式请求，代理转发至ragflow引擎的/api/v1/chats/{id}/completions接口；',
      '流式保持：代理转发过程中保持SSE流式特性，逐块转发数据；',
      '数据完整性：转发过程中数据不丢失、不损坏。'],
     '代理转发返回至少1个SSE数据块；数据块包含data事件内容。'),
    ('INT-REQ-004', '用户认证集成（JWT）', '1', 'B-测试',
     '系统应集成用户认证机制，保护API接口访问安全。',
     ['JWT令牌认证：应用服务应实现JWT令牌认证机制，API请求需携带有效令牌；',
      '令牌缓存：认证令牌应支持缓存复用，避免重复认证；',
      'ragflow认证代理：应用服务应代理ragflow引擎的认证请求，统一管理API Key；',
      '登录验证：支持用户名/密码登录，返回有效的用户信息。'],
     '登录接口返回有效用户信息（code=200或包含user字段）；连续两次API调用均成功（令牌缓存生效）。'),
    ('INT-REQ-005', '离线容器化部署', '1', 'B-测试',
     '系统应支持完全离线的Docker Compose容器化部署。',
     ['离线镜像包：支持将所有Docker镜像导出为离线包（offline-save），在目标环境加载（offline-load）；',
      '无外网依赖：系统运行过程中不依赖任何外网连接，包括LLM推理（智谱GLM-9B本地部署）；',
      '服务自愈：所有服务配置restart: unless-stopped，主机重启后自动恢复；',
      '核心服务可达：离线环境中ragflow引擎、应用服务、前端界面均正常响应。'],
     '离线环境中所有服务健康检查通过；ragflow引擎和 应用服务API正常响应。'),
    ('INT-REQ-006', '多客户环境隔离', '2', 'A-演示',
     '系统应支持多客户部署环境的隔离。',
     ['独立配置：每个客户拥有独立的.env配置文件（端口、密码等）；',
      '独立nginx路由：每个客户拥有独立的nginx配置（server_name、路由规则）；',
      '目录隔离：客户配置存放在docker/projects/<customer>/目录下，互不影响。'],
     '不同客户可使用不同端口访问；配置变更不影响其他客户环境。'),
    ('INT-REQ-007', '端到端知识问答流程（M1→M2→M3完整链路）', '1', 'B-测试',
     '系统应支持从知识入库到智能问答的端到端完整流程。',
     ['完整链路：M1（知识库创建→文档上传→文档解析）→M2（检索增强）→M3（聊天助手创建→对话→回答生成）全链路贯通；',
      '结果验证：对"雷达探测距离"类问题，系统应返回包含正确事实数据（如"60公里""探测"等关键词）的回答。'],
     '端到端流程返回非空回答；回答包含预期的事实数据关键词。'),
    ('INT-REQ-008', '多文档联合推理', '1', 'B-测试',
     '系统应支持基于多份不同格式文档的联合推理能力。',
     ['多文档上传：支持将TXT、DOCX、XLSX等不同格式的多份文档上传至同一知识库；',
      '联合推理：系统能综合多份文档中的信息进行推理，回答需要跨文档关联的问题；',
      '回答质量：多文档联合推理的回答应具有实质内容（长度大于30字符）。'],
     '多文档上传并解析成功；联合推理回答长度>30字符且回答涉及多份文档信息。'),
    ('INT-REQ-009', '数据安全隔离', '1', 'B-测试',
     '系统应确保不同知识库之间的数据严格隔离。',
     ['知识库隔离：不同知识库的文档和分块数据完全隔离，A知识库的文档不应出现在B知识库中；',
      '检索隔离：对A知识库的检索不应返回B知识库的内容；',
      '删除隔离：删除A知识库不影响B知识库的数据完整性。'],
     '知识库A有文档而知识库B文档数量为0；各知识库独立管理，无数据泄露。'),
    ('INT-REQ-010', '离线交付包制作', '2', 'A-演示',
     '系统应支持制作离线交付包，便于在无外网环境中部署。',
     ['镜像导出：提供offline-save.sh/ps1脚本，将所有Docker镜像导出为tar文件；',
      '镜像加载：提供offline-load.sh/ps1脚本，在目标环境加载镜像；',
      '版本标记：离线包应包含版本信息，便于追踪。'],
     '离线包包含所有必需Docker镜像；加载后可正常启动所有服务。'),
]:
    add_h4(f'{req_id}：{req_name}')
    add_body(f'标识：{req_id}。优先级：{priority}（{"必须" if priority=="1" else "应该"}实现）。合格性方法：{method}。')
    add_body(desc_brief)
    for b in bullets:
        add_list1(b)
    add_body(f'合格性标准：{criteria}')

# ── 3.3 CSCI外部接口需求 ──
add_h2('CSCI外部接口需求')

add_h3('接口标识和接口图')
add_body('DARPA-IQAS系统的外部接口关系如图 4所示。用户浏览器通过HTTP/HTTPS访问gaisoft-frontend（端口8899），gaisoft-frontend通过REST API访问gaisoft-server（端口8088），gaisoft-server通过ragflow REST API访问ragflow引擎（端口9380）。')
add_figure_caption('图 4  外部接口关系图')

add_body('DARPA-IQAS外部接口需求见表 3-4。')
add_table_caption('表 3-4  外部接口需求表')
make_table(
    ['接口标识', '接口名称', '接口类型', '源实体', '目标实体'],
    [
        ['IF-001', '浏览器-前端接口', 'HTTP/HTTPS', '用户浏览器', 'gaisoft-frontend'],
        ['IF-002', '前端-后端接口', 'REST API', 'gaisoft-frontend', 'gaisoft-server'],
        ['IF-003', '后端-RAG引擎接口', 'REST API', 'gaisoft-server', 'ragflow引擎'],
        ['IF-004', '后端-数据库接口', 'MySQL协议', 'gaisoft-server', 'MySQL'],
        ['IF-005', '后端-缓存接口', 'Redis协议', 'gaisoft-server', 'Valkey/Redis'],
    ]
)

add_h3('ragflow REST API接口（IF-003）')
add_h4('数据元素特性')
add_body('ragflow REST API接口的核心数据元素特性见表 3-5。')
add_table_caption('表 3-5  ragflow API数据元素特性')
make_table(
    ['数据元素', '名称', '类型', '说明'],
    [
        ['dataset_id', '知识库ID', 'string', '唯一标识知识库'],
        ['dataset_name', '知识库名称', 'string', '知识库显示名称'],
        ['chunk_method', '分块方法', 'string', 'naive/book/table/paper等'],
        ['document_id', '文档ID', 'string', '唯一标识文档'],
        ['chunk_id', '分块ID', 'string', '唯一标识分块'],
        ['content', '分块内容', 'string', '文本内容'],
        ['similarity', '相似度', 'float', '0~1范围'],
        ['chat_id', '聊天助手ID', 'string', '唯一标识聊天助手'],
        ['session_id', '会话ID', 'string', '唯一标识会话'],
    ]
)

add_body('ragflow REST API主要端点见表 3-6。')
add_table_caption('表 3-6  ragflow API端点一览表')
make_table(
    ['端点', '方法', '用途', '对应模块'],
    [
        ['/api/v1/datasets', 'POST', '创建知识库', 'M1'],
        ['/api/v1/datasets', 'GET', '列出知识库', 'M1'],
        ['/api/v1/datasets/{id}', 'DELETE', '删除知识库', 'M1'],
        ['/api/v1/datasets/{id}/documents', 'POST', '上传文档', 'M1'],
        ['/api/v1/datasets/{id}/documents', 'GET', '列出文档', 'M1'],
        ['/api/v1/datasets/{id}/chunks', 'POST/GET', '分块管理', 'M1'],
        ['/api/v1/retrieval', 'POST', '知识检索', 'M2'],
        ['/api/v1/chats', 'POST', '创建聊天助手', 'M3'],
        ['/api/v1/chats/{id}', 'PUT/DELETE', '更新/删除助手', 'M3'],
        ['/api/v1/chats/{id}/completions', 'POST', '对话补全', 'M3'],
        ['/api/v1/chats/{id}/sessions', 'POST/GET', '会话管理', 'M3'],
        ['/api/v1/system/healthz', 'GET', '健康检查', '系统'],
    ]
)

add_h3('应用服务 REST API接口（IF-002）')
add_body('应用服务（gaisoft-server）REST API主要端点见表 3-7。')
add_table_caption('表 3-7  应用服务API端点一览表')
make_table(
    ['端点', '方法', '用途', '对应需求'],
    [
        ['/api/getInfo', 'GET', '获取用户信息', 'INT-REQ-004'],
        ['/api/login', 'POST', '用户登录', 'INT-REQ-004'],
        ['/api/kb/sessions', 'GET/POST', '知识库会话管理', 'INT-REQ-001/002'],
        ['/api/kb/chats', 'GET', '知识库聊天记录', 'INT-REQ-002'],
        ['/api/stream/proxy', 'POST', '流式代理转发', 'INT-REQ-003'],
        ['/api/ragflow/common', 'GET', 'ragflow代理', 'INT-REQ-004'],
    ]
)

# ── 3.4 CSCI内部接口需求 ──
add_h2('CSCI内部接口需求')
add_body('DARPA-IQAS系统内部接口主要存在于应用服务层（gaisoft-server）内部，各业务模块通过Spring Boot内部调用进行交互，见表 3-8。')
add_table_caption('表 3-8  CSCI内部接口需求表')
make_table(
    ['内部接口', '源模块', '目标模块', '交互方式', '说明'],
    [
        ['知识库管理接口', 'M1控制器', 'ragflow API客户端', 'HTTP调用', '转发知识库CRUD操作'],
        ['文档管理接口', 'M1控制器', 'ragflow API客户端', 'HTTP调用', '转发文档上传解析操作'],
        ['检索服务接口', 'M2控制器', 'ragflow API客户端', 'HTTP调用', '转发检索请求'],
        ['对话管理接口', 'M3控制器', 'ragflow API客户端', 'HTTP调用', '转发对话请求'],
        ['认证服务接口', '认证模块', 'ragflow API客户端', 'HTTP调用', '管理API Key和认证令牌'],
        ['流式代理接口', 'M3控制器', 'ragflow SSE端点', 'SSE代理', '代理转发流式响应'],
    ]
)

# ── 3.5 CSCI内部数据需求 ──
add_h2('CSCI内部数据需求')
add_body('DARPA-IQAS系统的内部数据存储分为以下几类：')
add_list1('rag_flow数据库（MySQL）：由ragflow引擎自动管理，存储知识库元数据、文档元数据、分块数据、聊天助手配置、会话数据和用户数据；')
add_list1('darpa_iqas数据库（MySQL）：由gaisoft-server管理，存储用户信息、业务配置、操作日志；')
add_list1('Elasticsearch索引数据：存储文档分块的向量表示（用于语义检索）和文本索引（用于关键词检索）；')
add_list1('MinIO对象存储：存储用户上传的PDF/Word/Excel等原始文档文件。')

# ── 3.6 适应性需求 ──
add_h2('适应性需求')
add_list1('离线环境运行：系统在完全断开外网的环境中正常运行，包括LLM推理、文档解析、检索问答等所有功能；')
add_list1('端口可配置：所有服务端口通过.env文件配置，可根据部署环境调整（避免端口冲突）；')
add_list1('多客户环境：通过docker/projects/<customer>/目录支持多客户独立配置和部署；')
add_list1('资源适配：Elasticsearch内存限制（MEM_LIMIT）可根据服务器配置调整，默认8GB；')
add_list1('数据卷持久化：所有数据通过Docker卷持久化，容器重建不丢失数据；')
add_list1('配置热更新：后端jar包和前端html支持脚本化更新（build-mes.sh/ps1、build-ui.sh/ps1），无需重建镜像。')

# ── 3.7 保密性需求 ──
add_h2('保密性(Security)需求')
add_list1('军事文档数据隔离：不同知识库之间的文档数据严格隔离，防止越权访问（INT-REQ-009）；')
add_list1('用户认证与访问控制：通过JWT令牌实现用户认证，未认证请求被拒绝（INT-REQ-004）；')
add_list1('API Key保护：ragflow API Key由应用服务层统一管理，前端不直接暴露API Key；')
add_list1('离线无外网泄露风险：系统完全离线运行，不存在数据外传至外网的风险（INT-REQ-005）；')
add_list1('密码保护：系统密码通过环境变量配置，不硬编码在代码中；')
add_list1('数据库访问控制：MySQL、Redis、Elasticsearch仅开放内部端口，通过Docker网络隔离外部访问。')

# ── 3.8 安全性需求 ──
add_h2('安全性(Safety)需求')
add_list1('文件上传安全校验：对上传的文档文件进行格式校验，拒绝非法文件类型；')
add_list1('SQL注入防护：应用服务层采用参数化查询，防止SQL注入攻击；')
add_list1('XSS防护：前端界面应对用户输入进行转义处理，防止跨站脚本攻击；')
add_list1('输入内容过滤：对用户输入的查询内容进行基本校验，过滤恶意内容；')
add_list1('服务异常处理：各服务具备异常捕获和优雅降级能力，避免因单个请求导致服务崩溃；')
add_list1('资源限制：通过Docker资源限制（如ES内存限制）防止单一服务占用过多系统资源。')

# ── 3.9 环境适应性需求 ──
add_h2('CSCI环境适应性需求')
add_list1('离线Linux部署：支持在无外网连接的Linux服务器上部署运行；')
add_list1('Docker环境要求：需要Docker 24+和Docker Compose v2环境；')
add_list1('最低硬件配置：8核CPU、16GB内存、500GB硬盘；')
add_list1('网络环境：局域网TCP/IP环境，HTTP协议通信；')
add_list1('无特殊环境依赖：除Docker外不需要安装其他软件依赖。')

# ── 3.10 其他质量特性 ──
add_h2('其他质量特性')

add_h3('功能性要求')
add_list1('完整性：系统应实现本规格说明书规定的全部优先级1的功能需求；')
add_list1('准确性：检索增强模块top-5召回率不低于60%（M2-REQ-007）；问答系统应基于检索到的知识生成准确回答，不编造不存在的信息（M3-REQ-005）；')
add_list1('互操作性：系统应支持与OpenAI兼容的API调用方式（M3-REQ-006），支持多格式文档输入（M1-REQ-002），支持跨语言检索（M2-REQ-006）；')
add_list1('一致性：系统各模块间的数据流转应保持一致；')
add_list1('安全性：系统应具备用户认证（JWT）、数据隔离、API Key保护等安全机制。')

add_h3('可靠性要求')
add_list1('容错性：系统应具备异常处理能力，单个请求的异常不应导致整个服务崩溃；')
add_list1('自动恢复：所有Docker服务配置restart: unless-stopped策略，服务异常退出或主机重启后自动恢复；')
add_list1('健康检查：ragflow引擎、MySQL、Redis等核心服务配置Docker健康检查，确保服务可用性；')
add_list1('数据持久化：所有关键数据通过Docker卷持久化存储；')
add_list1('异常兜底：知识图谱不可用时自动降级为普通检索（M2-REQ-005）。')

add_h3('易用性要求')
add_list1('可操作性：系统提供Web界面操作，用户通过浏览器即可完成全部操作；')
add_list1('可理解性：界面功能应清晰明确，操作流程符合用户认知习惯；')
add_list1('反馈性：长时间操作（如文档解析）应提供进度状态反馈。')

add_h3('有效性要求')
add_list1('时效性：检索响应时间P95<5秒（100次请求基准测试）；流式输出首token延迟<3秒；')
add_list1('资源效率：Elasticsearch内存使用不超过配置限制（默认8GB）；Redis缓存使用不超过128MB。')

add_h3('可维护性要求')
add_list1('模块化设计：三级架构各模块独立，修改某一层级不影响其他层级；')
add_list1('Docker化部署：所有服务容器化，便于升级、回滚和扩展；')
add_list1('脚本化管理：提供start/stop/build-mes/build-ui等脚本，简化运维操作；')
add_list1('日志管理：各服务日志通过Docker卷挂载持久化，便于问题排查；')
add_list1('配置集中化：通过.env文件集中管理配置参数。')

add_h3('可移植性要求')
add_list1('适应性：系统可在任何安装Docker 24+的Linux服务器上运行；')
add_list1('安装性：通过离线镜像包加载+脚本启动的方式部署，安装过程简单；')
add_list1('遵循性：遵循Docker容器化标准，符合行业主流部署实践；')
add_list1('可替换性：核心组件（如LLM模型、嵌入模型）可替换升级，不影响整体架构。')

add_h3('保障性要求')
add_list1('计划性：提供完整的文档体系（需求规格、设计说明、用户手册、测试大纲、测试报告）；')
add_list1('全面性：48条自动化测试用例覆盖全部核心功能；')
add_list1('离线更新支持：支持通过脚本更新后端jar包和前端html。')

add_h3('可测试性要求')
add_list1('自动化测试：提供基于pytest的自动化测试框架，覆盖48条测试用例；')
add_list1('UI自动化：提供基于Playwright的UI自动化测试，覆盖4个核心界面；')
add_list1('测试容器化：测试框架通过test-runner容器运行，与生产环境隔离；')
add_list1('测试报告：生成Allure格式测试报告；')
add_list1('测试数据工厂：提供测试数据生成工具，自动生成各格式测试文档。')

# ── 3.11 计算机资源需求 ──
add_h2('计算机资源需求')

add_h3('计算机硬件需求')
add_table_caption('表 3-9  计算机硬件需求表')
make_table(
    ['资源项', '最低配置', '推荐配置'],
    [
        ['CPU', '4核 x86_64', '8核 x86_64'],
        ['内存', '16GB', '32GB'],
        ['硬盘', '200GB（SSD）', '500GB（SSD）'],
        ['网络', '局域网', '千兆局域网'],
    ]
)

add_h3('计算机硬件资源使用需求')
add_table_caption('表 3-10  硬件资源使用需求表')
make_table(
    ['资源', '服务', '使用限制', '说明'],
    [
        ['内存', 'Elasticsearch', '≤8GB', 'MEM_LIMIT配置，可通过.env调整'],
        ['内存', 'Redis', '≤128MB', 'maxmemory配置，LRU淘汰策略'],
        ['存储', 'MySQL', '动态增长', '建议预留50GB'],
        ['存储', 'MinIO', '动态增长', '取决于文档数量和大小'],
    ]
)

add_h3('计算机软件需求')
add_table_caption('表 3-11  计算机软件需求表')
make_table(
    ['软件项', '版本要求', '用途'],
    [
        ['操作系统', 'Linux（推荐Ubuntu 20.04+）', '服务器操作系统'],
        ['Docker Engine', '24.0+', '容器运行时'],
        ['Docker Compose', 'v2', '容器编排'],
        ['Web浏览器', 'Chrome/Firefox/Edge', '用户访问界面'],
    ]
)

add_h3('计算机通信需求')
add_list1('网络类型：局域网TCP/IP；')
add_list1('通信协议：HTTP/HTTPS；')
add_list1('通信接口：HTTP REST API；')
add_list1('运行时间：7×24小时持续运行；')
add_list1('数据传输内容：JSON格式数据、文档文件、SSE流式数据。')

# ── 3.12 设计和实现约束 ──
add_h2('设计和实现约束')
add_list1('离线运行约束：系统必须在完全离线环境中运行，所有依赖（含LLM模型）必须在部署前内置；')
add_list1('框架约束：RAG核心能力基于ragflow v0.18.0，通过其REST API进行集成，不修改ragflow源代码；')
add_list1('容器化约束：所有服务必须以Docker容器方式运行，通过Docker Compose编排管理；')
add_list1('军事文档兼容：文档解析引擎必须兼容军事领域常见的文档格式和排版特点；')
add_list1('LLM本地部署：使用智谱GLM-9B本地部署，不依赖外部LLM API服务；')
add_list1('嵌入模型：默认使用BAAI/bge-large-zh-v1.5中文嵌入模型，支持替换；')
add_list1('前端技术栈：使用Vue 3框架，通过nginx提供静态文件服务；')
add_list1('后端技术栈：使用Spring Boot 2.x框架；')
add_list1('数据库约束：使用MySQL 8.0存储结构化数据，Elasticsearch 8.11.3存储向量索引和全文索引。')

# ── 3.13 人员相关需求 ──
add_h2('人员相关需求')
add_table_caption('表 3-12  人员角色表')
make_table(
    ['角色', '职责', '技能要求', '培训需求'],
    [
        ['系统管理员', '系统部署、配置、运维管理', 'Linux基础、Docker操作', 'Docker部署培训、故障排查培训'],
        ['知识工程师', '知识库管理、文档上传、检索参数调优', '基本计算机操作能力', '系统操作培训（1天）'],
        ['普通用户', '智能问答对话、知识检索', '基本浏览器操作能力', '系统使用培训（半天）'],
    ]
)

# ── 3.14 训练相关需求 ──
add_h2('训练相关需求')
add_list1('系统管理员培训：提供Docker容器化部署培训，包括镜像加载、服务启停、配置修改、故障排查等；')
add_list1('知识工程师培训：提供知识库管理培训，包括文档上传、解析监控、分块管理、检索参数调优等；')
add_list1('普通用户培训：提供系统使用培训，包括智能问答操作、多轮对话、引用溯源查看等；')
add_list1('培训文档：提供《DARPA智能问答服务工具-软件用户手册》作为培训教材。')

# ── 3.15 软件保障需求 ──
add_h2('软件保障需求')
add_list1('离线更新：支持通过build-mes.sh/ps1更新后端jar包、通过build-ui.sh/ps1更新前端html，无需重建Docker镜像；')
add_list1('版本迭代：支持通过替换Docker镜像实现版本升级；')
add_list1('数据备份：支持通过Docker卷备份和MySQL导出进行数据备份；')
add_list1('故障恢复：通过Docker restart策略实现自动故障恢复，通过日志定位问题；')
add_list1('文档支持：提供完整的技术文档体系。')

# ── 3.16 包装需求 ──
add_h2('包装需求')
add_list1('离线交付包：通过offline-save.sh/ps1导出Docker镜像为tar文件，作为离线交付包；')
add_list1('交付内容：交付包应包含所有Docker镜像、部署脚本、配置文件模板和文档；')
add_list1('存储介质：交付包大小应适合存放于移动硬盘等存储介质；')
add_list1('安装脚本：提供offline-load.sh/ps1脚本用于目标环境加载镜像。')

# ── 3.17 其他需求 ──
add_h2('其他需求')
add_body('无其他未覆盖的特殊需求。')

# ── 3.18 需求的优先顺序和关键性 ──
add_h2('需求的优先顺序和关键性')
add_body('需求的优先次序分为1、2两种，优先级别越高表明该项需求是要优先实现的。优先级别为1的需求是软件的基本需求，必须在交付版本中实现；优先级别为2的需求是增强需求，提升系统完整性和用户体验。')
add_body('DARPA-IQAS需求的优先次序和关键程度见表 3-13。')
add_table_caption('表 3-13  需求的优先次序和关键程度')

# Build the full priority table
priority_rows = []
for prefix in ['M1-REQ', 'M2-REQ', 'M3-REQ', 'INT-REQ']:
    for num in range(1, 20):
        rid = f'{prefix}-{num:03d}'
        # lookup from our data - simplified inline
        if rid in [
            'M1-REQ-007','M1-REQ-008',
            'M2-REQ-004','M2-REQ-005','M2-REQ-006',
            'M3-REQ-006','M3-REQ-007',
            'INT-REQ-006','INT-REQ-010'
        ]:
            priority_rows.append([rid, '2', '一般', 'B-测试' if rid != 'M3-REQ-006' and rid != 'M3-REQ-007' else 'B-测试'])
        elif rid <= 'M2-REQ-008' or rid <= 'M3-REQ-008' or rid <= 'INT-REQ-010' or rid <= 'M1-REQ-008':
            # Check if valid req
            valid = True
            if prefix == 'M1-REQ' and num > 8: valid = False
            if prefix == 'M2-REQ' and num > 8: valid = False
            if prefix == 'M3-REQ' and num > 8: valid = False
            if prefix == 'INT-REQ' and num > 10: valid = False
            if valid:
                method = 'D-审查' if rid == 'M2-REQ-008' else ('A-演示' if rid in ['M3-REQ-008','INT-REQ-006','INT-REQ-010'] else 'B-测试')
                priority_rows.append([rid, '1', '关键', method])

make_table(
    ['需求标识', '优先级', '关键性', '合格性方法'],
    priority_rows
)

# ── Chapter 4: 合格性规定 ──
add_h1('合格性规定')
add_body('DARPA-IQAS系统采用以下合格性方法验证需求的满足程度：')
add_list1('演示（A）：通过操作演示，使用目视观察或专用测试设备，采集被证实的能力直接进行比较，判别CSCI或CSCI一部分是否合格；')
add_list1('测试（B）：使用专用测试设备、仪器或软件，对被测CSCI或CSCI一部分采集数据，验证功能是否正常使用；')
add_list1('分析（C）：使用合格的性约性检验方法、近似计算、仿真、接口端到端测量等手段验证；')
add_list1('审查（D）：对CSCI代码、文档等进行目视审查。')
add_body('DARPA-IQAS合格性规定见表 4-1。')
add_table_caption('表 4-1  合格性规定')
make_table(
    ['合格性方法', '适用需求', '说明'],
    [
        ['A-演示', 'M3-REQ-008、INT-REQ-006、INT-REQ-010', '端到端流程演示、多客户隔离演示、离线交付包演示'],
        ['B-测试', 'M1-REQ-001~008、M2-REQ-001~007、M3-REQ-001~007、INT-REQ-001~005/007~009', '48条自动化测试用例覆盖'],
        ['C-分析', '3.10质量特性', '性能指标分析、可靠性分析'],
        ['D-审查', 'M2-REQ-008、3.5内部数据需求', '框架适配审查、数据结构审查'],
    ]
)

# ── Chapter 5: 需求可追踪性 ──
add_h1('需求可追踪性')

add_body('本章描述：')
add_list1('从本规格说明中的每一个CSCI需求，到所涉及的系统/子系统需求的可追踪性；')
add_list1('从已分配给本CSCI的每一个系统/子系统需求，到所涉及的CSCI需求的可追踪性。')

add_body('DARPA-IQAS对需求的正向追踪见表 5-1。')
add_table_caption('表 5-1  需求的正向追踪性')
make_table(
    ['需求标识', '需求名称', '测试用例', '测试文件'],
    [
        ['M1-REQ-001', '知识库CRUD', 'TC-M1-001~004', 'test_dataset_crud.py'],
        ['M1-REQ-002', '文档上传', 'TC-M1-005~011', 'test_document_upload.py'],
        ['M1-REQ-003', '文档解析', 'TC-M1-012~016', 'test_document_parsing.py'],
        ['M1-REQ-004', '智能分块', 'TC-M1-017~019', 'test_chunk_management.py'],
        ['M1-REQ-005', '分块管理', 'TC-M1-020~022', 'test_chunk_management.py'],
        ['M1-REQ-006', '元数据过滤', 'TC-M1-023', 'test_metadata_filter.py'],
        ['M1-REQ-007', '文档状态监控', 'TC-M1-012~016', 'test_document_parsing.py'],
        ['M1-REQ-008', '多源数据整合', 'TC-M1-016', 'test_document_parsing.py'],
        ['M2-REQ-001', '向量检索', 'TC-M2-001~004', 'test_vector_search.py'],
        ['M2-REQ-002', '混合检索', 'TC-M2-005~007', 'test_hybrid_search.py'],
        ['M2-REQ-003', '相似度阈值', 'TC-M2-008~009', 'test_similarity_threshold.py'],
        ['M2-REQ-004', '重排序', 'TC-M2-010', 'test_reranking.py'],
        ['M2-REQ-005', '知识图谱', 'TC-M2-011', 'test_knowledge_graph.py'],
        ['M2-REQ-006', '跨语言检索', 'TC-M2-012~013', 'test_cross_language.py'],
        ['M2-REQ-007', '检索准确率', 'TC-M2-014~015', 'test_retrieval_accuracy.py'],
        ['M2-REQ-008', '框架领域适配', '—', 'D-审查'],
        ['M3-REQ-001', '聊天助手管理', 'TC-M3-001~003', 'test_chat_assistant.py'],
        ['M3-REQ-002', '系统提示词', 'TC-M3-004~006', 'test_system_prompt.py'],
        ['M3-REQ-003', '多轮对话', 'TC-M3-007~008', 'test_multi_turn_dialog.py'],
        ['M3-REQ-004', '流式响应', 'TC-M3-009', 'test_streaming_response.py'],
        ['M3-REQ-005', '引用溯源', 'TC-M3-010~011', 'test_reference_citation.py'],
        ['M3-REQ-006', '提示模板引擎', 'TC-M3-012', 'test_prompt_template_engine.py'],
        ['M3-REQ-007', '结构化约束', 'TC-M3-004/009', 'test_system_prompt.py'],
        ['M3-REQ-008', '意图知识对齐', 'TC-E2E-001', 'test_e2e_knowledge_to_answer.py'],
        ['INT-REQ-001', '知识库会话绑定', 'TC-INT-001~002', 'test_kb_session.py'],
        ['INT-REQ-002', '统一聊天接口', 'TC-INT-003~004', 'test_kb_chat.py'],
        ['INT-REQ-003', '流式代理转发', 'TC-INT-005', 'test_stream_proxy.py'],
        ['INT-REQ-004', '用户认证集成', 'TC-INT-006~007', 'test_auth_integration.py'],
        ['INT-REQ-005', '离线部署', 'TC-E2E-003', 'test_e2e_offline_deployment.py'],
        ['INT-REQ-006', '多客户隔离', '—', 'A-演示'],
        ['INT-REQ-007', '端到端问答', 'TC-E2E-001', 'test_e2e_knowledge_to_answer.py'],
        ['INT-REQ-008', '多文档推理', 'TC-E2E-002', 'test_e2e_multi_doc_reasoning.py'],
        ['INT-REQ-009', '数据安全隔离', 'TC-E2E-004', 'test_e2e_data_security.py'],
        ['INT-REQ-010', '离线交付包', '—', 'A-演示'],
    ]
)

add_body('DARPA-IQAS对需求的逆向追踪见表 5-2。')
add_table_caption('表 5-2  需求的逆向追踪性')
make_table(
    ['测试用例', '测试内容', '覆盖需求'],
    [
        ['TC-M1-001~004', '知识库CRUD', 'M1-REQ-001'],
        ['TC-M1-005~011', '文档上传（PDF/DOCX/XLSX/TXT/MD/大文件/批量）', 'M1-REQ-002'],
        ['TC-M1-012~016', '文档解析（naive/book/table/paper/多源）', 'M1-REQ-003/007/008'],
        ['TC-M1-017~022', '分块管理（增删改查）', 'M1-REQ-004/005'],
        ['TC-M1-023', '元数据过滤', 'M1-REQ-006'],
        ['TC-M2-001~004', '向量语义检索', 'M2-REQ-001'],
        ['TC-M2-005~007', '混合检索', 'M2-REQ-002'],
        ['TC-M2-008~009', '相似度阈值', 'M2-REQ-003'],
        ['TC-M2-010', '重排序', 'M2-REQ-004'],
        ['TC-M2-011', 'GraphRAG检索', 'M2-REQ-005'],
        ['TC-M2-012~013', '跨语言检索', 'M2-REQ-006'],
        ['TC-M2-014~015', '准确率/性能', 'M2-REQ-007'],
        ['TC-M3-001~003', '聊天助手CRUD', 'M3-REQ-001'],
        ['TC-M3-004~006', '系统提示词', 'M3-REQ-002/007'],
        ['TC-M3-007~008', '多轮对话', 'M3-REQ-003'],
        ['TC-M3-009', '流式SSE输出', 'M3-REQ-004'],
        ['TC-M3-010~011', '引用溯源', 'M3-REQ-005'],
        ['TC-M3-012', 'OpenAI兼容接口', 'M3-REQ-006'],
        ['TC-INT-001~002', '会话持久化/隔离', 'INT-REQ-001'],
        ['TC-INT-003~004', '聊天/会话列表', 'INT-REQ-002'],
        ['TC-INT-005', '流式代理', 'INT-REQ-003'],
        ['TC-INT-006~007', '认证/登录', 'INT-REQ-004'],
        ['TC-UI-001~004', 'UI界面测试', 'M1/M2/M3'],
        ['TC-E2E-001', '知识到答案全流程', 'INT-REQ-007/M3-REQ-008'],
        ['TC-E2E-002', '多文档推理', 'INT-REQ-008'],
        ['TC-E2E-003', '离线部署验证', 'INT-REQ-005'],
        ['TC-E2E-004', '数据安全隔离', 'INT-REQ-009'],
    ]
)

# ── Chapter 6: 注释 ──
add_h1('注释')
add_body('本章列出本文档涉及的术语和缩略语，见表 6-1。')
add_table_caption('表 6-1  术语和缩略语表')
make_table(
    ['术语/缩略语', '全称', '含义'],
    [
        ['RAG', 'Retrieval-Augmented Generation', '检索增强生成'],
        ['LLM', 'Large Language Model', '大语言模型'],
        ['DARPA', 'Defense Advanced Research Projects Agency', '美国国防高级研究计划局'],
        ['DARPA-IQAS', 'DARPA Intelligent Question Answering Service', 'DARPA智能问答服务工具标识'],
        ['CSCI', 'Computer Software Configuration Item', '计算机软件配置项'],
        ['SSE', 'Server-Sent Events', '服务器发送事件'],
        ['JWT', 'JSON Web Token', 'JSON网络令牌'],
        ['CRUD', 'Create, Read, Update, Delete', '创建、读取、更新、删除'],
        ['GraphRAG', 'Graph-based RAG', '基于知识图谱的检索增强生成'],
        ['OCR', 'Optical Character Recognition', '光学字符识别'],
        ['Embedding', '—', '将文本转换为高维向量表示的过程'],
        ['Top-K', '—', '检索结果中相似度最高的K个结果'],
        ['P95', '95th Percentile', '第95百分位数'],
        ['QA', 'Question and Answer', '问答对'],
        ['MCP', 'Model Context Protocol', '模型上下文协议'],
        ['BGE', 'BAAI General Embedding', '北京智源通用嵌入模型'],
        ['GLM-9B', 'General Language Model 9B', '智谱AI通用语言模型（90亿参数）'],
    ]
)

# ── Save ──
doc.save(str(OUT))
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size / 1024:.1f} KB')
